from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate
from django.contrib import messages
from django.http import JsonResponse
from django.db.models import Count
from django.core.files.base import ContentFile
from collections import defaultdict
import json
import base64
import uuid
from datetime import date, timedelta
from django.utils import timezone
from .models import (User, Course, CLO, PLO, Assessment, Question, SubQuestion, SubQuestionGrade,
                     Enrollment, Submission, QuestionGrade, Announcement, StudyMaterial, Notification,
                     CLOActionPlan, PLOActionPlan, Section)
from .notifications import (notify_grade_released, notify_new_assignment, notify_new_material, notify_announcement)
from .grace_period import check_submission_window, apply_late_deduction, recalculate_final_score


#Home Redirect 
def home(request):
    if not request.user.is_authenticated:
        return render(request, 'homepage.html')
    if request.user.role == 'faculty' or request.user.is_superuser:
        return redirect('faculty_dashboard')
    elif request.user.role == 'student':
        return redirect('student_dashboard')
    elif request.user.role == 'admin':
        return redirect('admin_dashboard')
    elif request.user.role == 'dao':
        return redirect('dao_dashboard')
    elif request.user.role == 'dept_head':
        return redirect('dept_head_dashboard')
    return render(request, 'homepage.html')


#Auth 

def sign_in_html(request):
    if request.user.is_authenticated:
        return redirect('home')
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '')
        try:
            user_obj = User.objects.get(email=email)
            user = authenticate(request, username=user_obj.username, password=password)
            if user and user.is_active:
                login(request, user)
                return redirect('home')
        except User.DoesNotExist:
            pass
        return render(request, 'sign_in.html', {'error': 'Invalid email or password.'})
    return render(request, 'sign_in.html')


def sign_up_html(request):
    if request.user.is_authenticated:
        return redirect('home')
    if request.method == 'POST':
        full_name = request.POST.get('full_name', '').strip()
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '')
        role = request.POST.get('role', 'student')

        # Basic validations
        if not full_name or not email or not password:
            return render(request, 'sign_up.html', {'error': 'All fields are required.'})
        if len(password) < 8:
            return render(request, 'sign_up.html', {'error': 'Password must be at least 8 characters.'})
        if User.objects.filter(email=email).exists():
            return render(request, 'sign_up.html', {'error': 'Email already registered.'})

        # --- EMAIL DOMAIN VALIDATION BASED ON ROLE ---
        import re
        if role == 'student':
            # Must be digits@uap-bd.edu
            if not re.match(r'^\d+@uap-bd\.edu$', email):
                return render(request, 'sign_up.html', {'error': 'Student email must be digits@uap-bd.edu (e.g., 20241001@uap-bd.edu).'})
        elif role in ('faculty', 'admin'):
            # Must be name (letters/dots/underscores)@uap-bd.edu or admin@uap-bd.edu
            if not re.match(r'^[A-Za-z][A-Za-z0-9._]*@uap-bd\.edu$', email):
                return render(request, 'sign_up.html', {'error': 'Faculty email must be name@uap-bd.edu (e.g., john.doe@uap-bd.edu).'})
        else:
            return render(request, 'sign_up.html', {'error': 'Invalid role selected.'})

        # Generate unique username from email local part
        username = email.split('@')[0]
        base = username
        i = 1
        while User.objects.filter(username=username).exists():
            username = f"{base}{i}"
            i += 1

        # Create user (your custom User model has `role` and `full_name` fields)
        user = User.objects.create_user(
            username=username,
            email=email,
            password=password,
            full_name=full_name,
            role=role,
        )
        return redirect('/signin/?registered=1')
    return render(request, 'sign_up.html')


def sign_out(request):
    logout(request)
    return redirect('sign_in_html')


# Faculty Required

def faculty_required(view_func):
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('sign_in_html')
        if request.user.role not in ('faculty', 'admin') and not request.user.is_superuser:
            return redirect('student_dashboard')
        return view_func(request, *args, **kwargs)
    wrapper.__name__ = view_func.__name__
    return wrapper


# Student Required 

def student_required(view_func):
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('sign_in_html')
        if request.user.role not in ('student', 'admin') and not request.user.is_superuser:
            return redirect('faculty_dashboard')
        return view_func(request, *args, **kwargs)
    wrapper.__name__ = view_func.__name__
    return wrapper



#Faculty Views

@faculty_required
def faculty_dashboard(request):
    now = timezone.now()
    courses = Course.objects.filter(faculty=request.user, is_active=True).prefetch_related('enrollments', 'clos')
    assessments = Assessment.objects.filter(course__in=courses)
    pending_subs = Submission.objects.filter(assessment__in=assessments, status='submitted')
    flagged_subs = Submission.objects.filter(assessment__in=assessments, status='flagged')
    graded_subs  = Submission.objects.filter(assessment__in=assessments, status='graded')
    total_students = Enrollment.objects.filter(course__in=courses).values('student').distinct().count()

    recent_submissions = Submission.objects.filter(
        assessment__in=assessments
    ).select_related('student', 'assessment', 'assessment__course').order_by('-submitted_at')[:6]

    announcements = Announcement.objects.filter(course__in=courses).order_by('-created_at')[:4]

    upcoming = assessments.filter(
        status='published', due_date__gte=now, due_date__lte=now + timedelta(days=7)
    ).select_related('course').order_by('due_date')

    courses_stats = []
    for course in courses:
        course.enrolled_count  = course.enrollments.count()
        course.pending_grading = pending_subs.filter(assessment__course=course).count()
        courses_stats.append(course)

    my_sections = Section.objects.filter(faculty=request.user).select_related('course').prefetch_related('students').order_by('course__code', 'batch', 'name')

    return render(request, 'faculty/dashboard.html', {
        'courses': courses_stats,
        'assessments_count': assessments.count(),
        'pending_count': pending_subs.count(),
        'flagged_count': flagged_subs.count(),
        'graded_count': graded_subs.count(),
        'total_students': total_students,
        'recent_submissions': recent_submissions,
        'announcements': announcements,
        'upcoming': upcoming,
        'my_sections': my_sections,
    })


@faculty_required
def faculty_enrolled_students(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True).prefetch_related(
        'enrollments', 'enrollments__student'
    ).order_by('code')
    search = request.GET.get('q', '').strip()
    course_filter = request.GET.get('course', '')
    all_enrollments = []
    for course in courses:
        for enroll in course.enrollments.all():
            student = enroll.student
            if search and search.lower() not in (student.full_name or '').lower() and search.lower() not in student.username.lower():
                continue
            if course_filter and str(course.id) != course_filter:
                continue
            all_enrollments.append({'enrollment': enroll, 'student': student, 'course': course})
    return render(request, 'faculty/enrolled_students.html', {
        'courses': courses,
        'all_enrollments': all_enrollments,
        'search': search,
        'course_filter': course_filter,
        'total_count': len(all_enrollments),
    })


@faculty_required
def faculty_courses(request):
    all_courses = Course.objects.filter(faculty=request.user, is_active=True).prefetch_related(
        'clos', 'clos__plos', 'enrollments', 'enrollments__student'
    )
    active_courses = []
    archived_courses = []
    for course in all_courses:
        plo_ids = set()
        for clo in course.clos.all():
            for plo in clo.plos.all():
                plo_ids.add(plo.id)
        course.plo_count = len(plo_ids)
        if course.is_archived:
            archived_courses.append(course)
        else:
            active_courses.append(course)

    all_enrollments = Enrollment.objects.filter(
        course__faculty=request.user, course__is_active=True
    ).select_related('student', 'course').order_by('student__full_name', 'course__code')

    plos = PLO.objects.all()
    my_sections = Section.objects.filter(faculty=request.user).select_related('course').prefetch_related('students').order_by('course__code', 'batch', 'name')
    return render(request, 'faculty/courses.html', {
        'courses': active_courses,
        'archived_courses': archived_courses,
        'plos': plos,
        'all_enrollments': all_enrollments,
        'my_sections': my_sections,
    })


@faculty_required
def archive_course(request, course_id):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    course.is_archived = not course.is_archived
    course.save()
    return JsonResponse({'success': True, 'is_archived': course.is_archived})
 


@faculty_required
def add_course(request):
    return JsonResponse({'error': 'Course creation is managed by DAO.'}, status=403)


@faculty_required
def add_clo(request, course_id):
    course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    if request.method == 'POST':
        data = json.loads(request.body)
        count = course.clos.count() + 1
        clo = CLO.objects.create(
            course=course, code=f"CLO{count}",
            description=data['description'],
            bloom_level=data['bloom_level']
        )
        if data.get('plo_ids'):
            clo.plos.set(PLO.objects.filter(id__in=data['plo_ids']))
        return JsonResponse({'success': True, 'id': clo.id, 'code': clo.code})
    return JsonResponse({'error': 'POST required'}, status=400)


@faculty_required
def delete_clo(request, clo_id):
    clo = get_object_or_404(CLO, id=clo_id, course__faculty=request.user)
    clo.delete()
    return JsonResponse({'success': True})


@faculty_required
def get_course_clos(request, course_id):
    course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    clos = list(course.clos.values('id', 'code', 'description', 'bloom_level'))
    return JsonResponse({'clos': clos})


@faculty_required
def add_student_to_course(request, course_id):
    course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    if request.method == 'POST':
        data = json.loads(request.body)
        try:
            student = User.objects.get(email=data['email'], role='student')
            Enrollment.objects.get_or_create(student=student, course=course)
            section_id = data.get('section_id')
            if section_id:
                section = Section.objects.filter(id=int(section_id), course=course, faculty=request.user).first()
                if section:
                    section.students.add(student)
            return JsonResponse({'success': True, 'name': student.full_name or student.username})
        except User.DoesNotExist:
            return JsonResponse({'error': 'Student not found'}, status=404)
    return JsonResponse({'error': 'POST required'}, status=400)


@faculty_required
def remove_student_from_course(request, course_id, student_id):
    course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    if request.method == 'POST':
        Enrollment.objects.filter(student_id=student_id, course=course).delete()
        return JsonResponse({'success': True})
    return JsonResponse({'error': 'POST required'}, status=400)
def add_students_by_range(request, course_id):
    course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    try:
        from_id = int(str(data.get('from_id', '')).strip())
        to_id   = int(str(data.get('to_id',   '')).strip())
    except (ValueError, TypeError):
        return JsonResponse({'error': 'Registration IDs must be numbers.'}, status=400)
    if from_id > to_id:
        return JsonResponse({'error': 'From ID must be less than or equal to To ID.'}, status=400)
    if to_id - from_id >= 500:
        return JsonResponse({'error': 'Range too large — maximum 500 students at once.'}, status=400)

    section = None
    section_id = data.get('section_id')
    if section_id:
        try:
            section = Section.objects.filter(id=int(section_id), course=course, faculty=request.user).first()
        except (ValueError, TypeError):
            pass

    added, already_enrolled, not_found = [], [], []
    for reg_id in range(from_id, to_id + 1):
        email = f'{reg_id}@uap-bd.edu'
        try:
            student = User.objects.get(email=email, role='student')
            _, created = Enrollment.objects.get_or_create(student=student, course=course)
            (added if created else already_enrolled).append(reg_id)
            if section:
                section.students.add(student)
        except User.DoesNotExist:
            not_found.append(reg_id)

    return JsonResponse({
        'success': True,
        'added': len(added),
        'already_enrolled': len(already_enrolled),
        'not_found': len(not_found),
        'not_found_ids': not_found[:30],
    })


@faculty_required
def faculty_assessments(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)
    assessments = Assessment.objects.filter(course__in=courses).prefetch_related(
        'questions__clos__plos'
    ).order_by('-created_at')
    return render(request, 'faculty/assessments.html', {
        'assessments': assessments, 'courses': courses
    })


@faculty_required
def create_assessment(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        course = get_object_or_404(Course, id=data['course_id'], faculty=request.user, is_active=True)
        assessment = Assessment.objects.create(
            course=course, title=data['title'],
            description=data.get('description', ''),
            assessment_type=data['assessment_type'],
            due_date=data['due_date'], status='published'
        )
        total = 0
        for i, q in enumerate(data.get('questions', []), 1):
            question = Question.objects.create(
                assessment=assessment, order=i,
                text=q['text'], max_marks=int(q['max_marks'])
            )
            if q.get('clo_ids'):
                question.clos.set(CLO.objects.filter(id__in=q['clo_ids']))
            total += int(q['max_marks'])
        Assessment.objects.filter(pk=assessment.pk).update(total_marks=total)
        section_ids = data.get('section_ids', [])
        if section_ids:
            assessment.sections.set(Section.objects.filter(id__in=section_ids, course=course))
        return JsonResponse({'success': True, 'id': assessment.id})
    return JsonResponse({'error': 'POST required'}, status=400)


@faculty_required
def faculty_grading(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)

    course_id = request.GET.get('course')
    if not course_id:
        # Level 1 — course cards with submission counts
        course_cards = []
        for c in courses:
            assessments = Assessment.objects.filter(course=c)
            subs = Submission.objects.filter(assessment__in=assessments)
            total   = subs.count()
            pending = subs.filter(status='submitted').count()
            graded  = subs.filter(status__in=['graded', 'flagged']).count()
            if total > 0:
                course_cards.append({
                    'course':  c,
                    'total':   total,
                    'pending': pending,
                    'graded':  graded,
                })
        return render(request, 'faculty/grading.html', {
            'course_cards': course_cards,
            'selected_course': None,
        })

    # Level 2 — submissions for selected course
    selected_course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    assessments = Assessment.objects.filter(course=selected_course)
    submissions = Submission.objects.filter(
        assessment__in=assessments
    ).select_related('student', 'assessment__course').order_by('-submitted_at')
    return render(request, 'faculty/grading.html', {
        'submissions':     submissions,
        'assessments':     assessments,
        'selected_course': selected_course,
        'pending':         submissions.filter(status='submitted').count(),
        'graded':          submissions.filter(status__in=['graded', 'flagged']).count(),
        'course_cards':    None,
    })


@faculty_required
def get_submission_detail(request, sub_id):
    sub = get_object_or_404(Submission, id=sub_id, assessment__course__faculty=request.user)
    questions = []
    for q in sub.assessment.questions.all().order_by('order'):
        try:
            obtained = QuestionGrade.objects.get(submission=sub, question=q).marks_obtained
        except QuestionGrade.DoesNotExist:
            obtained = 0

        sub_qs = []
        for sq in q.sub_questions.order_by('order'):
            try:
                sq_obtained = SubQuestionGrade.objects.get(submission=sub, sub_question=sq).marks_obtained
            except SubQuestionGrade.DoesNotExist:
                sq_obtained = 0
            sub_qs.append({
                'id': sq.id, 'order': sq.order, 'text': sq.text,
                'image_url': sq.image.url if sq.image else None,
                'max_marks': sq.max_marks, 'obtained': sq_obtained,
                'clos': [{'code': c.code} for c in sq.clos.all()],
                'plos': [{'code': p.code} for p in sq.plos.all()],
            })

        questions.append({
            'id': q.id, 'order': q.order, 'text': q.text,
            'image_url': q.image.url if q.image else None,
            'max_marks': q.max_marks, 'obtained': obtained,
            'clos': [{'code': c.code} for c in q.clos.all()],
            'plos': [{'code': p.code} for p in q.plos.all()],
            'sub_questions': sub_qs,
        })
    file_url = sub.submitted_file.url if sub.submitted_file else None
    file_name = sub.submitted_file.name.split('/')[-1] if sub.submitted_file else None
    return JsonResponse({
        'id': sub.id,
        'student_name': sub.student.full_name or sub.student.username,
        'assessment_title': sub.assessment.title,
        'assessment_type': sub.assessment.assessment_type,
        'total_marks': sub.assessment.total_marks,
        'content': sub.content,
        'file_url': file_url,
        'file_name': file_name,
        'plagiarism': sub.plagiarism_score,
        'ai_content': sub.ai_content_score,
        'status': sub.status,
        'total_score': sub.total_score,
        'feedback': sub.feedback,
        'questions': questions,
    })


@faculty_required
def grade_submission(request, sub_id):
    sub = get_object_or_404(Submission, id=sub_id, assessment__course__faculty=request.user)
    if request.method == 'POST':
        data = json.loads(request.body)
        total = 0

        # Sub-question grades
        sq_totals = defaultdict(float)  # question_id → accumulated sub marks
        for sqg in data.get('sub_question_grades', []):
            sq = get_object_or_404(SubQuestion, id=sqg['sub_question_id'])
            marks = min(float(sqg['marks']), sq.max_marks)
            SubQuestionGrade.objects.update_or_create(
                submission=sub, sub_question=sq, defaults={'marks_obtained': marks}
            )
            sq_totals[sq.question_id] += marks

        # Roll sub-question totals up to question level
        for q_id, q_marks in sq_totals.items():
            q = get_object_or_404(Question, id=q_id)
            QuestionGrade.objects.update_or_create(
                submission=sub, question=q, defaults={'marks_obtained': q_marks}
            )
            total += q_marks

        # Plain question grades (questions without sub-questions)
        for qg_data in data.get('question_grades', []):
            q = get_object_or_404(Question, id=qg_data['question_id'])
            if q.id in sq_totals:
                continue  # already handled above
            marks = min(float(qg_data['marks']), q.max_marks)
            QuestionGrade.objects.update_or_create(
                submission=sub, question=q, defaults={'marks_obtained': marks}
            )
            total += marks
 
        was_graded_before = sub.status in ('graded', 'flagged')  # ← নতুন
 
        status = 'graded'
        sub.total_score = total
        sub.feedback    = data.get('feedback', '')
        sub.status      = status
        sub.save()

        recalculate_final_score(sub)

        if not was_graded_before:
            notify_grade_released(sub)
 
        return JsonResponse({'success': True})
    return JsonResponse({'error': 'POST required'}, status=400)


@faculty_required
def faculty_analytics(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)
    selected_course = None
    grade_dist = []
    clo_attainment = []
    student_details = []
    plo_attainment = []
    escar_clos = []
    escar_plos = []

    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    elif courses.exists():
        selected_course = courses.first()

    if selected_course:
        SUB_TYPES = {'mid', 'final'}

        # Only published assessments — same scope as marks sheet
        assessments = list(
            Assessment.objects.filter(course=selected_course, status='published')
            .prefetch_related(
                'questions__clos', 'questions__plos',
                'questions__sub_questions__clos', 'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        # Build column list exactly like marks sheet:
        # mid/final → one column per sub-question; others → one column per question
        # is_final flag tracks which columns belong to the final exam
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub':    True,
                            'entity_id': sq.id,
                            'max_marks': sq.max_marks,
                            'clo_ids':   [c.id for c in sq.clos.all()],
                            'plo_ids':   [p.id for p in sq.plos.all()],
                            'is_final':  is_final,
                        })
                else:
                    all_columns.append({
                        'is_sub':    False,
                        'entity_id': q.id,
                        'max_marks': q.max_marks,
                        'clo_ids':   [c.id for c in q.clos.all()],
                        'plo_ids':   [p.id for p in q.plos.all()],
                        'is_final':  is_final,
                    })

        clos = list(selected_course.clos.all())

        # Fixed denominators across ALL published assessments (same as marks sheet)
        clo_max = {
            clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids'])
            for clo in clos
        }
        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_map = {p.id: p for p in plos}
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }
        total_max_overall = sum(col['max_marks'] for col in all_columns)

        # University grading scale buckets (label, lo_inclusive, hi_exclusive)
        GRADE_SCALE = [
            ('A+', 80, 200), ('A', 75, 80), ('A-', 70, 75),
            ('B+', 65, 70),  ('B', 60, 65), ('B-', 55, 60),
            ('C+', 50, 55),  ('C', 45, 50), ('D', 40, 45), ('F', 0, 40),
        ]
        grade_buckets = {label: 0 for label, _, _ in GRADE_SCALE}

        # Fetch all grades in two bulk queries (same as marks sheet)
        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map  = {}
        sq_grade_map = {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=selected_course,
            ).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=selected_course,
            ).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

        # All distinct enrolled students (no duplicate-enrollment risk)
        students = list(
            User.objects.filter(enrollments__course=selected_course)
            .distinct().order_by('full_name', 'username')
        )
        total_enrolled = len(students)

        # Determine present students: those with ANY grade record in the final exam.
        # If no final exam exists, fall back to anyone with any grades.
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(
                    col['entity_id'] in (sq_m if col['is_sub'] else q_m)
                    for col in all_columns if col['is_final']
                )
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)

        # Per-CLO/PLO: count all graded students who achieved >= 40% on that CLO/PLO
        clo_achieved = {clo.id: 0 for clo in clos}
        plo_achieved = {p.id: 0 for p in plos}
        graded_count = 0

        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue

            graded_count += 1

            def _mark(col):
                return sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0)

            # Grade distribution — all students who have any grades
            if total_max_overall > 0:
                total_raw_student = sum(_mark(col) for col in all_columns)
                student_pct = total_raw_student / total_max_overall * 100
                for lbl, lo, hi in GRADE_SCALE:
                    if lo <= student_pct < hi:
                        grade_buckets[lbl] += 1
                        break

            # Per-student CLO/PLO percentages — same rounded pct for display AND counting
            clo_data = []
            for clo in clos:
                mx = clo_max[clo.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
                pct = round(raw / mx * 100, 1)
                clo_data.append({'code': clo.code, 'pct': pct})
                if pct >= 40:
                    clo_achieved[clo.id] += 1
            plo_data = []
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                pct = round(raw / mx * 100, 1)
                plo_data.append({'code': p.code, 'pct': pct})
                if pct >= 40:
                    plo_achieved[p.id] += 1
            overall_pcts = [d['pct'] for d in clo_data]
            overall = round(sum(overall_pcts) / len(overall_pcts), 1) if overall_pcts else 0.0
            student_details.append({
                'name':     student.full_name or student.username,
                'clo_data': clo_data,
                'plo_data': plo_data,
                'overall':  overall,
            })

        # Grade distribution chart data
        grade_dist = [{'label': lbl, 'count': grade_buckets[lbl]} for lbl, _, _ in GRADE_SCALE]

        # CLO attainment: achieved = all graded students ≥40%; % uses present_count as denominator
        for clo in clos:
            achieved = clo_achieved[clo.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            clo_attainment.append({'code': clo.code, 'attainment': pct,
                                   'achieved': achieved, 'present': present_count})

        # PLO attainment
        for p in plos:
            achieved = plo_achieved[p.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            plo_attainment.append({'code': p.code, 'description': p.description,
                                   'attainment': pct, 'achieved': achieved, 'present': present_count})

        # eSCAR
        for clo in clos:
            achieved = clo_achieved[clo.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            attained = pct >= 60
            action = ('CLO attained. Maintain current teaching strategies.' if attained
                      else 'CLO not attained. Review teaching methods and provide additional practice.')
            escar_clos.append({'code': clo.code, 'achieved': achieved, 'present': present_count,
                               'pct': pct, 'attained': attained, 'action': action})

        for p in plos:
            achieved = plo_achieved[p.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            attained = pct >= 60
            action = ('PLO attained. Continue current curriculum alignment.' if attained
                      else 'PLO not attained. Strengthen curriculum alignment and CLO coverage.')
            escar_plos.append({'code': p.code, 'achieved': achieved, 'present': present_count,
                               'pct': pct, 'attained': attained, 'action': action})

        escar_rows = [
            {'clo': escar_clos[i] if i < len(escar_clos) else None,
             'plo': escar_plos[i] if i < len(escar_plos) else None}
            for i in range(max(len(escar_clos), len(escar_plos), 1))
        ]
    else:
        escar_rows = []

    return render(request, 'faculty/analytics.html', {
        'courses': courses,
        'selected_course': selected_course,
        'grade_dist': json.dumps(grade_dist),
        'clo_attainment': json.dumps(clo_attainment),
        'clo_attainment_list': clo_attainment,
        'student_details': student_details,
        'plo_attainment': json.dumps(plo_attainment),
        'plo_attainment_list': plo_attainment,
        'escar_rows': escar_rows,
    })


def _compute_escar(selected_course):
    """Compute ESCAR data for a course. Returns (rows, total_enrolled, total_participated, total_absent)."""
    if not selected_course:
        return [], 0, 0, 0

    assessments = Assessment.objects.filter(course=selected_course, status='published')
    enrolled_students = list(Enrollment.objects.filter(course=selected_course).select_related('student'))
    total_enrolled = len(enrolled_students)

    all_q_ids = list(Question.objects.filter(assessment__in=assessments).values_list('id', flat=True))
    all_grades = QuestionGrade.objects.filter(
        question_id__in=all_q_ids,
        submission__assessment__course=selected_course,
    ).values('submission__student_id', 'question_id', 'marks_obtained')

    grade_map = {}
    for g in all_grades:
        sid = g['submission__student_id']
        grade_map.setdefault(sid, {})[g['question_id']] = g['marks_obtained']

    participated_ids = set(grade_map.keys())
    total_participated = len(participated_ids)
    total_absent = total_enrolled - total_participated

    q_maxmarks = {q.id: q.max_marks for q in Question.objects.filter(id__in=all_q_ids)}

    def _attainment(q_ids, q_marks_map):
        s_passed = 0; s_assessed = 0
        for enrollment in enrolled_students:
            sid = enrollment.student_id
            if sid not in participated_ids:
                continue
            sg = grade_map.get(sid, {})
            student_q_grades = {qid: sg[qid] for qid in q_ids if qid in sg}
            if not student_q_grades:
                continue
            student_max = sum(q_marks_map.get(qid, 0) for qid in student_q_grades)
            if student_max == 0:
                continue
            obtained = sum(student_q_grades.values())
            s_assessed += 1
            if (obtained / student_max * 100) >= 40:
                s_passed += 1
        pct = round((s_passed / s_assessed * 100) if s_assessed > 0 else 0, 2)
        return s_passed, s_assessed, pct

    clos_data = []
    for clo in selected_course.clos.all():
        q_ids = list(Question.objects.filter(assessment__in=assessments, clos=clo).values_list('id', flat=True))
        q_marks_map = {qid: q_maxmarks[qid] for qid in q_ids if qid in q_maxmarks}
        s_passed, s_assessed, attainment_pct = _attainment(q_ids, q_marks_map) if q_ids else (0, 0, 0.0)
        plan_obj = CLOActionPlan.objects.filter(course=selected_course, clo=clo).first()
        clos_data.append({
            'id': clo.id, 'code': clo.code, 'description': clo.description,
            'students_assessed': s_assessed, 'students_passed': s_passed,
            'attainment_pct': attainment_pct, 'attained': attainment_pct >= 60,
            'action_plan': plan_obj.action_plan if plan_obj else '',
        })

    plo_q_map = defaultdict(set)
    plo_obj_map = {}
    for clo in selected_course.clos.all():
        q_ids = list(Question.objects.filter(assessment__in=assessments, clos=clo).values_list('id', flat=True))
        for plo in clo.plos.all():
            plo_q_map[plo.id].update(q_ids)
            plo_obj_map[plo.id] = plo

    plos_data = []
    for plo_id, q_id_set in plo_q_map.items():
        plo = plo_obj_map[plo_id]
        q_ids = list(q_id_set)
        if not q_ids:
            continue
        q_marks_map = {qid: q_maxmarks[qid] for qid in q_ids if qid in q_maxmarks}
        s_passed, s_assessed, attainment_pct = _attainment(q_ids, q_marks_map)
        plan_obj = PLOActionPlan.objects.filter(course=selected_course, plo=plo).first()
        plos_data.append({
            'id': plo.id, 'code': plo.code, 'description': plo.description,
            'students_assessed': s_assessed, 'students_passed': s_passed,
            'attainment_pct': attainment_pct, 'attained': attainment_pct >= 60,
            'action_plan': plan_obj.action_plan if plan_obj else '',
        })

    max_len = max(len(clos_data), len(plos_data), 1)
    rows = [{'clo': clos_data[i] if i < len(clos_data) else None,
             'plo': plos_data[i] if i < len(plos_data) else None}
            for i in range(max_len)]
    return rows, total_enrolled, total_participated, total_absent


@faculty_required
def faculty_plo_comparison(request):
    """
    Cross-semester PLO attainment comparison for faculty.
    Faculty picks a course code; the view computes per-semester PLO attainment
    (same metric as faculty_analytics: % of present students who scored ≥40% on each PLO)
    and returns a comparison table + chart data.
    """
    # Include archived so faculty can compare historical semesters
    all_faculty_courses = Course.objects.filter(faculty=request.user).order_by('code', 'semester')

    # Unique course codes in stable order
    seen_codes = set()
    course_codes = []
    for c in all_faculty_courses:
        if c.code not in seen_codes:
            seen_codes.add(c.code)
            course_codes.append({'code': c.code, 'name': c.name})

    selected_code = request.GET.get('course_code', '').strip()
    selected_ids  = request.GET.getlist('sem_ids')   # course IDs chosen for comparison

    courses_for_code = []
    if selected_code:
        courses_for_code = list(all_faculty_courses.filter(code=selected_code))
    elif course_codes:
        selected_code = course_codes[0]['code']
        courses_for_code = list(all_faculty_courses.filter(code=selected_code))

    # Which course instances to actually compare
    if selected_ids:
        compare_set = set(selected_ids)
        courses_to_compare = [c for c in courses_for_code if str(c.id) in compare_set]
    else:
        courses_to_compare = list(courses_for_code)

    SUB_TYPES = {'mid', 'final'}
    comparison_data = []   # [{course, semester, plo_attainment:{code->pct}, present_count}]
    all_plo_meta   = {}    # code -> description

    for course in courses_to_compare:
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related(
                'questions__plos',
                'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub': True, 'entity_id': sq.id,
                            'max_marks': sq.max_marks,
                            'plo_ids': [p.id for p in sq.plos.all()],
                            'is_final': is_final,
                        })
                else:
                    all_columns.append({
                        'is_sub': False, 'entity_id': q.id,
                        'max_marks': q.max_marks,
                        'plo_ids': [p.id for p in q.plos.all()],
                        'is_final': is_final,
                    })

        if not all_columns:
            comparison_data.append({
                'course': course, 'semester': course.semester,
                'plo_attainment': {}, 'present_count': 0,
            })
            continue

        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }
        for p in plos:
            all_plo_meta[p.code] = p.description

        # Bulk fetch grades
        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=course,
            ).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=course,
            ).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

        students = list(User.objects.filter(enrollments__course=course).distinct())

        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(
                    col['entity_id'] in (sq_m if col['is_sub'] else q_m)
                    for col in all_columns if col['is_final']
                )
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)

        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue

            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)

            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if present_count > 0 and (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1

        plo_attainment = {}
        for p in plos:
            plo_attainment[p.code] = (
                round(plo_achieved[p.id] / present_count * 100, 1) if present_count > 0 else 0.0
            )

        comparison_data.append({
            'course': course, 'semester': course.semester,
            'plo_attainment': plo_attainment, 'present_count': present_count,
        })

    # Sort semesters chronologically
    def _sem_key(entry):
        parts = entry['semester'].split()
        year  = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))
    comparison_data.sort(key=_sem_key)

    plo_codes = sorted(all_plo_meta.keys())

    # Table rows: one row per semester, one cell per PLO
    table_rows = []
    for entry in comparison_data:
        cells = []
        for code in plo_codes:
            att = entry['plo_attainment'].get(code)
            cells.append({'plo_code': code, 'attainment': att, 'has_data': att is not None})
        table_rows.append({
            'semester': entry['semester'],
            'course_id': entry['course'].id,
            'is_archived': entry['course'].is_archived,
            'present_count': entry['present_count'],
            'cells': cells,
        })

    # JSON for Chart.js (cells ordered to match plo_codes)
    chart_json = json.dumps({
        'plo_codes': plo_codes,
        'plo_descriptions': all_plo_meta,
        'semesters': [e['semester'] for e in comparison_data],
        'datasets': [
            {
                'semester': e['semester'],
                'values': [e['plo_attainment'].get(code) for code in plo_codes],
            }
            for e in comparison_data
        ],
    })

    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]

    return render(request, 'faculty/plo_comparison.html', {
        'course_codes':     course_codes,
        'selected_code':    selected_code,
        'courses_for_code': courses_for_code,
        'selected_ids':     [int(i) for i in selected_ids if i.isdigit()],
        'plo_codes':        plo_codes,
        'plo_list':         plo_list,
        'table_rows':       table_rows,
        'chart_json':       chart_json,
        'has_data':         bool(plo_codes and table_rows),
    })


@faculty_required
def faculty_semester_plo_comparison(request):
    """
    Semester-wise PLO attainment comparison for faculty.
    Aggregates class-level PLO attainment (% of present students ≥40%)
    across all courses taught by the faculty, grouped by semester.
    Table: rows = PLOs, columns = semesters.
    """
    SUB_TYPES = {'mid', 'final'}

    all_faculty_courses = list(
        Course.objects.filter(faculty=request.user).order_by('semester', 'code')
    )

    def _sem_key(s):
        parts = s.split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))

    all_semesters = sorted(set(c.semester for c in all_faculty_courses), key=_sem_key)
    selected_sems = request.GET.getlist('sems')
    if not selected_sems:
        selected_sems = list(all_semesters)
    selected_sems_set = set(selected_sems)

    # semester -> plo_code -> {description, total_achieved, total_present}
    semester_plo_data = {}
    all_plo_meta = {}
    # Track unique present student IDs per semester for the "Students Present" column
    semester_present_students = {}  # semester -> set of student ids

    for course in all_faculty_courses:
        if course.semester not in selected_sems_set:
            continue

        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related(
                'questions__plos',
                'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub': True, 'entity_id': sq.id,
                            'max_marks': sq.max_marks,
                            'plo_ids': [p.id for p in sq.plos.all()],
                            'is_final': is_final,
                        })
                else:
                    all_columns.append({
                        'is_sub': False, 'entity_id': q.id,
                        'max_marks': q.max_marks,
                        'plo_ids': [p.id for p in q.plos.all()],
                        'is_final': is_final,
                    })

        if not all_columns:
            continue

        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        if not plo_ids_used:
            continue

        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }
        for p in plos:
            all_plo_meta[p.code] = p.description

        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=course,
            ).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=course,
            ).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

        students = list(User.objects.filter(enrollments__course=course).distinct())

        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(
                    col['entity_id'] in (sq_m if col['is_sub'] else q_m)
                    for col in all_columns if col['is_final']
                )
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)

        if present_count == 0:
            continue

        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue

            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)

            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1

        semester = course.semester
        if semester not in semester_plo_data:
            semester_plo_data[semester] = {}
        semester_present_students.setdefault(semester, set()).update(present_student_ids)

        for p in plos:
            if plo_max[p.id] == 0:
                continue
            code = p.code
            if code not in semester_plo_data[semester]:
                semester_plo_data[semester][code] = {
                    'description': p.description,
                    'total_achieved': 0,
                    'total_present': 0,
                }
            semester_plo_data[semester][code]['total_achieved'] += plo_achieved[p.id]
            semester_plo_data[semester][code]['total_present'] += present_count

    plo_codes = sorted(all_plo_meta.keys())
    semesters_sorted = sorted(semester_plo_data.keys(), key=_sem_key)

    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]

    # Rows = semesters, columns = PLOs (standard comparison table orientation)
    table_rows = []
    for sem in semesters_sorted:
        cells = []
        for code in plo_codes:
            pe = semester_plo_data.get(sem, {}).get(code)
            if pe and pe['total_present'] > 0:
                att = round(pe['total_achieved'] / pe['total_present'] * 100, 1)
                cells.append({'plo_code': code, 'attainment': att, 'has_data': True})
            else:
                cells.append({'plo_code': code, 'attainment': 0, 'has_data': False})
        table_rows.append({
            'semester': sem,
            'present_count': len(semester_present_students.get(sem, set())),
            'cells': cells,
        })

    chart_json = json.dumps({
        'plo_codes': plo_codes,
        'semesters': semesters_sorted,
        'datasets': [
            {
                'plo': code,
                'values': [
                    (
                        round(semester_plo_data[sem][code]['total_achieved'] /
                              semester_plo_data[sem][code]['total_present'] * 100, 1)
                        if (sem in semester_plo_data and code in semester_plo_data[sem]
                            and semester_plo_data[sem][code]['total_present'] > 0)
                        else None
                    )
                    for sem in semesters_sorted
                ],
            }
            for code in plo_codes
        ],
    })

    return render(request, 'faculty/semester_plo_comparison.html', {
        'all_semesters':  all_semesters,
        'selected_sems':  selected_sems,
        'plo_codes':      plo_codes,
        'plo_list':       plo_list,
        'table_rows':     table_rows,
        'semesters':      semesters_sorted,
        'chart_json':     chart_json,
        'has_data':       bool(plo_codes and semesters_sorted),
    })


@faculty_required
def faculty_escar(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)
    selected_course = None
    clos_data = []
    plos_data = []
    rows = []
    total_enrolled = 0
    total_participated = 0
    total_absent = 0

    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    elif courses.exists():
        selected_course = courses.first()

    if selected_course:
        assessments = Assessment.objects.filter(course=selected_course, status='published')
        enrolled_students = list(Enrollment.objects.filter(course=selected_course).select_related('student'))
        total_enrolled = len(enrolled_students)

        # Read all grades the same way the marksheet does — no submission-status filter.
        all_q_ids = list(Question.objects.filter(assessment__in=assessments).values_list('id', flat=True))
        all_grades = QuestionGrade.objects.filter(
            question_id__in=all_q_ids,
            submission__assessment__course=selected_course,
        ).values('submission__student_id', 'question_id', 'marks_obtained')

        # grade_map mirrors the marksheet: {student_id: {question_id: marks}}
        grade_map = {}
        for g in all_grades:
            sid = g['submission__student_id']
            grade_map.setdefault(sid, {})[g['question_id']] = g['marks_obtained']

        participated_ids = set(grade_map.keys())
        total_participated = len(participated_ids)
        total_absent = total_enrolled - total_participated

        # Build question max-marks cache once for all questions in the course
        q_maxmarks = {q.id: q.max_marks for q in Question.objects.filter(id__in=all_q_ids)}

        def _attainment(q_ids, q_marks_map):
            """Return (students_passed, students_assessed, attainment_pct) for a set of question IDs."""
            s_passed = 0
            s_assessed = 0
            for enrollment in enrolled_students:
                sid = enrollment.student_id
                if sid not in participated_ids:
                    continue
                sg = grade_map.get(sid, {})
                # Only questions this student actually has a grade for
                student_q_grades = {qid: sg[qid] for qid in q_ids if qid in sg}
                if not student_q_grades:
                    continue
                student_max = sum(q_marks_map.get(qid, 0) for qid in student_q_grades)
                if student_max == 0:
                    continue
                obtained = sum(student_q_grades.values())
                s_assessed += 1
                if (obtained / student_max * 100) >= 40:
                    s_passed += 1
            pct = round((s_passed / s_assessed * 100) if s_assessed > 0 else 0, 2)
            return s_passed, s_assessed, pct

        for clo in selected_course.clos.all():
            q_ids = list(Question.objects.filter(assessment__in=assessments, clos=clo).values_list('id', flat=True))
            q_marks_map = {qid: q_maxmarks[qid] for qid in q_ids if qid in q_maxmarks}
            s_passed, s_assessed, attainment_pct = _attainment(q_ids, q_marks_map) if q_ids else (0, 0, 0.0)
            plan_obj = CLOActionPlan.objects.filter(course=selected_course, clo=clo).first()
            clos_data.append({
                'id': clo.id,
                'code': clo.code,
                'description': clo.description,
                'students_assessed': s_assessed,
                'students_passed': s_passed,
                'attainment_pct': attainment_pct,
                'attained': attainment_pct >= 60,
                'action_plan': plan_obj.action_plan if plan_obj else '',
            })

        # PLO mapping via CLO.plos (PLOs are assigned to CLOs, not questions directly)
        plo_q_map = defaultdict(set)
        plo_obj_map = {}
        for clo in selected_course.clos.all():
            q_ids = list(Question.objects.filter(assessment__in=assessments, clos=clo).values_list('id', flat=True))
            for plo in clo.plos.all():
                plo_q_map[plo.id].update(q_ids)
                plo_obj_map[plo.id] = plo

        for plo_id, q_id_set in plo_q_map.items():
            plo = plo_obj_map[plo_id]
            q_ids = list(q_id_set)
            if not q_ids:
                continue
            q_marks_map = {qid: q_maxmarks[qid] for qid in q_ids if qid in q_maxmarks}
            s_passed, s_assessed, attainment_pct = _attainment(q_ids, q_marks_map)
            plan_obj = PLOActionPlan.objects.filter(course=selected_course, plo=plo).first()
            plos_data.append({
                'id': plo.id,
                'code': plo.code,
                'description': plo.description,
                'students_assessed': s_assessed,
                'students_passed': s_passed,
                'attainment_pct': attainment_pct,
                'attained': attainment_pct >= 60,
                'action_plan': plan_obj.action_plan if plan_obj else '',
            })

        max_len = max(len(clos_data), len(plos_data), 1)
        for i in range(max_len):
            rows.append({
                'clo': clos_data[i] if i < len(clos_data) else None,
                'plo': plos_data[i] if i < len(plos_data) else None,
            })

    return render(request, 'faculty/escar.html', {
        'courses': courses,
        'selected_course': selected_course,
        'rows': rows,
        'total_enrolled': total_enrolled,
        'total_participated': total_participated,
        'total_absent': total_absent,
    })


@faculty_required
def save_escar_plan(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    course = get_object_or_404(Course, id=data['course_id'], faculty=request.user, is_active=True)
    plan_type = data.get('type')
    plan_text = data.get('action_plan', '')
    if plan_type == 'clo':
        clo = get_object_or_404(CLO, id=data['id'], course=course)
        obj, _ = CLOActionPlan.objects.get_or_create(course=course, clo=clo)
        obj.action_plan = plan_text
        obj.save()
    elif plan_type == 'plo':
        plo = get_object_or_404(PLO, id=data['id'])
        obj, _ = PLOActionPlan.objects.get_or_create(course=course, plo=plo)
        obj.action_plan = plan_text
        obj.save()
    else:
        return JsonResponse({'error': 'Invalid type'}, status=400)
    return JsonResponse({'success': True})


@faculty_required
def faculty_announcements(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)
    course_id = request.GET.get('course')
    selected_course = None
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    announcements = Announcement.objects.filter(
        course__in=courses
    ).select_related('course').order_by('-created_at')
    all_faculty_sections = Section.objects.filter(faculty=request.user).select_related('course').order_by('course__code', 'name')
    return render(request, 'faculty/announcements.html', {
        'announcements': announcements,
        'courses': courses,
        'selected_course': selected_course,
        'all_faculty_sections': all_faculty_sections,
    })


@faculty_required
def create_announcement(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        course = get_object_or_404(Course, id=data['course_id'], faculty=request.user, is_active=True)
        ann = Announcement.objects.create(
            course=course, title=data['title'], content=data['content'],
            priority=data.get('priority', 'medium'), created_by=request.user
        )
        section_ids = data.get('section_ids', [])
        if section_ids:
            ann.sections.set(Section.objects.filter(id__in=section_ids, course=course))
        notify_announcement(ann)
        return JsonResponse({'success': True, 'id': ann.id})
    return JsonResponse({'error': 'POST required'}, status=400)


@faculty_required
def delete_announcement(request, ann_id):
    ann = get_object_or_404(Announcement, id=ann_id, created_by=request.user)
    ann.delete()
    return JsonResponse({'success': True})


# Student Views

@student_required
def student_dashboard(request):
    now = timezone.now()
    enrollments = Enrollment.objects.filter(student=request.user).select_related('course')
    courses = [e.course for e in enrollments]
    assessments = Assessment.objects.filter(course__in=courses, status='published')
    submissions = Submission.objects.filter(student=request.user, assessment__in=assessments)
    submitted_ids = submissions.values_list('assessment_id', flat=True)

    pending_qs = assessments.exclude(id__in=submitted_ids).select_related('course').order_by('due_date')
    pending_count = pending_qs.count()
    upcoming = pending_qs.filter(due_date__isnull=False, due_date__gte=now)[:5]

    graded = submissions.filter(status='graded').select_related('assessment', 'assessment__course')
    avg_grade = 0
    if graded.exists():
        total_pct = sum(
            s.total_score / s.assessment.total_marks * 100
            for s in graded if s.assessment.total_marks > 0
        )
        avg_grade = round(total_pct / graded.count(), 1)

    recent_grades = graded.order_by('-submitted_at')[:5]
    from django.db.models import Q as _Q
    student_section_ids = Section.objects.filter(
        course__in=courses, students=request.user
    ).values_list('id', flat=True)
    announcements = Announcement.objects.filter(course__in=courses).filter(
        _Q(sections__isnull=True) | _Q(sections__id__in=student_section_ids)
    ).distinct().order_by('-created_at')[:4]

    my_sections = Section.objects.filter(students=request.user).select_related('course').prefetch_related('faculty').order_by('course__code', 'batch', 'name')

    return render(request, 'student/dashboard.html', {
        'courses': courses,
        'submissions_count': submissions.count(),
        'pending_count': pending_count,
        'graded_count': graded.count(),
        'avg_grade': avg_grade,
        'recent_grades': recent_grades,
        'announcements': announcements,
        'upcoming': upcoming,
        'my_sections': my_sections,
    })


@student_required
def student_courses(request):
    enrollments = Enrollment.objects.filter(student=request.user).select_related('course')
    active_courses = []
    archived_courses = []
    for e in enrollments:
        c = e.course
        c.clos_list = c.clos.prefetch_related('plos').all()
        c.assignment_count = Assessment.objects.filter(course=c, status='published').count()
        if c.is_archived:
            archived_courses.append(c)
        else:
            active_courses.append(c)
    my_sections = Section.objects.filter(students=request.user).select_related('course').prefetch_related('faculty')
    section_map = {s.course_id: s for s in my_sections}
    for c in active_courses:
        c.my_section = section_map.get(c.id)
    for c in archived_courses:
        c.my_section = section_map.get(c.id)

    return render(request, 'student/courses.html', {
        'courses': active_courses,
        'archived_courses': archived_courses,
    })


@student_required
def enroll_via_code(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        code = data.get('code', '').strip().upper()
        try:
            course = Course.objects.get(enrollment_code=code)
            if course.is_archived:
                return JsonResponse({'error': 'This course is archived and no longer accepting new enrollments.'}, status=400)
            _, created = Enrollment.objects.get_or_create(student=request.user, course=course)
            if created:
                return JsonResponse({'success': True, 'course': f"{course.code}: {course.name}"})
            return JsonResponse({'error': 'You are already enrolled in this course.'}, status=400)
        except Course.DoesNotExist:
            return JsonResponse({'error': 'Invalid enrollment code. Please check and try again.'}, status=404)
    return JsonResponse({'error': 'POST required'}, status=400)


@student_required
def unenroll_course(request, course_id):
    if request.method == 'POST':
        enrollment = get_object_or_404(Enrollment, student=request.user, course_id=course_id)
        enrollment.delete()
        return JsonResponse({'success': True})
    return JsonResponse({'error': 'POST required'}, status=400)


@student_required
def enroll_course(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    Enrollment.objects.get_or_create(student=request.user, course=course)
    return JsonResponse({'success': True})


@student_required
def student_submissions(request):
    enrollments = Enrollment.objects.filter(student=request.user)
    courses = [e.course for e in enrollments]
    assessments = Assessment.objects.filter(course__in=courses, status='published')
    submissions = Submission.objects.filter(
        student=request.user, assessment__in=assessments
    ).select_related('assessment__course').order_by('-submitted_at')
    submitted_ids = submissions.values_list('assessment_id', flat=True)
    pending = assessments.exclude(id__in=submitted_ids)
    today = date.today()
    todo_assessments = pending.filter(due_date__gte=today) | pending.filter(due_date__isnull=True)
    missing_assessments = pending.filter(due_date__lt=today)
    return render(request, 'student/submissions.html', {
        'submissions': submissions,
        'todo_assessments': todo_assessments,
        'missing_assessments': missing_assessments,
        'pending_assessments': todo_assessments | missing_assessments,
    })


@student_required
def submit_assessment(request, assessment_id):
    assessment = get_object_or_404(Assessment, id=assessment_id, status='published')
    if not Enrollment.objects.filter(student=request.user, course=assessment.course).exists():
        return JsonResponse({'error': 'Not enrolled'}, status=403)
    if assessment.course.is_archived:
        return JsonResponse({'error': 'This course is archived. Submissions are no longer accepted.'}, status=403)
    if request.method == 'POST':
        data = json.loads(request.body)
        content = data.get('content', '')
        sub, created = Submission.objects.get_or_create(
            student=request.user, assessment=assessment,
            defaults={
                'content': content,
                'plagiarism_score': round(len(content) % 50, 1),
                'ai_content_score': round(len(content) % 20, 1),
            }
        )
        if not created:
            return JsonResponse({'error': 'Already submitted'}, status=400)
        return JsonResponse({'success': True})
    return JsonResponse({'error': 'POST required'}, status=400)


@student_required
def student_clo_results(request):
    enrollments = Enrollment.objects.filter(student=request.user).select_related('course')
    results = []
    program_plo_data = {}   # plo_id -> {code, description, total_max, total_raw}
    semester_plo_data = {}  # semester -> {plo_id: {code, description, total_max, total_raw}}
    # semester -> [{course_code, course_name, max_plo_att, best_plo_code, best_plo_desc, plo_results}]
    semester_course_ranking = {}
    SUB_TYPES = {'mid', 'final'}

    for e in enrollments:
        course = e.course
        semester = course.semester
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related(
                'questions__clos', 'questions__plos',
                'questions__sub_questions__clos', 'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        # Build columns exactly like faculty analytics
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub':    True,
                            'entity_id': sq.id,
                            'max_marks': sq.max_marks,
                            'clo_ids':   [c.id for c in sq.clos.all()],
                            'plo_ids':   [p.id for p in sq.plos.all()],
                        })
                else:
                    all_columns.append({
                        'is_sub':    False,
                        'entity_id': q.id,
                        'max_marks': q.max_marks,
                        'clo_ids':   [c.id for c in q.clos.all()],
                        'plo_ids':   [p.id for p in q.plos.all()],
                    })

        clos = list(course.clos.all())
        clo_max = {
            clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids'])
            for clo in clos
        }
        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }
        total_max_overall = sum(col['max_marks'] for col in all_columns)

        # Fetch this student's grades in two bulk queries
        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_m, sq_m = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=course,
                submission__student=request.user,
            ):
                q_m[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=course,
                submission__student=request.user,
            ):
                sq_m[g.sub_question_id] = g.marks_obtained

        def _mark(col):
            return sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0)

        has_grades = bool(q_m or sq_m)

        # Overall percentage using same denominator as faculty
        total_raw = sum(_mark(col) for col in all_columns)
        avg_pct = round(total_raw / total_max_overall * 100, 1) if (total_max_overall > 0 and has_grades) else 0.0

        grade = 'F'
        if avg_pct >= 80:   grade = 'A+'
        elif avg_pct >= 75: grade = 'A'
        elif avg_pct >= 70: grade = 'A-'
        elif avg_pct >= 65: grade = 'B+'
        elif avg_pct >= 60: grade = 'B'
        elif avg_pct >= 55: grade = 'B-'
        elif avg_pct >= 50: grade = 'C+'
        elif avg_pct >= 45: grade = 'C'
        elif avg_pct >= 40: grade = 'D'

        # CLO attainment — same formula as faculty per-student calculation
        clo_results = []
        for clo in clos:
            mx = clo_max[clo.id]
            raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
            att = round(raw / mx * 100, 1) if mx > 0 else 0.0
            clo_results.append({
                'code': clo.code, 'bloom': clo.bloom_level,
                'description': clo.description,
                'obtained': round(raw, 1), 'total': mx,
                'attainment': att,
            })

        # PLO attainment
        plo_results = []
        for p in plos:
            mx = plo_max[p.id]
            if mx == 0:
                continue
            raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
            att = round(raw / mx * 100, 1)
            plo_results.append({'code': p.code, 'description': p.description, 'attainment': att})
            # accumulate for program-level summary
            if p.id not in program_plo_data:
                program_plo_data[p.id] = {
                    'code': p.code, 'description': p.description,
                    'total_max': 0.0, 'total_raw': 0.0,
                }
            program_plo_data[p.id]['total_max'] += mx
            program_plo_data[p.id]['total_raw'] += raw
            # accumulate for semester-level summary
            if semester not in semester_plo_data:
                semester_plo_data[semester] = {}
            if p.id not in semester_plo_data[semester]:
                semester_plo_data[semester][p.id] = {
                    'code': p.code, 'description': p.description,
                    'total_max': 0.0, 'total_raw': 0.0,
                }
            semester_plo_data[semester][p.id]['total_max'] += mx
            semester_plo_data[semester][p.id]['total_raw'] += raw

        # Build per-semester course ranking by max PLO attainment
        if plo_results:
            best = max(plo_results, key=lambda p: p['attainment'])
            semester_course_ranking.setdefault(semester, []).append({
                'course_code': course.code,
                'course_name': course.name,
                'max_plo_att': best['attainment'],
                'best_plo_code': best['code'],
                'best_plo_desc': best['description'],
                'plo_results': sorted(plo_results, key=lambda p: p['attainment'], reverse=True),
            })

        graded_count = Submission.objects.filter(
            student=request.user, assessment__in=assessments,
            status__in=['graded', 'flagged'],
        ).count()

        results.append({
            'course': course, 'grade': grade,
            'avg_pct': avg_pct, 'graded_count': graded_count,
            'clo_results': clo_results,
            'plo_results': plo_results,
        })

    program_plo_attainment = sorted(
        [
            {
                'code': d['code'],
                'description': d['description'],
                'attainment': round(d['total_raw'] / d['total_max'] * 100, 1),
                'obtained': round(d['total_raw'], 1),
                'total': round(d['total_max'], 1),
            }
            for d in program_plo_data.values() if d['total_max'] > 0
        ],
        key=lambda x: x['code'],
    )

    # Build semester-wise PLO comparison table
    def _sem_sort_key(s):
        parts = s.split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        term_order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, term_order.get(parts[0].lower() if parts else '', 5))

    _all_plo_info = {}
    for _sd in semester_plo_data.values():
        for _info in _sd.values():
            if _info['code'] not in _all_plo_info:
                _all_plo_info[_info['code']] = _info['description']

    _plo_codes = sorted(_all_plo_info.keys())
    _semesters = sorted(semester_plo_data.keys(), key=_sem_sort_key)

    _comparison_rows = []
    for _code in _plo_codes:
        _cells = []
        for _sem in _semesters:
            _sd = semester_plo_data.get(_sem, {})
            _pe = next((i for i in _sd.values() if i['code'] == _code), None)
            if _pe and _pe['total_max'] > 0:
                _att = round(_pe['total_raw'] / _pe['total_max'] * 100, 1)
                _cells.append({'semester': _sem, 'attainment': _att, 'has_data': True})
            else:
                _cells.append({'semester': _sem, 'attainment': 0, 'has_data': False})
        _comparison_rows.append({
            'code': _code,
            'description': _all_plo_info[_code],
            'cells': _cells,
        })

    semester_plo_comparison = {
        'semesters': _semesters,
        'rows': _comparison_rows,
    }

    # Sort each semester's courses by max PLO attainment descending
    semester_plo_ranking = []
    for _sem in _semesters:
        _courses = semester_course_ranking.get(_sem)
        if not _courses:
            continue
        _courses.sort(key=lambda c: c['max_plo_att'], reverse=True)
        semester_plo_ranking.append({'semester': _sem, 'courses': _courses})

    return render(request, 'student/clo_results.html', {
        'results': results,
        'program_plo_attainment': program_plo_attainment,
        'semester_plo_comparison': semester_plo_comparison,
        'semester_plo_ranking': semester_plo_ranking,
    })


@student_required
def student_notifications(request):
    notifs      = Notification.objects.filter(
        recipient=request.user
    ).select_related('course', 'assessment')
    unread_count = notifs.filter(is_read=False).count()
    # mark all read when page opens
    notifs.filter(is_read=False).update(is_read=True)
    return render(request, 'student/notifications.html', {
        'notifs':       notifs,
        'unread_count': unread_count,
    })



# Study Material Views 

@faculty_required
def faculty_materials(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)
 
    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
        materials = StudyMaterial.objects.filter(
            course=selected_course
        ).order_by('-uploaded_at')
        faculty_sections = list(Section.objects.filter(course=selected_course, faculty=request.user).order_by('name'))
        return render(request, 'faculty/materials.html', {
            'selected_course': selected_course,
            'courses': courses,
            'materials': materials,
            'material_count': materials.count(),
            'faculty_sections': faculty_sections,
        })
 
    # Course list view — annotate counts
    for c in courses:
        c.material_count = StudyMaterial.objects.filter(course=c).count()
        c.visible_count  = StudyMaterial.objects.filter(course=c, is_visible=True).count()
        c.hidden_count   = StudyMaterial.objects.filter(course=c, is_visible=False).count()
 
    return render(request, 'faculty/materials.html', {
        'selected_course': None,
        'courses': courses,
    })


@faculty_required
def upload_material(request):
    if request.method == 'POST':
        course_id     = request.POST.get('course_id')
        title         = request.POST.get('title', '').strip()
        description   = request.POST.get('description', '').strip()
        material_type = request.POST.get('material_type', 'lecture_note')
        video_url     = request.POST.get('video_url', '').strip()
        uploaded_file = request.FILES.get('file')
 
        course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
 
        # ── Permission check ──
        if not title:
            return JsonResponse({'error': 'Title is required.'}, status=400)
        if material_type == 'video' and not video_url and not uploaded_file:
            return JsonResponse({'error': 'Please provide a video URL or upload a video file.'}, status=400)
        if material_type != 'video' and not uploaded_file and not video_url:
            return JsonResponse({'error': 'Please select a file to upload.'}, status=400)
 
        material = StudyMaterial.objects.create(
            course=course,
            title=title,
            description=description,
            material_type=material_type,
            file=uploaded_file,
            video_url=video_url,
            uploaded_by=request.user,
            is_visible=True,
        )
        section_ids = request.POST.getlist('section_ids')
        if section_ids:
            material.sections.set(Section.objects.filter(id__in=section_ids, course=course))
        return JsonResponse({
            'success':       True,
            'id':            material.id,
            'title':         material.title,
            'description':   material.description,
            'material_type': material.material_type,
            'type_label':    material.get_material_type_display(),
            'file_url':      material.file.url if material.file else '',
            'filename':      material.filename(),
            'video_url':     material.video_url,
            'embed_url':     material.embed_url(),
            'is_video':      material.is_video(),
            'is_visible':    material.is_visible,
            'uploaded_at':   material.uploaded_at.strftime('%b %d, %Y'),
        })
    return JsonResponse({'error': 'POST required'}, status=400)

@faculty_required
def toggle_material_visibility(request, material_id):
    """Faculty can show/hide a material from students."""
    material = get_object_or_404(StudyMaterial, id=material_id, course__faculty=request.user)
    material.is_visible = not material.is_visible
    material.save()
    return JsonResponse({'success': True, 'is_visible': material.is_visible})


@faculty_required
def delete_material(request, material_id):
    material = get_object_or_404(StudyMaterial, id=material_id, course__faculty=request.user)
    if material.file:
        material.file.delete(save=False)
    material.delete()
    return JsonResponse({'success': True})


@student_required
def student_materials(request):
    enrollments = Enrollment.objects.filter(student=request.user)
    courses     = [e.course for e in enrollments]
 
    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=int(course_id))
 
        # Permission check — must be enrolled
        if selected_course not in courses:
            from django.shortcuts import redirect
            return redirect('student_materials')
 
        student_sections = Section.objects.filter(course=selected_course, students=request.user)
        from django.db.models import Q
        materials = StudyMaterial.objects.filter(
            course=selected_course,
            is_visible=True,
        ).filter(
            Q(sections__isnull=True) | Q(sections__in=student_sections)
        ).distinct().order_by('-uploaded_at')
 
        return render(request, 'student/materials.html', {
            'selected_course': selected_course,
            'materials':       materials,
            'material_count':  materials.count(),
        })
 
    # Course list
    for c in courses:
        c.material_count = StudyMaterial.objects.filter(course=c, is_visible=True).count()
 
    return render(request, 'student/materials.html', {
        'selected_course': None,
        'courses':         courses,
    })

@faculty_required
def add_plo(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        desc = data.get('description', '').strip()
        if not desc:
            return JsonResponse({'error': 'Description required'}, status=400)
        count = PLO.objects.count() + 1
        plo = PLO.objects.create(
            code=f"PLO{count}",
            description=desc,
            created_by=request.user
        )
        return JsonResponse({'success': True, 'id': plo.id, 'code': plo.code})
    return JsonResponse({'error': 'POST required'}, status=400)


# Faculty Assignment Views 

@faculty_required
def faculty_assignments(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)
    plos = PLO.objects.all()

    course_id = request.GET.get('course')
    if course_id:
        # Course detail view 
        selected_course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
        base_qs = Assessment.objects.filter(course=selected_course).select_related('course').prefetch_related(
            'questions', 'questions__clos', 'questions__plos'
        )
        drafts = list(base_qs.filter(status='draft').order_by('-created_at'))
        published_list = list(base_qs.filter(status='published').order_by('-created_at'))
        for a in drafts + published_list:
            a.submission_count = a.submissions.count()
            a.question_count = a.questions.count()
        faculty_sections = list(Section.objects.filter(course=selected_course, faculty=request.user).order_by('name'))
        return render(request, 'faculty/assignments.html', {
            'selected_course': selected_course,
            'courses': courses,
            'drafts': drafts,
            'published_list': published_list,
            'plos': plos,
            'draft_count': len(drafts),
            'published_count': len(published_list),
            'total_subs': sum(a.submission_count for a in published_list),
            'faculty_sections': faculty_sections,
        })

    # Course list view 
    for c in courses:
        c.total_count = Assessment.objects.filter(course=c).count()
        c.draft_count = Assessment.objects.filter(course=c, status='draft').count()
        c.published_count = Assessment.objects.filter(course=c, status='published').count()
    return render(request, 'faculty/assignments.html', {
        'selected_course': None,
        'courses': courses,
        'plos': plos,
    })


def _decode_image(b64_str, prefix='img'):
    """Decode a base64 data-URL string into a ContentFile. Returns None on failure."""
    if not b64_str or ',' not in b64_str:
        return None
    try:
        header, data = b64_str.split(',', 1)
        ext = header.split('/')[1].split(';')[0]
        return ContentFile(base64.b64decode(data), name=f'{prefix}_{uuid.uuid4().hex[:8]}.{ext}')
    except Exception:
        return None


@faculty_required
def create_assignment(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        course = get_object_or_404(Course, id=data['course_id'], faculty=request.user, is_active=True)
        assessment_type = data.get('assessment_type', 'assignment')
        # Faculty can explicitly choose publish or draft; assignment defaults to publish
        default_publish = assessment_type == 'assignment'
        publish_immediately = data.get('publish_immediately', default_publish)
        status = 'published' if publish_immediately else 'draft'
        # due_date is optional for non-assignment types
        due_date = data.get('due_date') or None
        assignment = Assessment.objects.create(
            course=course,
            title=data['title'],
            description=data.get('description', ''),
            assessment_type=assessment_type,
            due_date=due_date,
            status=status,
            total_marks=0,
            grace_period_hours   = int(data.get('grace_period_hours', 0)),
            late_deduction_type  = data.get('late_deduction_type', 'percent'),
            late_deduction_value = float(data.get('late_deduction_value', 0)),
            max_late_days        = int(data.get('max_late_days', 0)),
        )
        total = 0
        for i, q in enumerate(data.get('questions', []), 1):
            question = Question.objects.create(
                assessment=assignment,
                order=i,
                text=q['text'],
                max_marks=0,
            )
            # Question-level image
            img_file = _decode_image(q.get('image_b64'), f'q{i}')
            if img_file:
                question.image.save(img_file.name, img_file, save=False)

            sub_questions = q.get('sub_questions', [])
            if sub_questions:
                all_clo_ids, all_plo_ids, q_total = set(), set(), 0
                for j, sq in enumerate(sub_questions, 1):
                    sub = SubQuestion.objects.create(
                        question=question,
                        order=j,
                        text=sq['text'],
                        max_marks=int(sq.get('max_marks', 5)),
                    )
                    sq_img = _decode_image(sq.get('image_b64'), f'sq{i}_{j}')
                    if sq_img:
                        sub.image.save(sq_img.name, sq_img, save=False)
                    if sq.get('clo_ids'):
                        sub.clos.set(CLO.objects.filter(id__in=sq['clo_ids']))
                        all_clo_ids.update(sq['clo_ids'])
                    if sq.get('plo_ids'):
                        sub.plos.set(PLO.objects.filter(id__in=sq['plo_ids']))
                        all_plo_ids.update(sq['plo_ids'])
                    sub.save()
                    q_total += int(sq.get('max_marks', 5))
                question.max_marks = q_total
                if all_clo_ids:
                    question.clos.set(CLO.objects.filter(id__in=all_clo_ids))
                if all_plo_ids:
                    question.plos.set(PLO.objects.filter(id__in=all_plo_ids))
                total += q_total
            else:
                # Fallback: question-level marks/CLO/PLO (backward compat)
                question.max_marks = int(q.get('max_marks', 10))
                if q.get('clo_ids'):
                    question.clos.set(CLO.objects.filter(id__in=q['clo_ids']))
                if q.get('plo_ids'):
                    question.plos.set(PLO.objects.filter(id__in=q['plo_ids']))
                total += question.max_marks
            question.save()
        # Use manual total_marks if provided and no questions, else sum of questions
        manual_total = int(data.get('total_marks', 0))
        assignment.total_marks = total if total > 0 else manual_total
        assignment.save()

        section_ids = data.get('section_ids', [])
        if section_ids:
            assignment.sections.set(Section.objects.filter(id__in=section_ids, course=course))

        if status == 'published':
            notify_new_assignment(assignment)

        type_labels = dict(Assessment.TYPE_CHOICES)
        return JsonResponse({
            'success': True,
            'id': assignment.id,
            'title': assignment.title,
            'type_label': type_labels.get(assessment_type, assessment_type),
            'assessment_type': assessment_type,
            'status': status,
            'course_name': f"{course.code}: {course.name}",
            'due_date': str(assignment.due_date),
            'total_marks': assignment.total_marks,
            'description': assignment.description,
            'question_count': len(data.get('questions', [])),
        })
    return JsonResponse({'error': 'POST required'}, status=400)


@faculty_required
def delete_assignment(request, assignment_id):
    assignment = get_object_or_404(
        Assessment, id=assignment_id, course__faculty=request.user
    )
    assignment.delete()
    return JsonResponse({'success': True})


@faculty_required
def get_assessment_data(request, assignment_id):
    assessment = get_object_or_404(Assessment, id=assignment_id, course__faculty=request.user)
    questions = []
    for q in assessment.questions.order_by('order'):
        subs = []
        for sq in q.sub_questions.order_by('order'):
            subs.append({
                'id': sq.id,
                'text': sq.text,
                'max_marks': sq.max_marks,
                'image_url':  request.build_absolute_uri(sq.image.url) if sq.image else None,
                'image_name': sq.image.name if sq.image else None,
                'clo_ids': list(sq.clos.values_list('id', flat=True)),
                'plo_ids': list(sq.plos.values_list('id', flat=True)),
            })
        questions.append({
            'id': q.id,
            'text': q.text,
            'max_marks': q.max_marks,
            'image_url':  request.build_absolute_uri(q.image.url) if q.image else None,
            'image_name': q.image.name if q.image else None,
            'clo_ids': list(q.clos.values_list('id', flat=True)),
            'plo_ids': list(q.plos.values_list('id', flat=True)),
            'sub_questions': subs,
        })
    return JsonResponse({
        'id': assessment.id,
        'title': assessment.title,
        'description': assessment.description,
        'assessment_type': assessment.assessment_type,
        'due_date': str(assessment.due_date) if assessment.due_date else '',
        'status': assessment.status,
        'grace_period_hours':   assessment.grace_period_hours,
        'max_late_days':        assessment.max_late_days,
        'late_deduction_type':  assessment.late_deduction_type,
        'late_deduction_value': assessment.late_deduction_value,
        'section_ids': list(assessment.sections.values_list('id', flat=True)),
        'questions': questions,
    })


@faculty_required
def edit_assessment(request, assignment_id):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    assessment = get_object_or_404(Assessment, id=assignment_id, course__faculty=request.user)
    data = json.loads(request.body)
    title = data.get('title', '').strip()
    if not title:
        return JsonResponse({'error': 'Title is required.'}, status=400)

    assessment.title                = title
    assessment.description          = data.get('description', '').strip()
    assessment.due_date             = data.get('due_date') or None
    assessment.grace_period_hours   = int(data.get('grace_period_hours', 0) or 0)
    assessment.max_late_days        = int(data.get('max_late_days', 0) or 0)
    assessment.late_deduction_type  = data.get('late_deduction_type', 'percent')
    assessment.late_deduction_value = float(data.get('late_deduction_value', 0) or 0)
    assessment.save()

    # Rebuild all questions from submitted data
    assessment.questions.all().delete()
    total = 0
    for i, q in enumerate(data.get('questions', []), 1):
        question = Question.objects.create(
            assessment=assessment, order=i, text=q['text'], max_marks=0,
        )
        img_file = _decode_image(q.get('image_b64'), f'q{i}')
        if img_file:
            question.image.save(img_file.name, img_file, save=False)
        elif q.get('existing_image_name'):
            question.image.name = q['existing_image_name']

        sub_questions = q.get('sub_questions', [])
        if sub_questions:
            all_clo_ids, all_plo_ids, q_total = set(), set(), 0
            for j, sq in enumerate(sub_questions, 1):
                sub = SubQuestion.objects.create(
                    question=question, order=j,
                    text=sq['text'], max_marks=int(sq.get('max_marks', 5)),
                )
                sq_img = _decode_image(sq.get('image_b64'), f'sq{i}_{j}')
                if sq_img:
                    sub.image.save(sq_img.name, sq_img, save=False)
                elif sq.get('existing_image_name'):
                    sub.image.name = sq['existing_image_name']
                if sq.get('clo_ids'):
                    sub.clos.set(CLO.objects.filter(id__in=sq['clo_ids']))
                    all_clo_ids.update(sq['clo_ids'])
                if sq.get('plo_ids'):
                    sub.plos.set(PLO.objects.filter(id__in=sq['plo_ids']))
                    all_plo_ids.update(sq['plo_ids'])
                sub.save()
                q_total += int(sq.get('max_marks', 5))
            question.max_marks = q_total
            if all_clo_ids:
                question.clos.set(CLO.objects.filter(id__in=all_clo_ids))
            if all_plo_ids:
                question.plos.set(PLO.objects.filter(id__in=all_plo_ids))
            total += q_total
        else:
            question.max_marks = int(q.get('max_marks', 0))
            if q.get('clo_ids'):
                question.clos.set(CLO.objects.filter(id__in=q['clo_ids']))
            if q.get('plo_ids'):
                question.plos.set(PLO.objects.filter(id__in=q['plo_ids']))
            total += question.max_marks
        question.save()

    manual_total = int(data.get('total_marks', 0))
    assessment.total_marks = total if total > 0 else manual_total
    assessment.save()
    section_ids = data.get('section_ids', [])
    assessment.sections.set(Section.objects.filter(id__in=section_ids, course=assessment.course))
    return JsonResponse({'success': True})


@faculty_required
def publish_assessment(request, assignment_id):
    assessment = get_object_or_404(
        Assessment, id=assignment_id, course__faculty=request.user
    )
    assessment.status = 'published'
    assessment.save()
    notify_new_assignment(assessment)
 
    return JsonResponse({'success': True})


# Student Assignment Views 

@student_required
def student_assignments(request):
    enrollments = Enrollment.objects.filter(student=request.user).select_related('course')
    course_cards = []
    for e in enrollments:
        course = e.course
        total = Assessment.objects.filter(course=course, status='published').count()
        submitted = Submission.objects.filter(
            student=request.user, assessment__course=course
        ).count()
        if total > 0:
            course_cards.append({
                'course':    course,
                'total':     total,
                'submitted': submitted,
                'pending':   total - submitted,
            })
    return render(request, 'student/assignments.html', {'course_cards': course_cards})


@student_required
def student_course_assignments(request, course_id):
    enrolled_ids = list(Enrollment.objects.filter(
        student=request.user
    ).values_list('course_id', flat=True))
    if course_id not in enrolled_ids:
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden()
    course = get_object_or_404(Course, id=course_id)
    student_sections = Section.objects.filter(course=course, students=request.user)
    from django.db.models import Q
    assessments = Assessment.objects.filter(
        course=course, status='published'
    ).filter(
        Q(sections__isnull=True) | Q(sections__in=student_sections)
    ).distinct().prefetch_related('questions__clos', 'questions__plos').order_by('-created_at')
    submissions = {
        s.assessment_id: s
        for s in Submission.objects.filter(student=request.user, assessment__in=assessments)
    }
    assignments_with_status = []
    for a in assessments:
        sub    = submissions.get(a.id)
        window = check_submission_window(a) if not sub else None
        assignments_with_status.append((a, sub, window))
    return render(request, 'student/course_assignments.html', {
        'course':                  course,
        'assignments_with_status': assignments_with_status,
    })


@student_required
def submit_assignment(request, assignment_id):
    from django.shortcuts import redirect as _redirect
    assignment = get_object_or_404(Assessment, id=assignment_id, status='published')

    if not Enrollment.objects.filter(student=request.user, course=assignment.course).exists():
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden()

    if assignment.course.is_archived:
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("This course is archived. Submissions are no longer accepted.")

    existing_sub = Submission.objects.filter(student=request.user, assessment=assignment).first()
    window       = check_submission_window(assignment)
    questions    = assignment.questions.prefetch_related(
        'clos', 'plos',
        'sub_questions__clos', 'sub_questions__plos',
    ).order_by('order')

    error = None

    if request.method == 'POST':
        if existing_sub:
            error = 'You have already submitted this assignment.'
        elif not window['can_submit']:
            error = window['window_msg']
        else:
            content       = request.POST.get('content', '').strip()
            uploaded_file = request.FILES.get('submitted_file')
            if not content and not uploaded_file:
                error = 'Please provide an answer or upload a file.'
            else:
                sub = Submission.objects.create(
                    student=request.user,
                    assessment=assignment,
                    content=content,
                    submitted_file=uploaded_file,
                )
                apply_late_deduction(sub)
                return _redirect('student_assignments')

    return render(request, 'student/submit_assignment.html', {
        'assignment':   assignment,
        'questions':    questions,
        'existing_sub': existing_sub,
        'window':       window,
        'error':        error,
    })


@student_required
def unsubmit_assignment(request, assignment_id):
    from django.shortcuts import redirect as _redirect
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)

    assignment = get_object_or_404(Assessment, id=assignment_id, status='published')

    if assignment.course.is_archived:
        return JsonResponse({'error': 'This course is archived.'}, status=403)

    sub = Submission.objects.filter(student=request.user, assessment=assignment).first()

    if not sub:
        return JsonResponse({'error': 'No submission found.'}, status=404)
    if sub.status == 'graded':
        return JsonResponse({'error': 'Cannot unsubmit a graded submission.'}, status=403)

    sub.delete()
    return _redirect('submit_assignment', assignment_id=assignment_id)


@faculty_required
def faculty_marks_sheet(request):
    courses = Course.objects.filter(faculty=request.user, is_active=True)
    selected_course = None

    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id, faculty=request.user, is_active=True)
    elif courses.exists():
        selected_course = courses.first()

    if not selected_course:
        return render(request, 'faculty/marks_sheet.html', {'courses': courses, 'selected_course': None})

    SUB_TYPES = {'mid', 'final'}
    ALPHA = 'abcdefghijklmnopqrstuvwxyz'

    assessments = list(
        Assessment.objects.filter(course=selected_course, status='published')
        .prefetch_related(
            'questions__clos', 'questions__plos',
            'questions__sub_questions__clos', 'questions__sub_questions__plos',
        )
        .order_by('assessment_type', 'created_at')
    )

    all_columns = []
    assessment_groups = []
    for a in assessments:
        questions = list(a.questions.all().order_by('order'))
        if not questions:
            continue
        use_subs = a.assessment_type in SUB_TYPES
        cols_for_a = []
        for q in questions:
            subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
            if use_subs and subs:
                for idx, sq in enumerate(subs):
                    col = {
                        'label':     f"Q{q.order} ({ALPHA[idx]})",
                        'is_sub':    True,
                        'col_id':    f"sq_{sq.id}",
                        'col_type':  'sub',
                        'entity_id': sq.id,
                        'question':  q,
                        'sub_q':     sq,
                        'max_marks': sq.max_marks,
                        'clo_ids':   [c.id for c in sq.clos.all()],
                        'plo_ids':   [p.id for p in sq.plos.all()],
                        'clo_codes': [c.code for c in sq.clos.all()],
                        'plo_codes': [p.code for p in sq.plos.all()],
                    }
                    cols_for_a.append(col)
                    all_columns.append(col)
            else:
                col = {
                    'label':     f"Q{q.order}",
                    'is_sub':    False,
                    'col_id':    f"q_{q.id}",
                    'col_type':  'q',
                    'entity_id': q.id,
                    'question':  q,
                    'sub_q':     None,
                    'max_marks': q.max_marks,
                    'clo_ids':   [c.id for c in q.clos.all()],
                    'plo_ids':   [p.id for p in q.plos.all()],
                    'clo_codes': [c.code for c in q.clos.all()],
                    'plo_codes': [p.code for p in q.plos.all()],
                }
                cols_for_a.append(col)
                all_columns.append(col)
        assessment_groups.append({'assessment': a, 'questions': questions, 'col_count': len(cols_for_a)})

    students = list(
        User.objects.filter(enrollments__course=selected_course)
        .distinct().order_by('full_name', 'username')
    )

    # Fetch question-level grades (non-sub columns)
    q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
    q_grade_map = {}
    if q_ids:
        for g in QuestionGrade.objects.filter(
            question_id__in=q_ids,
            submission__assessment__course=selected_course
        ).select_related('submission__student'):
            q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained

    # Fetch sub-question-level grades (mid/final sub columns)
    sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
    sq_grade_map = {}
    if sq_ids:
        for g in SubQuestionGrade.objects.filter(
            sub_question_id__in=sq_ids,
            submission__assessment__course=selected_course
        ).select_related('submission__student'):
            sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

    clos = list(selected_course.clos.all().order_by('code'))
    plo_ids_used = set()
    for col in all_columns:
        plo_ids_used.update(col['plo_ids'])
    plos = list(PLO.objects.filter(id__in=plo_ids_used).order_by('code'))

    clo_max = {
        clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids'])
        for clo in clos
    }
    plo_max = {
        plo.id: sum(col['max_marks'] for col in all_columns if plo.id in col['plo_ids'])
        for plo in plos
    }
    total_max = sum(col['max_marks'] for col in all_columns)

    rows = []
    for i, student in enumerate(students):
        sq_m = sq_grade_map.get(student.id, {})
        q_m  = q_grade_map.get(student.id, {})

        def _val(col):
            return sq_m.get(col['entity_id']) if col['is_sub'] else q_m.get(col['entity_id'])

        cells = [
            {
                'col_id':    col['col_id'],
                'col_type':  col['col_type'],
                'entity_id': col['entity_id'],
                'max_marks': col['max_marks'],
                'value':     _val(col),
            }
            for col in all_columns
        ]
        total = sum(c['value'] for c in cells if c['value'] is not None)

        clo_cells = [
            {'clo_id': clo.id, 'code': clo.code, 'max': clo_max[clo.id],
             'raw': sum(
                 (sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0))
                 for col in all_columns if clo.id in col['clo_ids']
             )}
            for clo in clos
        ]
        plo_cells = [
            {'plo_id': plo.id, 'code': plo.code, 'max': plo_max[plo.id],
             'raw': sum(
                 (sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0))
                 for col in all_columns if plo.id in col['plo_ids']
             )}
            for plo in plos
        ]
        rows.append({
            'sl': i + 1,
            'student': student,
            'cells':   cells,
            'total':   total,
            'clo_cells': clo_cells,
            'plo_cells': plo_cells,
        })

    clo_max_list = [{'clo': c, 'max': clo_max[c.id]} for c in clos]
    plo_max_list = [{'plo': p, 'max': plo_max[p.id]} for p in plos]
    t2_colspan = 2 + 3 * (len(clos) + len(plos))

    js_columns = json.dumps([
        {
            'col_id':    col['col_id'],
            'col_type':  col['col_type'],
            'entity_id': col['entity_id'],
            'max':       col['max_marks'],
            'clo_ids':   col['clo_ids'],
            'plo_ids':   col['plo_ids'],
        }
        for col in all_columns
    ])
    js_clos = json.dumps([{'id': c.id, 'code': c.code, 'max': clo_max[c.id]} for c in clos])
    js_plos = json.dumps([{'id': p.id, 'code': p.code, 'max': plo_max[p.id]} for p in plos])

    return render(request, 'faculty/marks_sheet.html', {
        'courses':           courses,
        'selected_course':   selected_course,
        'assessment_groups': assessment_groups,
        'all_columns':       all_columns,
        'rows':              rows,
        'clos':              clos,
        'plos':              plos,
        'clo_max':           clo_max,
        'plo_max':           plo_max,
        'total_max':         total_max,
        'clo_max_list':      clo_max_list,
        'plo_max_list':      plo_max_list,
        't2_colspan':        t2_colspan,
        'js_columns':        js_columns,
        'js_clos':           js_clos,
        'js_plos':           js_plos,
    })


def _compute_marks_sheet_readonly(selected_course, section=None):
    """Returns the marks sheet context dict for read-only views (DAO/Admin/DeptHead)."""
    SUB_TYPES = {'mid', 'final'}
    ALPHA = 'abcdefghijklmnopqrstuvwxyz'

    assessments = list(
        Assessment.objects.filter(course=selected_course, status='published')
        .prefetch_related(
            'questions__clos', 'questions__plos',
            'questions__sub_questions__clos', 'questions__sub_questions__plos',
        )
        .order_by('assessment_type', 'created_at')
    )

    all_columns = []
    assessment_groups = []
    for a in assessments:
        questions = list(a.questions.all().order_by('order'))
        if not questions:
            continue
        use_subs = a.assessment_type in SUB_TYPES
        cols_for_a = []
        for q in questions:
            subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
            if use_subs and subs:
                for idx, sq in enumerate(subs):
                    col = {
                        'label': f"Q{q.order} ({ALPHA[idx]})", 'is_sub': True,
                        'col_id': f"sq_{sq.id}", 'entity_id': sq.id,
                        'max_marks': sq.max_marks,
                        'clo_ids': [c.id for c in sq.clos.all()],
                        'plo_ids': [p.id for p in sq.plos.all()],
                        'clo_codes': [c.code for c in sq.clos.all()],
                        'plo_codes': [p.code for p in sq.plos.all()],
                    }
                    cols_for_a.append(col)
                    all_columns.append(col)
            else:
                col = {
                    'label': f"Q{q.order}", 'is_sub': False,
                    'col_id': f"q_{q.id}", 'entity_id': q.id,
                    'max_marks': q.max_marks,
                    'clo_ids': [c.id for c in q.clos.all()],
                    'plo_ids': [p.id for p in q.plos.all()],
                    'clo_codes': [c.code for c in q.clos.all()],
                    'plo_codes': [p.code for p in q.plos.all()],
                }
                cols_for_a.append(col)
                all_columns.append(col)
        assessment_groups.append({'assessment': a, 'questions': questions, 'col_count': len(cols_for_a)})

    if section is not None:
        students = list(section.students.all().order_by('full_name', 'username'))
    else:
        students = list(
            User.objects.filter(enrollments__course=selected_course)
            .distinct().order_by('full_name', 'username')
        )

    q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
    q_grade_map = {}
    if q_ids:
        for g in QuestionGrade.objects.filter(
            question_id__in=q_ids,
            submission__assessment__course=selected_course
        ).select_related('submission__student'):
            q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained

    sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
    sq_grade_map = {}
    if sq_ids:
        for g in SubQuestionGrade.objects.filter(
            sub_question_id__in=sq_ids,
            submission__assessment__course=selected_course
        ).select_related('submission__student'):
            sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

    clos = list(selected_course.clos.all().order_by('code'))
    plo_ids_used = set()
    for col in all_columns:
        plo_ids_used.update(col['plo_ids'])
    plos = list(PLO.objects.filter(id__in=plo_ids_used).order_by('code'))

    clo_max = {clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids']) for clo in clos}
    plo_max = {plo.id: sum(col['max_marks'] for col in all_columns if plo.id in col['plo_ids']) for plo in plos}
    total_max = sum(col['max_marks'] for col in all_columns)

    rows = []
    for i, student in enumerate(students):
        sq_m = sq_grade_map.get(student.id, {})
        q_m  = q_grade_map.get(student.id, {})

        def _val(col):
            return sq_m.get(col['entity_id']) if col['is_sub'] else q_m.get(col['entity_id'])

        cells = [{'col_id': col['col_id'], 'max_marks': col['max_marks'], 'value': _val(col)} for col in all_columns]
        total = sum(c['value'] for c in cells if c['value'] is not None)

        clo_cells = []
        for clo in clos:
            raw = sum(
                (sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0))
                for col in all_columns if clo.id in col['clo_ids']
            )
            mx = clo_max[clo.id]
            pct = round(raw / mx * 100, 2) if mx > 0 else None
            att = ('Yes' if pct >= 40 else 'No') if pct is not None else '—'
            clo_cells.append({'clo_id': clo.id, 'code': clo.code, 'max': mx, 'raw': raw, 'pct': pct, 'att': att})

        plo_cells = []
        for plo in plos:
            raw = sum(
                (sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0))
                for col in all_columns if plo.id in col['plo_ids']
            )
            mx = plo_max[plo.id]
            pct = round(raw / mx * 100, 2) if mx > 0 else None
            att = ('Yes' if pct >= 40 else 'No') if pct is not None else '—'
            plo_cells.append({'plo_id': plo.id, 'code': plo.code, 'max': mx, 'raw': raw, 'pct': pct, 'att': att})

        rows.append({'sl': i + 1, 'student': student, 'cells': cells, 'total': total,
                     'clo_cells': clo_cells, 'plo_cells': plo_cells})

    clo_max_list = [{'clo': c, 'max': clo_max[c.id]} for c in clos]
    plo_max_list = [{'plo': p, 'max': plo_max[p.id]} for p in plos]
    t2_colspan = 2 + 3 * (len(clos) + len(plos))

    return {
        'assessment_groups': assessment_groups,
        'all_columns': all_columns,
        'rows': rows,
        'clos': clos,
        'plos': plos,
        'total_max': total_max,
        'clo_max_list': clo_max_list,
        'plo_max_list': plo_max_list,
        't2_colspan': t2_colspan,
    }


@faculty_required
def update_question_grade(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    data = json.loads(request.body)
    try:
        marks = float(data.get('marks', 0))
    except (TypeError, ValueError):
        marks = 0

    student = get_object_or_404(User, id=data.get('student_id'))
    col_type = data.get('col_type', 'q')

    if col_type == 'sub':
        sub_q = get_object_or_404(SubQuestion, id=data.get('entity_id'))
        if request.user not in sub_q.question.assessment.course.faculty.all():
            return JsonResponse({'error': 'Forbidden'}, status=403)
        marks = min(max(marks, 0), sub_q.max_marks)
        submission, _ = Submission.objects.get_or_create(
            student=student, assessment=sub_q.question.assessment,
            defaults={'content': '', 'status': 'graded', 'total_score': 0,
                      'plagiarism_score': 0, 'ai_content_score': 0}
        )
        SubQuestionGrade.objects.update_or_create(
            submission=submission, sub_question=sub_q,
            defaults={'marks_obtained': marks}
        )
        # Roll up sub-question totals to QuestionGrade
        sq_total = sum(
            g.marks_obtained
            for g in SubQuestionGrade.objects.filter(
                submission=submission, sub_question__question=sub_q.question
            )
        )
        QuestionGrade.objects.update_or_create(
            submission=submission, question=sub_q.question,
            defaults={'marks_obtained': sq_total}
        )
    else:
        # Backward-compatible: accept entity_id or question_id
        qid = data.get('entity_id') or data.get('question_id')
        question = get_object_or_404(Question, id=qid,
                                      assessment__course__faculty=request.user)
        marks = min(max(marks, 0), question.max_marks)
        submission, _ = Submission.objects.get_or_create(
            student=student, assessment=question.assessment,
            defaults={'content': '', 'status': 'graded', 'total_score': 0,
                      'plagiarism_score': 0, 'ai_content_score': 0}
        )
        QuestionGrade.objects.update_or_create(
            submission=submission, question=question,
            defaults={'marks_obtained': marks}
        )

    total = sum(g.marks_obtained for g in QuestionGrade.objects.filter(submission=submission))
    submission.total_score = total
    if submission.status == 'submitted':
        submission.status = 'graded'
    submission.save()
    return JsonResponse({'success': True, 'marks': marks, 'total': total})


@student_required
def get_unread_count(request):
    count = Notification.objects.filter(recipient=request.user, is_read=False).count()
    return JsonResponse({'count': count})
 
 
@student_required
def mark_all_read(request):
    Notification.objects.filter(recipient=request.user, is_read=False).update(is_read=True)
    return JsonResponse({'success': True})




# QUESTION BANK VIEWS



@faculty_required
def faculty_question_bank(request):
    from .models import PastPaper
    courses = Course.objects.filter(faculty=request.user, is_active=True)
    papers  = PastPaper.objects.filter(
        uploaded_by=request.user
    ).prefetch_related('questions', 'allowed_courses')

    # Build assessment question bank: course → type groups → assessments → questions
    TYPE_ORDER  = ['assignment', 'quiz', 'mid', 'ct', 'final', 'lab']
    TYPE_LABELS = dict(Assessment.TYPE_CHOICES)

    def type_sort_key(t):
        try:
            return TYPE_ORDER.index(t)
        except ValueError:
            return len(TYPE_ORDER)

    course_bank = []
    for course in courses:
        assessments = Assessment.objects.filter(course=course).prefetch_related(
            'questions__clos', 'questions__plos',
            'questions__sub_questions__clos',
            'questions__sub_questions__plos',
        ).order_by('assessment_type', 'created_at')

        type_map = defaultdict(list)
        for a in assessments:
            qs = list(a.questions.all())
            if qs:
                type_map[a.assessment_type].append({'assessment': a, 'questions': qs})

        if not type_map:
            continue

        # Add num_label so template can show "Quiz 1", "Quiz 2" etc.
        for t, entries in type_map.items():
            lbl = TYPE_LABELS.get(t, t.replace('_', ' ').title())
            total = len(entries)
            for i, entry in enumerate(entries, 1):
                entry['num_label'] = f"{lbl} {i}" if total > 1 else lbl

        type_groups = [
            {
                'type_key':   t,
                'type_label': TYPE_LABELS.get(t, t.replace('_', ' ').title()),
                'entries':    type_map[t],
                'q_count':    sum(len(e['questions']) for e in type_map[t]),
            }
            for t in sorted(type_map.keys(), key=type_sort_key)
        ]
        course_bank.append({
            'course':      course,
            'type_groups': type_groups,
            'total_q':     sum(g['q_count'] for g in type_groups),
        })

    return render(request, 'faculty/question_bank.html', {
        'papers':      papers,
        'courses':     courses,
        'course_bank': course_bank,
    })


@faculty_required
def create_past_paper(request):
    from .models import PastPaper, PastPaperQuestion
    if request.method == 'POST':
        data = json.loads(request.body)
        paper = PastPaper.objects.create(
            title         = data['title'],
            course_code   = data['course_code'],
            course_name   = data['course_name'],
            semester      = data['semester'],
            exam_type     = data['exam_type'],
            total_marks   = int(data.get('total_marks', 0)),
            duration_mins = int(data.get('duration_mins', 0)),
            description   = data.get('description', ''),
            is_public     = data.get('is_public', False),
            uploaded_by   = request.user,
        )
        course_ids = data.get('allowed_course_ids', [])
        if course_ids:
            paper.allowed_courses.set(
                Course.objects.filter(id__in=course_ids, faculty=request.user, is_active=True)
            )
        total = 0
        for i, q in enumerate(data.get('questions', []), 1):
            PastPaperQuestion.objects.create(
                paper       = paper,
                order       = i,
                text        = q['text'],
                marks       = int(q.get('marks', 0)),
                answer_hint = q.get('answer_hint', ''),
                show_hint   = q.get('show_hint', False),
                topic_tag   = q.get('topic_tag', ''),
                difficulty  = q.get('difficulty', ''),
            )
            total += int(q.get('marks', 0))
        if total > 0:
            paper.total_marks = total
            paper.save(update_fields=['total_marks'])
        return JsonResponse({'success': True, 'id': paper.id})
    return JsonResponse({'error': 'POST required'}, status=400)


@faculty_required
def delete_past_paper(request, paper_id):
    from .models import PastPaper
    paper = get_object_or_404(PastPaper, id=paper_id, uploaded_by=request.user)
    paper.delete()
    return JsonResponse({'success': True})


@faculty_required
def toggle_paper_visibility(request, paper_id):
    from .models import PastPaper
    paper = get_object_or_404(PastPaper, id=paper_id, uploaded_by=request.user)
    paper.is_public = not paper.is_public
    paper.save(update_fields=['is_public'])
    return JsonResponse({'success': True, 'is_public': paper.is_public})


@faculty_required
def toggle_hint_visibility(request, question_id):
    from .models import PastPaperQuestion
    q = get_object_or_404(PastPaperQuestion, id=question_id,
                          paper__uploaded_by=request.user)
    q.show_hint = not q.show_hint
    q.save(update_fields=['show_hint'])
    return JsonResponse({'success': True, 'show_hint': q.show_hint})


@student_required
def student_question_bank(request):
    from .models import PastPaper
    from django.db.models import Q

    enrollments = Enrollment.objects.filter(student=request.user).select_related('course')
    enrolled_ids = [e.course_id for e in enrollments]

    # Build course bank for the Assessment Questions tab
    course_bank = []
    for e in enrollments:
        course = e.course
        assessments = Assessment.objects.filter(
            course=course, status='published'
        ).prefetch_related('questions').order_by('assessment_type', 'created_at')
        type_set = set()
        q_total = 0
        for a in assessments:
            qs = a.questions.count()
            if qs > 0:
                type_set.add(a.assessment_type)
                q_total += qs
        if q_total > 0:
            course_bank.append({
                'course':      course,
                'type_count':  len(type_set),
                'q_total':     q_total,
            })

    # Past papers (existing logic unchanged)
    papers = PastPaper.objects.filter(
        Q(is_public=True) | Q(allowed_courses__id__in=enrolled_ids)
    ).distinct().prefetch_related('questions')

    search      = request.GET.get('q', '').strip()
    exam_type   = request.GET.get('type', '')
    semester    = request.GET.get('semester', '')
    course_code = request.GET.get('course', '')
    difficulty  = request.GET.get('difficulty', '')

    if search:
        papers = papers.filter(
            Q(title__icontains=search) |
            Q(course_code__icontains=search) |
            Q(course_name__icontains=search) |
            Q(questions__text__icontains=search) |
            Q(questions__topic_tag__icontains=search)
        ).distinct()
    if exam_type:
        papers = papers.filter(exam_type=exam_type)
    if semester:
        papers = papers.filter(semester__icontains=semester)
    if course_code:
        papers = papers.filter(course_code__icontains=course_code)
    if difficulty:
        papers = papers.filter(questions__difficulty=difficulty).distinct()

    all_accessible = PastPaper.objects.filter(
        Q(is_public=True) | Q(allowed_courses__id__in=enrolled_ids)
    ).distinct()

    return render(request, 'student/question_bank.html', {
        'course_bank':  course_bank,
        'papers':       papers,
        'total_count':  papers.count(),
        'course_codes': sorted(set(all_accessible.values_list('course_code', flat=True))),
        'semesters':    sorted(set(all_accessible.values_list('semester', flat=True)), reverse=True),
        'search':       search,
        'exam_type':    exam_type,
        'semester':     semester,
        'course_code':  course_code,
        'difficulty':   difficulty,
        'active_tab':   request.GET.get('tab', 'assessments'),
    })


@student_required
def student_view_paper(request, paper_id):
    from .models import PastPaper
    from django.db.models import Q
    enrolled_ids = Enrollment.objects.filter(
        student=request.user
    ).values_list('course_id', flat=True)

    paper = get_object_or_404(PastPaper, id=paper_id)

    # Permission check
    if not paper.is_public:
        if not paper.allowed_courses.filter(id__in=enrolled_ids).exists():
            from django.http import HttpResponseForbidden
            return HttpResponseForbidden("You don't have access to this paper.")

    return render(request, 'student/view_paper.html', {
        'paper':     paper,
        'questions': paper.questions.all(),
    })


@student_required
def student_qbank_course(request, course_id):
    enrolled_ids = list(Enrollment.objects.filter(
        student=request.user
    ).values_list('course_id', flat=True))

    if course_id not in enrolled_ids:
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden()

    course = get_object_or_404(Course, id=course_id)

    TYPE_ORDER  = ['assignment', 'quiz', 'mid', 'ct', 'final', 'lab']
    TYPE_LABELS = dict(Assessment.TYPE_CHOICES)

    def type_sort_key(t):
        try:
            return TYPE_ORDER.index(t)
        except ValueError:
            return len(TYPE_ORDER)

    type_map = {}
    for a in Assessment.objects.filter(course=course, status='published').prefetch_related('questions'):
        qs = a.questions.count()
        if qs > 0:
            if a.assessment_type not in type_map:
                type_map[a.assessment_type] = {'a_count': 0, 'q_count': 0}
            type_map[a.assessment_type]['a_count'] += 1
            type_map[a.assessment_type]['q_count'] += qs

    type_groups = [
        {
            'type_key':   t,
            'type_label': TYPE_LABELS.get(t, t.replace('_', ' ').title()),
            'a_count':    type_map[t]['a_count'],
            'q_count':    type_map[t]['q_count'],
        }
        for t in sorted(type_map.keys(), key=type_sort_key)
    ]

    return render(request, 'student/qbank_course.html', {
        'course':      course,
        'type_groups': type_groups,
    })


@student_required
def student_qbank_type(request, course_id, atype):
    enrolled_ids = list(Enrollment.objects.filter(
        student=request.user
    ).values_list('course_id', flat=True))

    if course_id not in enrolled_ids:
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden()

    course = get_object_or_404(Course, id=course_id)
    TYPE_LABELS = dict(Assessment.TYPE_CHOICES)
    type_label  = TYPE_LABELS.get(atype, atype.replace('_', ' ').title())

    assessments = Assessment.objects.filter(
        course=course, assessment_type=atype, status='published'
    ).prefetch_related(
        'questions__clos',
        'questions__plos',
        'questions__sub_questions__clos',
        'questions__sub_questions__plos',
    ).order_by('created_at')

    all_with_q = [a for a in assessments if a.questions.exists()]
    total = len(all_with_q)
    entries = []
    for idx, a in enumerate(all_with_q, 1):
        entries.append({
            'assessment': a,
            'label':      f"{type_label} {idx}" if total > 1 else type_label,
            'questions':  list(a.questions.all().order_by('order')),
        })

    return render(request, 'student/qbank_type.html', {
        'course':     course,
        'atype':      atype,
        'type_label': type_label,
        'entries':    entries,
    })

# ADMIN PORTAL VIEWS

def admin_required(view_func):
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('sign_in_html')
        if request.user.role != 'admin' and not request.user.is_superuser:
            return redirect('home')
        return view_func(request, *args, **kwargs)
    wrapper.__name__ = view_func.__name__
    return wrapper


@admin_required
def admin_dashboard(request):
    from datetime import date as dt_date
    return render(request, 'admin_portal/dashboard.html', {
        'today':                dt_date.today().strftime('%B %d, %Y'),
        'total_users':          User.objects.count(),
        'total_faculty':        User.objects.filter(role='faculty').count(),
        'total_students':       User.objects.filter(role='student').count(),
        'total_courses':        Course.objects.count(),
        'total_enrollments':    Enrollment.objects.count(),
        'total_assessments':    Assessment.objects.count(),
        'published_assessments':Assessment.objects.filter(status='published').count(),
        'total_submissions':    Submission.objects.count(),
        'pending_submissions':  Submission.objects.filter(status='submitted').count(),
        'recent_users':         User.objects.order_by('-date_joined')[:6],
        'recent_courses':       Course.objects.prefetch_related('enrollments').order_by('-created_at')[:6],
        'recent_submissions':   Submission.objects.select_related('student', 'assessment__course').order_by('-submitted_at')[:8],
    })


@admin_required
def admin_users(request):
    users = User.objects.all().order_by('-date_joined')
    return render(request, 'admin_portal/users.html', {'users': users})


@admin_required
def admin_create_user(request):
    if request.method == 'POST':
        full_name = request.POST.get('full_name', '').strip()
        email     = request.POST.get('email', '').strip()
        password  = request.POST.get('password', '')
        role      = request.POST.get('role', 'student')
        if User.objects.filter(email=email).exists():
            messages.error(request, 'Email already registered.')
            return redirect('admin_users')
        if len(password) < 8:
            messages.error(request, 'Password must be at least 8 characters.')
            return redirect('admin_users')
        username = email.split('@')[0]
        base = username; i = 1
        while User.objects.filter(username=username).exists():
            username = f"{base}{i}"; i += 1
        user = User.objects.create_user(
            username=username, email=email, password=password,
            full_name=full_name, role=role,
        )
        if role == 'admin':
            user.is_staff = True
            user.save()
        messages.success(request, f'User {full_name} ({role}) created successfully.')
    return redirect('admin_users')


@admin_required
def admin_edit_user(request, user_id):
    u = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        u.full_name = request.POST.get('full_name', u.full_name).strip()
        u.email     = request.POST.get('email', u.email).strip()
        u.role      = request.POST.get('role', u.role)
        pw = request.POST.get('password', '').strip()
        if pw:
            if len(pw) < 8:
                messages.error(request, 'Password must be at least 8 characters.')
                return redirect('admin_users')
            u.set_password(pw)
        u.is_staff = (u.role == 'admin')
        u.save()
        messages.success(request, f'User {u.full_name} updated.')
    return redirect('admin_users')


@admin_required
def admin_toggle_user(request, user_id):
    if request.method == 'POST':
        u = get_object_or_404(User, id=user_id)
        u.is_active = not u.is_active
        u.save()
        status = 'activated' if u.is_active else 'deactivated'
        messages.success(request, f'User {u.full_name or u.username} {status}.')
    return redirect('admin_users')


@admin_required
def admin_delete_user(request, user_id):
    if request.method == 'POST':
        u = get_object_or_404(User, id=user_id)
        name = u.full_name or u.username
        u.delete()
        messages.success(request, f'User {name} deleted.')
    return redirect('admin_users')


@admin_required
def admin_courses(request):
    courses = Course.objects.prefetch_related('faculty', 'enrollments', 'assessments').order_by('-created_at')
    semesters = Course.objects.values_list('semester', flat=True).distinct().order_by('semester')
    return render(request, 'admin_portal/courses.html', {
        'courses': courses, 'semesters': semesters,
    })


@admin_required
def admin_delete_course(request, course_id):
    if request.method == 'POST':
        c = get_object_or_404(Course, id=course_id)
        name = c.code
        c.delete()
        messages.success(request, f'Course {name} deleted.')
    return redirect('admin_courses')


@admin_required
def admin_assessments(request):
    assessments = Assessment.objects.select_related('course').prefetch_related('course__faculty', 'submissions').order_by('-created_at')
    return render(request, 'admin_portal/assessments.html', {'assessments': assessments})


@admin_required
def admin_delete_assessment(request, assessment_id):
    if request.method == 'POST':
        a = get_object_or_404(Assessment, id=assessment_id)
        title = a.title
        a.delete()
        messages.success(request, f'Assessment "{title}" deleted.')
    return redirect('admin_assessments')


@admin_required
def admin_submissions(request):
    submissions = Submission.objects.select_related('student', 'assessment__course').order_by('-submitted_at')
    return render(request, 'admin_portal/submissions.html', {'submissions': submissions})


@admin_required
def admin_announcements(request):
    announcements = Announcement.objects.select_related('course', 'created_by').order_by('-created_at')
    return render(request, 'admin_portal/announcements.html', {'announcements': announcements})


@admin_required
def admin_delete_announcement(request, ann_id):
    if request.method == 'POST':
        a = get_object_or_404(Announcement, id=ann_id)
        a.delete()
        messages.success(request, 'Announcement deleted.')
    return redirect('admin_announcements')


@admin_required
def admin_materials(request):
    materials = StudyMaterial.objects.select_related('course', 'uploaded_by').order_by('-uploaded_at')
    return render(request, 'admin_portal/materials.html', {'materials': materials})


@admin_required
def admin_delete_material(request, material_id):
    if request.method == 'POST':
        m = get_object_or_404(StudyMaterial, id=material_id)
        m.delete()
        messages.success(request, 'Material deleted.')
    return redirect('admin_materials')


@admin_required
def admin_escar(request):
    courses = Course.objects.all().order_by('code')
    selected_course = None
    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
    elif courses.exists():
        selected_course = courses.first()
    rows, total_enrolled, total_participated, total_absent = _compute_escar(selected_course)
    return render(request, 'admin_portal/escar.html', {
        'courses': courses,
        'selected_course': selected_course,
        'rows': rows,
        'total_enrolled': total_enrolled,
        'total_participated': total_participated,
        'total_absent': total_absent,
    })


@admin_required
def admin_analytics(request):
    from django.db.models import Q
    all_semesters = Course.objects.values_list('semester', flat=True).distinct().order_by('semester')
    course_name_q = request.GET.get('course_name', '').strip()
    semester_q    = request.GET.get('semester', '').strip()
    courses = Course.objects.all().order_by('code')
    if course_name_q:
        courses = courses.filter(Q(name__icontains=course_name_q) | Q(code__icontains=course_name_q))
    if semester_q:
        courses = courses.filter(semester=semester_q)
    selected_course = None
    plo_attainment = []
    clo_attainment = []
    total_enrolled = 0
    present_count = 0
    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
    elif courses.exists():
        selected_course = courses.first()
    if selected_course:
        SUB_TYPES = {'mid', 'final'}
        assessments = list(
            Assessment.objects.filter(course=selected_course, status='published')
            .prefetch_related(
                'questions__clos', 'questions__plos',
                'questions__sub_questions__clos', 'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'clo_ids': [c.id for c in sq.clos.all()],
                                            'plo_ids': [p.id for p in sq.plos.all()], 'is_final': is_final})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'clo_ids': [c.id for c in q.clos.all()],
                                        'plo_ids': [p.id for p in q.plos.all()], 'is_final': is_final})
        clos = list(selected_course.clos.all())
        clo_max = {clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids']) for clo in clos}
        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=selected_course).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=selected_course).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained
        students = list(User.objects.filter(enrollments__course=selected_course).distinct().order_by('full_name', 'username'))
        total_enrolled = len(students)
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(col['entity_id'] in (sq_m if col['is_sub'] else q_m) for col in all_columns if col['is_final'])
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)
        clo_achieved = {clo.id: 0 for clo in clos}
        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue
            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
            for clo in clos:
                mx = clo_max[clo.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
                if round(raw / mx * 100, 1) >= 40:
                    clo_achieved[clo.id] += 1
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if round(raw / mx * 100, 1) >= 40:
                    plo_achieved[p.id] += 1
        for clo in clos:
            achieved = clo_achieved[clo.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            clo_attainment.append({'code': clo.code, 'description': clo.description,
                                    'attainment': pct, 'achieved': achieved, 'present': present_count})
        for p in plos:
            achieved = plo_achieved[p.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            plo_attainment.append({'code': p.code, 'description': p.description,
                                    'attainment': pct, 'achieved': achieved, 'present': present_count})
    return render(request, 'admin_portal/analytics.html', {
        'courses': courses, 'selected_course': selected_course,
        'plo_attainment': plo_attainment, 'clo_attainment': clo_attainment,
        'total_enrolled': total_enrolled, 'present_count': present_count,
        'all_semesters': all_semesters, 'course_name_q': course_name_q, 'semester_q': semester_q,
    })


@admin_required
def admin_batch_analytics(request):
    all_batches = list(Section.objects.values_list('batch', flat=True).distinct().order_by('batch'))
    selected_batch = request.GET.get('batch', '').strip()
    if not selected_batch and all_batches:
        selected_batch = all_batches[0]
    result = None
    if selected_batch:
        result = _compute_batch_plo_attainment(selected_batch)
    return render(request, 'admin_portal/batch_analytics.html', {
        'all_batches': all_batches, 'selected_batch': selected_batch, 'result': result,
    })


@admin_required
def admin_plo_comparison(request):
    all_courses = Course.objects.all().order_by('code', 'semester')
    seen_codes = set()
    course_codes = []
    for c in all_courses:
        if c.code not in seen_codes:
            seen_codes.add(c.code)
            course_codes.append({'code': c.code, 'name': c.name})
    selected_code = request.GET.get('course_code', '').strip()
    selected_ids  = request.GET.getlist('sem_ids')
    courses_for_code = []
    if selected_code:
        courses_for_code = list(all_courses.filter(code=selected_code))
    elif course_codes:
        selected_code = course_codes[0]['code']
        courses_for_code = list(all_courses.filter(code=selected_code))
    if selected_ids:
        compare_set = set(selected_ids)
        courses_to_compare = [c for c in courses_for_code if str(c.id) in compare_set]
    else:
        courses_to_compare = list(courses_for_code)
    SUB_TYPES = {'mid', 'final'}
    comparison_data = []
    all_plo_meta = {}
    for course in courses_to_compare:
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related('questions__plos', 'questions__sub_questions__plos')
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'plo_ids': [p.id for p in sq.plos.all()], 'is_final': is_final})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'plo_ids': [p.id for p in q.plos.all()], 'is_final': is_final})
        if not all_columns:
            comparison_data.append({'course': course, 'semester': course.semester, 'plo_attainment': {}, 'present_count': 0})
            continue
        plo_ids_used = set(pid for col in all_columns for pid in col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        for p in plos:
            all_plo_meta[p.code] = p.description
        q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=course).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=course).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained
        students = list(User.objects.filter(enrollments__course=course).distinct())
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(col['entity_id'] in (sq_m if col['is_sub'] else q_m) for col in all_columns if col['is_final'])
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)
        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue
            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if present_count > 0 and (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1
        plo_attainment = {p.code: (round(plo_achieved[p.id] / present_count * 100, 1) if present_count > 0 else 0.0) for p in plos}
        comparison_data.append({'course': course, 'semester': course.semester, 'plo_attainment': plo_attainment, 'present_count': present_count})
    def _sem_key(entry):
        parts = entry['semester'].split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))
    comparison_data.sort(key=_sem_key)
    plo_codes = sorted(all_plo_meta.keys())
    table_rows = []
    for entry in comparison_data:
        cells = [{'plo_code': code, 'attainment': entry['plo_attainment'].get(code),
                  'has_data': entry['plo_attainment'].get(code) is not None} for code in plo_codes]
        table_rows.append({'semester': entry['semester'], 'course_id': entry['course'].id,
                           'is_archived': entry['course'].is_archived, 'present_count': entry['present_count'], 'cells': cells})
    chart_json = json.dumps({
        'plo_codes': plo_codes, 'plo_descriptions': all_plo_meta,
        'semesters': [e['semester'] for e in comparison_data],
        'datasets': [{'semester': e['semester'], 'values': [e['plo_attainment'].get(code) for code in plo_codes]} for e in comparison_data],
    })
    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]
    return render(request, 'admin_portal/plo_comparison.html', {
        'course_codes': course_codes, 'selected_code': selected_code,
        'courses_for_code': courses_for_code, 'selected_ids': [int(i) for i in selected_ids if i.isdigit()],
        'plo_codes': plo_codes, 'plo_list': plo_list, 'table_rows': table_rows,
        'chart_json': chart_json, 'has_data': bool(plo_codes and table_rows),
    })


@admin_required
def admin_semester_plo_comparison(request):
    SUB_TYPES = {'mid', 'final'}
    all_courses = list(Course.objects.all().order_by('semester', 'code'))
    def _sem_key(s):
        parts = s.split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))
    all_semesters = sorted(set(c.semester for c in all_courses), key=_sem_key)
    selected_sems = request.GET.getlist('sems')
    if not selected_sems:
        selected_sems = list(all_semesters)
    selected_sems_set = set(selected_sems)
    semester_plo_data = {}
    all_plo_meta = {}
    semester_present_students = {}
    for course in all_courses:
        if course.semester not in selected_sems_set:
            continue
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related('questions__plos', 'questions__sub_questions__plos')
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'plo_ids': [p.id for p in sq.plos.all()], 'is_final': is_final})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'plo_ids': [p.id for p in q.plos.all()], 'is_final': is_final})
        if not all_columns:
            continue
        plo_ids_used = set(pid for col in all_columns for pid in col['plo_ids'])
        if not plo_ids_used:
            continue
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        for p in plos:
            all_plo_meta[p.code] = p.description
        q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=course).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=course).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained
        students = list(User.objects.filter(enrollments__course=course).distinct())
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(col['entity_id'] in (sq_m if col['is_sub'] else q_m) for col in all_columns if col['is_final'])
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)
        if present_count == 0:
            continue
        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue
            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1
        semester = course.semester
        if semester not in semester_plo_data:
            semester_plo_data[semester] = {}
        semester_present_students.setdefault(semester, set()).update(present_student_ids)
        for p in plos:
            if plo_max[p.id] == 0:
                continue
            code = p.code
            if code not in semester_plo_data[semester]:
                semester_plo_data[semester][code] = {'description': p.description, 'total_achieved': 0, 'total_present': 0}
            semester_plo_data[semester][code]['total_achieved'] += plo_achieved[p.id]
            semester_plo_data[semester][code]['total_present'] += present_count
    plo_codes = sorted(all_plo_meta.keys())
    semesters_sorted = sorted(semester_plo_data.keys(), key=_sem_key)
    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]
    table_rows = []
    for sem in semesters_sorted:
        cells = []
        for code in plo_codes:
            pe = semester_plo_data.get(sem, {}).get(code)
            if pe and pe['total_present'] > 0:
                att = round(pe['total_achieved'] / pe['total_present'] * 100, 1)
                cells.append({'plo_code': code, 'attainment': att, 'has_data': True})
            else:
                cells.append({'plo_code': code, 'attainment': 0, 'has_data': False})
        table_rows.append({'semester': sem, 'present_count': len(semester_present_students.get(sem, set())), 'cells': cells})
    chart_json = json.dumps({
        'plo_codes': plo_codes, 'semesters': semesters_sorted,
        'datasets': [
            {'plo': code, 'values': [
                (round(semester_plo_data[sem][code]['total_achieved'] / semester_plo_data[sem][code]['total_present'] * 100, 1)
                 if (sem in semester_plo_data and code in semester_plo_data[sem] and semester_plo_data[sem][code]['total_present'] > 0)
                 else None)
                for sem in semesters_sorted
            ]}
            for code in plo_codes
        ],
    })
    return render(request, 'admin_portal/semester_plo_comparison.html', {
        'all_semesters': all_semesters, 'selected_sems': selected_sems,
        'plo_codes': plo_codes, 'plo_list': plo_list, 'table_rows': table_rows,
        'semesters': semesters_sorted, 'chart_json': chart_json,
        'has_data': bool(plo_codes and semesters_sorted),
    })


@admin_required
def admin_students(request):
    from django.db.models import Q, Count
    q = request.GET.get('q', '').strip()
    course_id = request.GET.get('course', '').strip()
    courses = Course.objects.all().order_by('code')
    selected_course = None
    students = User.objects.filter(role='student')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
        students = students.filter(enrollments__course=selected_course).distinct()
    if q:
        students = students.filter(Q(username__icontains=q) | Q(full_name__icontains=q))
    students = students.annotate(enrolled_count=Count('enrollments', distinct=True)).order_by('username')
    return render(request, 'admin_portal/students.html', {
        'students': students, 'courses': courses,
        'selected_course': selected_course, 'q': q, 'course_id': course_id,
    })


@admin_required
def admin_student_attainment(request, student_id):
    student = get_object_or_404(User, id=student_id, role='student')
    enrollments = Enrollment.objects.filter(student=student).select_related('course')
    results = []
    program_plo_data = {}
    SUB_TYPES = {'mid', 'final'}
    for e in enrollments:
        course = e.course
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related('questions__clos', 'questions__plos',
                              'questions__sub_questions__clos', 'questions__sub_questions__plos')
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'clo_ids': [c.id for c in sq.clos.all()], 'plo_ids': [p.id for p in sq.plos.all()]})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'clo_ids': [c.id for c in q.clos.all()], 'plo_ids': [p.id for p in q.plos.all()]})
        clos = list(course.clos.all())
        clo_max = {clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids']) for clo in clos}
        plo_ids_used = set(pid for col in all_columns for pid in col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        total_max_overall = sum(col['max_marks'] for col in all_columns)
        q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_m, sq_m = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=course, submission__student=student):
                q_m[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=course, submission__student=student):
                sq_m[g.sub_question_id] = g.marks_obtained
        def _mark(col, _sq=sq_m, _q=q_m):
            return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
        has_grades = bool(q_m or sq_m)
        total_raw = sum(_mark(col) for col in all_columns)
        avg_pct = round(total_raw / total_max_overall * 100, 1) if (total_max_overall > 0 and has_grades) else 0.0
        grade = 'F'
        if avg_pct >= 80: grade = 'A+'
        elif avg_pct >= 75: grade = 'A'
        elif avg_pct >= 70: grade = 'A-'
        elif avg_pct >= 65: grade = 'B+'
        elif avg_pct >= 60: grade = 'B'
        elif avg_pct >= 55: grade = 'B-'
        elif avg_pct >= 50: grade = 'C+'
        elif avg_pct >= 45: grade = 'C'
        elif avg_pct >= 40: grade = 'D'
        clo_results = []
        for clo in clos:
            mx = clo_max[clo.id]
            raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
            att = round(raw / mx * 100, 1) if mx > 0 else 0.0
            clo_results.append({'code': clo.code, 'bloom': clo.bloom_level, 'description': clo.description,
                                 'obtained': round(raw, 1), 'total': mx, 'attainment': att})
        plo_results = []
        for p in plos:
            mx = plo_max[p.id]
            if mx == 0:
                continue
            raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
            att = round(raw / mx * 100, 1)
            plo_results.append({'code': p.code, 'description': p.description, 'attainment': att})
            if p.id not in program_plo_data:
                program_plo_data[p.id] = {'code': p.code, 'description': p.description, 'total_max': 0.0, 'total_raw': 0.0}
            program_plo_data[p.id]['total_max'] += mx
            program_plo_data[p.id]['total_raw'] += raw
        graded_count = Submission.objects.filter(student=student, assessment__in=assessments, status__in=['graded', 'flagged']).count()
        results.append({'course': course, 'grade': grade, 'avg_pct': avg_pct,
                        'graded_count': graded_count, 'clo_results': clo_results, 'plo_results': plo_results})
    program_plo_attainment = sorted(
        [{'code': d['code'], 'description': d['description'],
          'attainment': round(d['total_raw'] / d['total_max'] * 100, 1),
          'obtained': round(d['total_raw'], 1), 'total': round(d['total_max'], 1)}
         for d in program_plo_data.values() if d['total_max'] > 0],
        key=lambda x: x['code'],
    )
    return render(request, 'admin_portal/student_attainment.html', {
        'student': student, 'results': results, 'program_plo_attainment': program_plo_attainment,
    })


# DAO PORTAL VIEWS

def dao_required(view_func):
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('sign_in_html')
        if request.user.role != 'dao':
            return redirect('home')
        return view_func(request, *args, **kwargs)
    wrapper.__name__ = view_func.__name__
    return wrapper


@dao_required
def dao_dashboard(request):
    from datetime import date as dt_date
    return render(request, 'dao_portal/dashboard.html', {
        'today':          dt_date.today().strftime('%B %d, %Y'),
        'total_users':    User.objects.count(),
        'total_faculty':  User.objects.filter(role='faculty').count(),
        'total_students': User.objects.filter(role='student').count(),
        'total_courses':  Course.objects.count(),
        'total_enrollments': Enrollment.objects.count(),
        'recent_users':   User.objects.order_by('-date_joined')[:6],
        'recent_courses': Course.objects.prefetch_related('enrollments').order_by('-created_at')[:6],
    })


@dao_required
def dao_users(request):
    users = User.objects.all().order_by('-date_joined')
    return render(request, 'dao_portal/users.html', {'users': users})


@dao_required
def dao_create_user(request):
    if request.method == 'POST':
        full_name = request.POST.get('full_name', '').strip()
        email     = request.POST.get('email', '').strip()
        password  = request.POST.get('password', '')
        role      = request.POST.get('role', 'student')
        if role not in ('faculty', 'student'):
            role = 'student'
        if User.objects.filter(email=email).exists():
            messages.error(request, 'Email already registered.')
            return redirect('dao_users')
        if len(password) < 8:
            messages.error(request, 'Password must be at least 8 characters.')
            return redirect('dao_users')
        username = email.split('@')[0]
        base = username; i = 1
        while User.objects.filter(username=username).exists():
            username = f"{base}{i}"; i += 1
        User.objects.create_user(
            username=username, email=email, password=password,
            full_name=full_name, role=role,
        )
        messages.success(request, f'User {full_name} created successfully.')
    return redirect('dao_users')


@dao_required
def dao_edit_user(request, user_id):
    u = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        u.full_name = request.POST.get('full_name', u.full_name).strip()
        u.email     = request.POST.get('email', u.email).strip()
        new_role    = request.POST.get('role', u.role)
        if new_role in ('faculty', 'student'):
            u.role = new_role
        pw = request.POST.get('password', '').strip()
        if pw:
            if len(pw) < 8:
                messages.error(request, 'Password must be at least 8 characters.')
                return redirect('dao_users')
            u.set_password(pw)
        u.save()
        messages.success(request, f'User {u.full_name} updated.')
    return redirect('dao_users')


@dao_required
def dao_toggle_user(request, user_id):
    if request.method == 'POST':
        u = get_object_or_404(User, id=user_id)
        u.is_active = not u.is_active
        u.save()
        status = 'activated' if u.is_active else 'deactivated'
        messages.success(request, f'User {u.full_name or u.username} {status}.')
    return redirect('dao_users')


@dao_required
def dao_delete_user(request, user_id):
    if request.method == 'POST':
        u = get_object_or_404(User, id=user_id)
        name = u.full_name or u.username
        u.delete()
        messages.success(request, f'User {name} deleted.')
    return redirect('dao_users')


SEMESTER_OPTIONS = [
    'Fall 2023', 'Spring 2024', 'Fall 2024',
    'Spring 2025', 'Fall 2025', 'Spring 2026', 'Fall 2026',
]

@dao_required
def dao_courses(request):
    courses = Course.objects.prefetch_related('faculty', 'enrollments', 'assessments', 'sections').order_by('-created_at')
    semesters = Course.objects.values_list('semester', flat=True).distinct().order_by('semester')
    faculty_list = User.objects.filter(role='faculty').order_by('full_name')
    return render(request, 'dao_portal/courses.html', {
        'courses': courses,
        'semesters': semesters,
        'faculty_list': faculty_list,
        'semester_options': SEMESTER_OPTIONS,
    })


@dao_required
def dao_create_course(request):
    if request.method == 'POST':
        code         = request.POST.get('code', '').strip()
        name         = request.POST.get('name', '').strip()
        description  = request.POST.get('description', '').strip()
        credit_hours = int(request.POST.get('credit_hours', 3) or 3)
        semester     = request.POST.get('semester', 'Fall 2025')
        if not code or not name:
            messages.error(request, 'Course code and name are required.')
            return redirect('dao_courses')
        course = Course.objects.create(
            code=code, name=name, description=description,
            credit_hours=credit_hours, semester=semester,
        )
        # Create inline sections submitted with the form
        sec_names = request.POST.getlist('section_name')
        sec_sems  = request.POST.getlist('section_semester')
        created_sections = 0
        for sec_name, sec_sem in zip(sec_names, sec_sems):
            sec_name = sec_name.strip().upper()
            sec_sem  = sec_sem.strip()
            if sec_name and sec_sem:
                Section.objects.get_or_create(course=course, name=sec_name, batch=sec_sem)
                created_sections += 1
        msg = f'Course {code} created'
        msg += f' with {created_sections} section(s).' if created_sections else '.'
        msg += ' Assign faculty to sections below.'
        messages.success(request, msg)
        return redirect('dao_assign_faculty_page', course_id=course.id)
    return redirect('dao_courses')


@dao_required
def dao_assign_faculty_page(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    faculty_list = User.objects.filter(role='faculty').order_by('full_name')
    sections = list(course.sections.prefetch_related('faculty').order_by('batch', 'name'))

    if request.method == 'POST':
        for section in sections:
            faculty_id = request.POST.get(f'faculty_{section.id}', '').strip()
            if faculty_id:
                faculty = User.objects.filter(id=faculty_id, role='faculty').first()
                if faculty:
                    section.faculty.set([faculty])  # replaces any previous assignment
            else:
                section.faculty.clear()
        # Sync course.faculty to the union of all section faculty
        all_faculty_ids = set()
        for section in sections:
            all_faculty_ids.update(section.faculty.values_list('id', flat=True))
        course.faculty.set(User.objects.filter(id__in=all_faculty_ids))
        messages.success(request, f'Faculty assignments saved for {course.code}. All course content is preserved.')
        return redirect('dao_courses')

    # Attach current faculty directly to each section object for the template
    for sec in sections:
        sec.current_faculty = sec.faculty.first()

    return render(request, 'dao_portal/assign_faculty.html', {
        'course': course,
        'faculty_list': faculty_list,
        'sections': sections,
    })


@dao_required
def dao_assign_faculty(request, course_id):
    if request.method == 'POST':
        c = get_object_or_404(Course, id=course_id)
        faculty_ids = request.POST.getlist('faculty_ids')
        c.faculty.set(User.objects.filter(id__in=faculty_ids, role='faculty'))
        messages.success(request, f'Faculty assignment for {c.code} updated.')
    return redirect('dao_courses')


@dao_required
def dao_toggle_course_active(request, course_id):
    if request.method == 'POST':
        c = get_object_or_404(Course, id=course_id)
        c.is_active = not c.is_active
        c.save()
        status = 'activated' if c.is_active else 'deactivated'
        messages.success(request, f'Course {c.code} {status}.')
    return redirect('dao_courses')


@dao_required
def dao_delete_course(request, course_id):
    if request.method == 'POST':
        c = get_object_or_404(Course, id=course_id)
        name = c.code
        c.delete()
        messages.success(request, f'Course {name} deleted.')
    return redirect('dao_courses')


@dao_required
def dao_analytics(request):
    from django.db.models import Q
    all_semesters = Course.objects.values_list('semester', flat=True).distinct().order_by('semester')

    course_name_q = request.GET.get('course_name', '').strip()
    semester_q    = request.GET.get('semester', '').strip()

    courses = Course.objects.all().order_by('code')
    if course_name_q:
        courses = courses.filter(Q(name__icontains=course_name_q) | Q(code__icontains=course_name_q))
    if semester_q:
        courses = courses.filter(semester=semester_q)

    selected_course = None
    plo_attainment = []
    clo_attainment = []
    total_enrolled = 0
    present_count = 0

    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
    elif courses.exists():
        selected_course = courses.first()

    if selected_course:
        SUB_TYPES = {'mid', 'final'}
        assessments = list(
            Assessment.objects.filter(course=selected_course, status='published')
            .prefetch_related(
                'questions__clos', 'questions__plos',
                'questions__sub_questions__clos', 'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                            'clo_ids': [c.id for c in sq.clos.all()],
                            'plo_ids': [p.id for p in sq.plos.all()],
                            'is_final': is_final,
                        })
                else:
                    all_columns.append({
                        'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                        'clo_ids': [c.id for c in q.clos.all()],
                        'plo_ids': [p.id for p in q.plos.all()],
                        'is_final': is_final,
                    })

        clos = list(selected_course.clos.all())
        clo_max = {
            clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids'])
            for clo in clos
        }
        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }

        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=selected_course,
            ).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=selected_course,
            ).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

        students = list(
            User.objects.filter(enrollments__course=selected_course)
            .distinct().order_by('full_name', 'username')
        )
        total_enrolled = len(students)

        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(
                    col['entity_id'] in (sq_m if col['is_sub'] else q_m)
                    for col in all_columns if col['is_final']
                )
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)

        clo_achieved = {clo.id: 0 for clo in clos}
        plo_achieved = {p.id: 0 for p in plos}

        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue

            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)

            for clo in clos:
                mx = clo_max[clo.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
                if round(raw / mx * 100, 1) >= 40:
                    clo_achieved[clo.id] += 1

            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if round(raw / mx * 100, 1) >= 40:
                    plo_achieved[p.id] += 1

        for clo in clos:
            achieved = clo_achieved[clo.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            clo_attainment.append({
                'code': clo.code, 'description': clo.description,
                'attainment': pct, 'achieved': achieved, 'present': present_count,
            })

        for p in plos:
            achieved = plo_achieved[p.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            plo_attainment.append({
                'code': p.code, 'description': p.description,
                'attainment': pct, 'achieved': achieved, 'present': present_count,
            })

    return render(request, 'dao_portal/analytics.html', {
        'courses': courses,
        'selected_course': selected_course,
        'plo_attainment': plo_attainment,
        'clo_attainment': clo_attainment,
        'total_enrolled': total_enrolled,
        'present_count': present_count,
        'all_semesters': all_semesters,
        'course_name_q': course_name_q,
        'semester_q': semester_q,
    })


@dao_required
def dao_escar(request):
    courses = Course.objects.all().order_by('code')
    selected_course = None
    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
    elif courses.exists():
        selected_course = courses.first()
    rows, total_enrolled, total_participated, total_absent = _compute_escar(selected_course)
    return render(request, 'dao_portal/escar.html', {
        'courses': courses,
        'selected_course': selected_course,
        'rows': rows,
        'total_enrolled': total_enrolled,
        'total_participated': total_participated,
        'total_absent': total_absent,
    })


@dao_required
def dao_students(request):
    from django.db.models import Q, Count
    q = request.GET.get('q', '').strip()
    course_id = request.GET.get('course', '').strip()

    courses = Course.objects.all().order_by('code')
    selected_course = None

    students = User.objects.filter(role='student')

    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
        students = students.filter(enrollments__course=selected_course).distinct()

    if q:
        students = students.filter(Q(username__icontains=q) | Q(full_name__icontains=q))

    students = students.annotate(
        enrolled_count=Count('enrollments', distinct=True)
    ).order_by('username')

    return render(request, 'dao_portal/students.html', {
        'students': students,
        'courses': courses,
        'selected_course': selected_course,
        'q': q,
        'course_id': course_id,
    })


@dao_required
def dao_student_attainment(request, student_id):
    student = get_object_or_404(User, id=student_id, role='student')
    enrollments = Enrollment.objects.filter(student=student).select_related('course')
    results = []
    program_plo_data = {}
    SUB_TYPES = {'mid', 'final'}

    for e in enrollments:
        course = e.course
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related(
                'questions__clos', 'questions__plos',
                'questions__sub_questions__clos', 'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                            'clo_ids': [c.id for c in sq.clos.all()],
                            'plo_ids': [p.id for p in sq.plos.all()],
                        })
                else:
                    all_columns.append({
                        'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                        'clo_ids': [c.id for c in q.clos.all()],
                        'plo_ids': [p.id for p in q.plos.all()],
                    })

        clos = list(course.clos.all())
        clo_max = {
            clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids'])
            for clo in clos
        }
        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }
        total_max_overall = sum(col['max_marks'] for col in all_columns)

        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_m, sq_m = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=course,
                submission__student=student,
            ):
                q_m[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=course,
                submission__student=student,
            ):
                sq_m[g.sub_question_id] = g.marks_obtained

        def _mark(col, _sq=sq_m, _q=q_m):
            return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)

        has_grades = bool(q_m or sq_m)
        total_raw = sum(_mark(col) for col in all_columns)
        avg_pct = round(total_raw / total_max_overall * 100, 1) if (total_max_overall > 0 and has_grades) else 0.0

        grade = 'F'
        if avg_pct >= 80:   grade = 'A+'
        elif avg_pct >= 75: grade = 'A'
        elif avg_pct >= 70: grade = 'A-'
        elif avg_pct >= 65: grade = 'B+'
        elif avg_pct >= 60: grade = 'B'
        elif avg_pct >= 55: grade = 'B-'
        elif avg_pct >= 50: grade = 'C+'
        elif avg_pct >= 45: grade = 'C'
        elif avg_pct >= 40: grade = 'D'

        clo_results = []
        for clo in clos:
            mx = clo_max[clo.id]
            raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
            att = round(raw / mx * 100, 1) if mx > 0 else 0.0
            clo_results.append({
                'code': clo.code, 'bloom': clo.bloom_level,
                'description': clo.description,
                'obtained': round(raw, 1), 'total': mx,
                'attainment': att,
            })

        plo_results = []
        for p in plos:
            mx = plo_max[p.id]
            if mx == 0:
                continue
            raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
            att = round(raw / mx * 100, 1)
            plo_results.append({'code': p.code, 'description': p.description, 'attainment': att})
            if p.id not in program_plo_data:
                program_plo_data[p.id] = {
                    'code': p.code, 'description': p.description,
                    'total_max': 0.0, 'total_raw': 0.0,
                }
            program_plo_data[p.id]['total_max'] += mx
            program_plo_data[p.id]['total_raw'] += raw

        graded_count = Submission.objects.filter(
            student=student, assessment__in=assessments,
            status__in=['graded', 'flagged'],
        ).count()

        results.append({
            'course': course, 'grade': grade,
            'avg_pct': avg_pct, 'graded_count': graded_count,
            'clo_results': clo_results,
            'plo_results': plo_results,
        })

    program_plo_attainment = sorted(
        [
            {
                'code': d['code'],
                'description': d['description'],
                'attainment': round(d['total_raw'] / d['total_max'] * 100, 1),
                'obtained': round(d['total_raw'], 1),
                'total': round(d['total_max'], 1),
            }
            for d in program_plo_data.values() if d['total_max'] > 0
        ],
        key=lambda x: x['code'],
    )

    return render(request, 'dao_portal/student_attainment.html', {
        'student': student,
        'results': results,
        'program_plo_attainment': program_plo_attainment,
    })


# ── DAO Section Management ────────────────────────────────────────────────────

@dao_required
def dao_sections(request):
    SEMESTER_OPTIONS = [
        'Fall 2023', 'Spring 2024', 'Fall 2024',
        'Spring 2025', 'Fall 2025', 'Spring 2026', 'Fall 2026',
    ]
    courses = Course.objects.all().order_by('code')
    # existing semesters already used in sections (for the filter dropdown)
    existing_semesters = list(
        Section.objects.values_list('batch', flat=True).distinct().order_by('batch')
    )
    # merge with predefined list, deduplicate, keep order
    all_semesters = sorted(set(SEMESTER_OPTIONS + existing_semesters))

    course_id  = request.GET.get('course', '').strip()
    semester_q = request.GET.get('batch', '').strip()

    sections = Section.objects.select_related('course').prefetch_related('faculty', 'students').order_by('course__code', 'batch', 'name')
    if course_id:
        sections = sections.filter(course_id=course_id)
    if semester_q:
        sections = sections.filter(batch=semester_q)

    selected_course = Course.objects.filter(id=course_id).first() if course_id else None

    return render(request, 'dao_portal/sections.html', {
        'sections': sections,
        'courses': courses,
        'semester_options': all_semesters,
        'selected_course': selected_course,
        'course_id': course_id,
        'batch_q': semester_q,
    })


@dao_required
def dao_create_section(request):
    if request.method == 'POST':
        course_id = request.POST.get('course_id', '').strip()
        name      = request.POST.get('name', '').strip().upper()
        batch     = request.POST.get('batch', '').strip()
        if not course_id or not name or not batch:
            messages.error(request, 'Course, section name, and batch are required.')
            return redirect('dao_sections')
        course = get_object_or_404(Course, id=course_id)
        if Section.objects.filter(course=course, name=name, batch=batch).exists():
            messages.error(request, f'Section {name} for batch {batch} already exists in {course.code}.')
            return redirect('dao_sections')
        section = Section.objects.create(course=course, name=name, batch=batch)
        messages.success(request, f'Section {name} created for {course.code} (Batch {batch}). Assign faculty and students below.')
        return redirect('dao_section_detail', section_id=section.id)
    return redirect('dao_sections')


@dao_required
def dao_section_detail(request, section_id):
    section = get_object_or_404(Section, id=section_id)
    faculty_list = User.objects.filter(role='faculty').order_by('full_name')
    assigned_faculty_ids = set(section.faculty.values_list('id', flat=True))
    students = section.students.all().order_by('username')
    return render(request, 'dao_portal/section_detail.html', {
        'section': section,
        'faculty_list': faculty_list,
        'assigned_faculty_ids': assigned_faculty_ids,
        'students': students,
    })


@dao_required
def dao_section_assign_faculty(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        faculty_ids = request.POST.getlist('faculty_ids')
        section.faculty.set(User.objects.filter(id__in=faculty_ids, role='faculty'))
        messages.success(request, f'Faculty updated for {section.course.code} — Section {section.name}.')
    return redirect('dao_section_detail', section_id=section_id)


@dao_required
def dao_section_assign_students(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        start_id = request.POST.get('start_reg_id', '').strip()
        end_id   = request.POST.get('end_reg_id', '').strip()
        if not start_id or not end_id:
            messages.error(request, 'Both start and end registration IDs are required.')
            return redirect('dao_section_detail', section_id=section_id)
        matched = User.objects.filter(role='student', username__gte=start_id, username__lte=end_id)
        if not matched.exists():
            messages.error(request, f'No students found with registration ID between "{start_id}" and "{end_id}".')
            return redirect('dao_section_detail', section_id=section_id)
        before = section.students.count()
        section.students.add(*matched)
        added = section.students.count() - before
        messages.success(request, f'{added} new student(s) added to Section {section.name} ({matched.count()} matched the range).')
    return redirect('dao_section_detail', section_id=section_id)


@dao_required
def dao_section_add_single_student(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        reg_id = request.POST.get('reg_id', '').strip()
        if not reg_id:
            messages.error(request, 'Registration ID is required.')
            return redirect('dao_section_detail', section_id=section_id)
        try:
            student = User.objects.get(role='student', username=reg_id)
        except User.DoesNotExist:
            messages.error(request, f'No student found with registration ID "{reg_id}".')
            return redirect('dao_section_detail', section_id=section_id)
        if section.students.filter(id=student.id).exists():
            messages.warning(request, f'{student.full_name or student.username} is already in this section.')
        else:
            section.students.add(student)
            messages.success(request, f'{student.full_name or student.username} added to Section {section.name}.')
    return redirect('dao_section_detail', section_id=section_id)


@dao_required
def dao_section_remove_student(request, section_id, student_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        student = get_object_or_404(User, id=student_id)
        section.students.remove(student)
        messages.success(request, f'{student.full_name or student.username} removed from Section {section.name}.')
    return redirect('dao_section_detail', section_id=section_id)


@dao_required
def dao_delete_section(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        label = f"{section.course.code} — Section {section.name} (Batch {section.batch})"
        section.delete()
        messages.success(request, f'Section deleted: {label}.')
    return redirect('dao_sections')


@dao_required
def dao_plo_comparison(request):
    all_courses = Course.objects.all().order_by('code', 'semester')
    seen_codes = set()
    course_codes = []
    for c in all_courses:
        if c.code not in seen_codes:
            seen_codes.add(c.code)
            course_codes.append({'code': c.code, 'name': c.name})
    selected_code = request.GET.get('course_code', '').strip()
    selected_ids  = request.GET.getlist('sem_ids')
    courses_for_code = []
    if selected_code:
        courses_for_code = list(all_courses.filter(code=selected_code))
    elif course_codes:
        selected_code = course_codes[0]['code']
        courses_for_code = list(all_courses.filter(code=selected_code))
    if selected_ids:
        compare_set = set(selected_ids)
        courses_to_compare = [c for c in courses_for_code if str(c.id) in compare_set]
    else:
        courses_to_compare = list(courses_for_code)
    SUB_TYPES = {'mid', 'final'}
    comparison_data = []
    all_plo_meta = {}
    for course in courses_to_compare:
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related('questions__plos', 'questions__sub_questions__plos')
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'plo_ids': [p.id for p in sq.plos.all()], 'is_final': is_final})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'plo_ids': [p.id for p in q.plos.all()], 'is_final': is_final})
        if not all_columns:
            comparison_data.append({'course': course, 'semester': course.semester, 'plo_attainment': {}, 'present_count': 0})
            continue
        plo_ids_used = set(pid for col in all_columns for pid in col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        for p in plos:
            all_plo_meta[p.code] = p.description
        q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=course).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=course).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained
        students = list(User.objects.filter(enrollments__course=course).distinct())
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(col['entity_id'] in (sq_m if col['is_sub'] else q_m) for col in all_columns if col['is_final'])
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)
        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue
            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if present_count > 0 and (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1
        plo_attainment = {p.code: (round(plo_achieved[p.id] / present_count * 100, 1) if present_count > 0 else 0.0) for p in plos}
        comparison_data.append({'course': course, 'semester': course.semester, 'plo_attainment': plo_attainment, 'present_count': present_count})
    def _sem_key(entry):
        parts = entry['semester'].split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))
    comparison_data.sort(key=_sem_key)
    plo_codes = sorted(all_plo_meta.keys())
    table_rows = []
    for entry in comparison_data:
        cells = [{'plo_code': code, 'attainment': entry['plo_attainment'].get(code),
                  'has_data': entry['plo_attainment'].get(code) is not None} for code in plo_codes]
        table_rows.append({'semester': entry['semester'], 'course_id': entry['course'].id,
                           'is_archived': entry['course'].is_archived, 'present_count': entry['present_count'], 'cells': cells})
    chart_json = json.dumps({
        'plo_codes': plo_codes, 'plo_descriptions': all_plo_meta,
        'semesters': [e['semester'] for e in comparison_data],
        'datasets': [{'semester': e['semester'], 'values': [e['plo_attainment'].get(code) for code in plo_codes]} for e in comparison_data],
    })
    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]
    return render(request, 'dao_portal/plo_comparison.html', {
        'course_codes': course_codes, 'selected_code': selected_code,
        'courses_for_code': courses_for_code, 'selected_ids': [int(i) for i in selected_ids if i.isdigit()],
        'plo_codes': plo_codes, 'plo_list': plo_list, 'table_rows': table_rows,
        'chart_json': chart_json, 'has_data': bool(plo_codes and table_rows),
    })


@dao_required
def dao_semester_plo_comparison(request):
    SUB_TYPES = {'mid', 'final'}
    all_courses = list(Course.objects.all().order_by('semester', 'code'))
    def _sem_key(s):
        parts = s.split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))
    all_semesters = sorted(set(c.semester for c in all_courses), key=_sem_key)
    selected_sems = request.GET.getlist('sems')
    if not selected_sems:
        selected_sems = list(all_semesters)
    selected_sems_set = set(selected_sems)
    semester_plo_data = {}
    all_plo_meta = {}
    semester_present_students = {}
    for course in all_courses:
        if course.semester not in selected_sems_set:
            continue
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related('questions__plos', 'questions__sub_questions__plos')
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'plo_ids': [p.id for p in sq.plos.all()], 'is_final': is_final})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'plo_ids': [p.id for p in q.plos.all()], 'is_final': is_final})
        if not all_columns:
            continue
        plo_ids_used = set(pid for col in all_columns for pid in col['plo_ids'])
        if not plo_ids_used:
            continue
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        for p in plos:
            all_plo_meta[p.code] = p.description
        q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=course).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=course).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained
        students = list(User.objects.filter(enrollments__course=course).distinct())
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(col['entity_id'] in (sq_m if col['is_sub'] else q_m) for col in all_columns if col['is_final'])
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)
        if present_count == 0:
            continue
        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue
            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1
        semester = course.semester
        if semester not in semester_plo_data:
            semester_plo_data[semester] = {}
        semester_present_students.setdefault(semester, set()).update(present_student_ids)
        for p in plos:
            if plo_max[p.id] == 0:
                continue
            code = p.code
            if code not in semester_plo_data[semester]:
                semester_plo_data[semester][code] = {'description': p.description, 'total_achieved': 0, 'total_present': 0}
            semester_plo_data[semester][code]['total_achieved'] += plo_achieved[p.id]
            semester_plo_data[semester][code]['total_present'] += present_count
    plo_codes = sorted(all_plo_meta.keys())
    semesters_sorted = sorted(semester_plo_data.keys(), key=_sem_key)
    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]
    table_rows = []
    for sem in semesters_sorted:
        cells = []
        for code in plo_codes:
            pe = semester_plo_data.get(sem, {}).get(code)
            if pe and pe['total_present'] > 0:
                att = round(pe['total_achieved'] / pe['total_present'] * 100, 1)
                cells.append({'plo_code': code, 'attainment': att, 'has_data': True})
            else:
                cells.append({'plo_code': code, 'attainment': 0, 'has_data': False})
        table_rows.append({'semester': sem, 'present_count': len(semester_present_students.get(sem, set())), 'cells': cells})
    chart_json = json.dumps({
        'plo_codes': plo_codes, 'semesters': semesters_sorted,
        'datasets': [
            {'plo': code, 'values': [
                (round(semester_plo_data[sem][code]['total_achieved'] / semester_plo_data[sem][code]['total_present'] * 100, 1)
                 if (sem in semester_plo_data and code in semester_plo_data[sem] and semester_plo_data[sem][code]['total_present'] > 0)
                 else None)
                for sem in semesters_sorted
            ]}
            for code in plo_codes
        ],
    })
    return render(request, 'dao_portal/semester_plo_comparison.html', {
        'all_semesters': all_semesters, 'selected_sems': selected_sems,
        'plo_codes': plo_codes, 'plo_list': plo_list, 'table_rows': table_rows,
        'semesters': semesters_sorted, 'chart_json': chart_json,
        'has_data': bool(plo_codes and semesters_sorted),
    })


# ── DEPARTMENT HEAD PORTAL VIEWS ──────────────────────────────────────────────

def dept_head_required(view_func):
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('sign_in_html')
        if request.user.role != 'dept_head':
            return redirect('home')
        return view_func(request, *args, **kwargs)
    wrapper.__name__ = view_func.__name__
    return wrapper


def _marks_sheet_filter_ctx(request):
    """Shared filter logic for read-only marks sheet views."""
    all_semesters = (
        Course.objects.values_list('semester', flat=True)
        .distinct().order_by('semester')
    )
    selected_semester = request.GET.get('semester', '')
    selected_section_id = request.GET.get('section', '')
    course_id = request.GET.get('course', '')

    courses = Course.objects.all().order_by('semester', 'code')
    if selected_semester:
        courses = courses.filter(semester=selected_semester)

    selected_section = None
    selected_course = None

    if selected_section_id:
        selected_section = get_object_or_404(Section, id=selected_section_id)
        selected_course = selected_section.course
    elif course_id:
        selected_course = get_object_or_404(Course, id=course_id)
    elif courses.exists():
        selected_course = courses.first()

    # Sections only for the currently selected course
    all_sections = (
        Section.objects.filter(course=selected_course).order_by('batch', 'name')
        if selected_course else Section.objects.none()
    )

    return {
        'courses': courses,
        'all_semesters': all_semesters,
        'selected_semester': selected_semester,
        'all_sections': all_sections,
        'selected_section': selected_section,
        'selected_course': selected_course,
    }


@dao_required
def dao_marks_sheet(request):
    ctx = _marks_sheet_filter_ctx(request)
    if ctx['selected_course']:
        ctx.update(_compute_marks_sheet_readonly(ctx['selected_course'], ctx['selected_section']))
    return render(request, 'dao_portal/marks_sheet.html', ctx)


@admin_required
def admin_marks_sheet(request):
    ctx = _marks_sheet_filter_ctx(request)
    if ctx['selected_course']:
        ctx.update(_compute_marks_sheet_readonly(ctx['selected_course'], ctx['selected_section']))
    return render(request, 'admin_portal/marks_sheet.html', ctx)


@dept_head_required
def dept_head_marks_sheet(request):
    ctx = _marks_sheet_filter_ctx(request)
    if ctx['selected_course']:
        ctx.update(_compute_marks_sheet_readonly(ctx['selected_course'], ctx['selected_section']))
    return render(request, 'dept_head_portal/marks_sheet.html', ctx)


@dept_head_required
def dept_head_dashboard(request):
    from datetime import date as dt_date
    return render(request, 'dept_head_portal/dashboard.html', {
        'today':             dt_date.today().strftime('%B %d, %Y'),
        'total_users':       User.objects.count(),
        'total_faculty':     User.objects.filter(role='faculty').count(),
        'total_students':    User.objects.filter(role='student').count(),
        'total_courses':     Course.objects.count(),
        'total_enrollments': Enrollment.objects.count(),
        'recent_users':      User.objects.order_by('-date_joined')[:6],
        'recent_courses':    Course.objects.prefetch_related('enrollments').order_by('-created_at')[:6],
    })


@dept_head_required
def dept_head_users(request):
    users = User.objects.all().order_by('-date_joined')
    return render(request, 'dept_head_portal/users.html', {'users': users})


@dept_head_required
def dept_head_create_user(request):
    if request.method == 'POST':
        full_name = request.POST.get('full_name', '').strip()
        email     = request.POST.get('email', '').strip()
        password  = request.POST.get('password', '')
        role      = request.POST.get('role', 'student')
        if role not in ('faculty', 'student'):
            role = 'student'
        if User.objects.filter(email=email).exists():
            messages.error(request, 'Email already registered.')
            return redirect('dept_head_users')
        if len(password) < 8:
            messages.error(request, 'Password must be at least 8 characters.')
            return redirect('dept_head_users')
        username = email.split('@')[0]
        base = username; i = 1
        while User.objects.filter(username=username).exists():
            username = f"{base}{i}"; i += 1
        User.objects.create_user(
            username=username, email=email, password=password,
            full_name=full_name, role=role,
        )
        messages.success(request, f'User {full_name} created successfully.')
    return redirect('dept_head_users')


@dept_head_required
def dept_head_edit_user(request, user_id):
    u = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        u.full_name = request.POST.get('full_name', u.full_name).strip()
        u.email     = request.POST.get('email', u.email).strip()
        new_role    = request.POST.get('role', u.role)
        if new_role in ('faculty', 'student'):
            u.role = new_role
        pw = request.POST.get('password', '').strip()
        if pw:
            if len(pw) < 8:
                messages.error(request, 'Password must be at least 8 characters.')
                return redirect('dept_head_users')
            u.set_password(pw)
        u.save()
        messages.success(request, f'User {u.full_name} updated.')
    return redirect('dept_head_users')


@dept_head_required
def dept_head_toggle_user(request, user_id):
    if request.method == 'POST':
        u = get_object_or_404(User, id=user_id)
        u.is_active = not u.is_active
        u.save()
        status = 'activated' if u.is_active else 'deactivated'
        messages.success(request, f'User {u.full_name or u.username} {status}.')
    return redirect('dept_head_users')


@dept_head_required
def dept_head_delete_user(request, user_id):
    if request.method == 'POST':
        u = get_object_or_404(User, id=user_id)
        name = u.full_name or u.username
        u.delete()
        messages.success(request, f'User {name} deleted.')
    return redirect('dept_head_users')


@dept_head_required
def dept_head_courses(request):
    courses = Course.objects.prefetch_related('faculty', 'enrollments', 'assessments', 'sections').order_by('-created_at')
    semesters = Course.objects.values_list('semester', flat=True).distinct().order_by('semester')
    faculty_list = User.objects.filter(role='faculty').order_by('full_name')
    return render(request, 'dept_head_portal/courses.html', {
        'courses': courses,
        'semesters': semesters,
        'faculty_list': faculty_list,
        'semester_options': SEMESTER_OPTIONS,
    })


@dept_head_required
def dept_head_create_course(request):
    if request.method == 'POST':
        code         = request.POST.get('code', '').strip()
        name         = request.POST.get('name', '').strip()
        description  = request.POST.get('description', '').strip()
        credit_hours = int(request.POST.get('credit_hours', 3) or 3)
        semester     = request.POST.get('semester', 'Fall 2025')
        if not code or not name:
            messages.error(request, 'Course code and name are required.')
            return redirect('dept_head_courses')
        course = Course.objects.create(
            code=code, name=name, description=description,
            credit_hours=credit_hours, semester=semester,
        )
        sec_names = request.POST.getlist('section_name')
        sec_sems  = request.POST.getlist('section_semester')
        created_sections = 0
        for sec_name, sec_sem in zip(sec_names, sec_sems):
            sec_name = sec_name.strip().upper()
            sec_sem  = sec_sem.strip()
            if sec_name and sec_sem:
                Section.objects.get_or_create(course=course, name=sec_name, batch=sec_sem)
                created_sections += 1
        msg = f'Course {code} created'
        msg += f' with {created_sections} section(s).' if created_sections else '.'
        messages.success(request, msg)
        return redirect('dept_head_assign_faculty_page', course_id=course.id)
    return redirect('dept_head_courses')


@dept_head_required
def dept_head_assign_faculty_page(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    faculty_list = User.objects.filter(role='faculty').order_by('full_name')
    sections = list(course.sections.prefetch_related('faculty').order_by('batch', 'name'))

    if request.method == 'POST':
        for section in sections:
            faculty_id = request.POST.get(f'faculty_{section.id}', '').strip()
            if faculty_id:
                faculty = User.objects.filter(id=faculty_id, role='faculty').first()
                if faculty:
                    section.faculty.set([faculty])
            else:
                section.faculty.clear()
        all_faculty_ids = set()
        for section in sections:
            all_faculty_ids.update(section.faculty.values_list('id', flat=True))
        course.faculty.set(User.objects.filter(id__in=all_faculty_ids))
        messages.success(request, f'Faculty assignments saved for {course.code}.')
        return redirect('dept_head_courses')

    for sec in sections:
        sec.current_faculty = sec.faculty.first()

    return render(request, 'dept_head_portal/assign_faculty.html', {
        'course': course,
        'faculty_list': faculty_list,
        'sections': sections,
    })


@dept_head_required
def dept_head_assign_faculty(request, course_id):
    if request.method == 'POST':
        c = get_object_or_404(Course, id=course_id)
        faculty_ids = request.POST.getlist('faculty_ids')
        c.faculty.set(User.objects.filter(id__in=faculty_ids, role='faculty'))
        messages.success(request, f'Faculty assignment for {c.code} updated.')
    return redirect('dept_head_courses')


@dept_head_required
def dept_head_toggle_course_active(request, course_id):
    if request.method == 'POST':
        c = get_object_or_404(Course, id=course_id)
        c.is_active = not c.is_active
        c.save()
        status = 'activated' if c.is_active else 'deactivated'
        messages.success(request, f'Course {c.code} {status}.')
    return redirect('dept_head_courses')


@dept_head_required
def dept_head_delete_course(request, course_id):
    if request.method == 'POST':
        c = get_object_or_404(Course, id=course_id)
        name = c.code
        c.delete()
        messages.success(request, f'Course {name} deleted.')
    return redirect('dept_head_courses')


@dept_head_required
def dept_head_analytics(request):
    from django.db.models import Q
    all_semesters = Course.objects.values_list('semester', flat=True).distinct().order_by('semester')

    course_name_q = request.GET.get('course_name', '').strip()
    semester_q    = request.GET.get('semester', '').strip()

    courses = Course.objects.all().order_by('code')
    if course_name_q:
        courses = courses.filter(Q(name__icontains=course_name_q) | Q(code__icontains=course_name_q))
    if semester_q:
        courses = courses.filter(semester=semester_q)

    selected_course = None
    plo_attainment = []
    clo_attainment = []
    total_enrolled = 0
    present_count = 0

    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
    elif courses.exists():
        selected_course = courses.first()

    if selected_course:
        SUB_TYPES = {'mid', 'final'}
        assessments = list(
            Assessment.objects.filter(course=selected_course, status='published')
            .prefetch_related(
                'questions__clos', 'questions__plos',
                'questions__sub_questions__clos', 'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                            'clo_ids': [c.id for c in sq.clos.all()],
                            'plo_ids': [p.id for p in sq.plos.all()],
                            'is_final': is_final,
                        })
                else:
                    all_columns.append({
                        'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                        'clo_ids': [c.id for c in q.clos.all()],
                        'plo_ids': [p.id for p in q.plos.all()],
                        'is_final': is_final,
                    })

        clos = list(selected_course.clos.all())
        clo_max = {
            clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids'])
            for clo in clos
        }
        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }

        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=selected_course,
            ).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=selected_course,
            ).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

        students = list(
            User.objects.filter(enrollments__course=selected_course)
            .distinct().order_by('full_name', 'username')
        )
        total_enrolled = len(students)

        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(
                    col['entity_id'] in (sq_m if col['is_sub'] else q_m)
                    for col in all_columns if col['is_final']
                )
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)

        clo_achieved = {clo.id: 0 for clo in clos}
        plo_achieved = {p.id: 0 for p in plos}

        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue

            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)

            for clo in clos:
                mx = clo_max[clo.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
                if round(raw / mx * 100, 1) >= 40:
                    clo_achieved[clo.id] += 1

            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if round(raw / mx * 100, 1) >= 40:
                    plo_achieved[p.id] += 1

        for clo in clos:
            achieved = clo_achieved[clo.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            clo_attainment.append({
                'code': clo.code, 'description': clo.description,
                'attainment': pct, 'achieved': achieved, 'present': present_count,
            })

        for p in plos:
            achieved = plo_achieved[p.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            plo_attainment.append({
                'code': p.code, 'description': p.description,
                'attainment': pct, 'achieved': achieved, 'present': present_count,
            })

    return render(request, 'dept_head_portal/analytics.html', {
        'courses': courses,
        'selected_course': selected_course,
        'plo_attainment': plo_attainment,
        'clo_attainment': clo_attainment,
        'total_enrolled': total_enrolled,
        'present_count': present_count,
        'all_semesters': all_semesters,
        'course_name_q': course_name_q,
        'semester_q': semester_q,
    })


@dept_head_required
def dept_head_students(request):
    from django.db.models import Q, Count
    q = request.GET.get('q', '').strip()
    course_id = request.GET.get('course', '').strip()

    courses = Course.objects.all().order_by('code')
    selected_course = None
    students = User.objects.filter(role='student')

    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
        students = students.filter(enrollments__course=selected_course).distinct()

    if q:
        students = students.filter(Q(username__icontains=q) | Q(full_name__icontains=q))

    students = students.annotate(
        enrolled_count=Count('enrollments', distinct=True)
    ).order_by('username')

    return render(request, 'dept_head_portal/students.html', {
        'students': students,
        'courses': courses,
        'selected_course': selected_course,
        'q': q,
        'course_id': course_id,
    })


@dept_head_required
def dept_head_student_attainment(request, student_id):
    student = get_object_or_404(User, id=student_id, role='student')
    enrollments = Enrollment.objects.filter(student=student).select_related('course')
    results = []
    program_plo_data = {}
    SUB_TYPES = {'mid', 'final'}

    for e in enrollments:
        course = e.course
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related(
                'questions__clos', 'questions__plos',
                'questions__sub_questions__clos', 'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                            'clo_ids': [c.id for c in sq.clos.all()],
                            'plo_ids': [p.id for p in sq.plos.all()],
                        })
                else:
                    all_columns.append({
                        'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                        'clo_ids': [c.id for c in q.clos.all()],
                        'plo_ids': [p.id for p in q.plos.all()],
                    })

        clos = list(course.clos.all())
        clo_max = {
            clo.id: sum(col['max_marks'] for col in all_columns if clo.id in col['clo_ids'])
            for clo in clos
        }
        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {
            p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids'])
            for p in plos
        }
        total_max_overall = sum(col['max_marks'] for col in all_columns)

        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_m, sq_m = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=course,
                submission__student=student,
            ):
                q_m[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=course,
                submission__student=student,
            ):
                sq_m[g.sub_question_id] = g.marks_obtained

        def _mark(col, _sq=sq_m, _q=q_m):
            return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)

        has_grades = bool(q_m or sq_m)
        total_raw = sum(_mark(col) for col in all_columns)
        avg_pct = round(total_raw / total_max_overall * 100, 1) if (total_max_overall > 0 and has_grades) else 0.0

        grade = 'F'
        if avg_pct >= 80:   grade = 'A+'
        elif avg_pct >= 75: grade = 'A'
        elif avg_pct >= 70: grade = 'A-'
        elif avg_pct >= 65: grade = 'B+'
        elif avg_pct >= 60: grade = 'B'
        elif avg_pct >= 55: grade = 'B-'
        elif avg_pct >= 50: grade = 'C+'
        elif avg_pct >= 45: grade = 'C'
        elif avg_pct >= 40: grade = 'D'

        clo_results = []
        for clo in clos:
            mx = clo_max[clo.id]
            raw = sum(_mark(col) for col in all_columns if clo.id in col['clo_ids'])
            att = round(raw / mx * 100, 1) if mx > 0 else 0.0
            clo_results.append({
                'code': clo.code, 'bloom': clo.bloom_level,
                'description': clo.description,
                'obtained': round(raw, 1), 'total': mx,
                'attainment': att,
            })

        plo_results = []
        for p in plos:
            mx = plo_max[p.id]
            if mx == 0:
                continue
            raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
            att = round(raw / mx * 100, 1)
            plo_results.append({'code': p.code, 'description': p.description, 'attainment': att})
            if p.id not in program_plo_data:
                program_plo_data[p.id] = {
                    'code': p.code, 'description': p.description,
                    'total_max': 0.0, 'total_raw': 0.0,
                }
            program_plo_data[p.id]['total_max'] += mx
            program_plo_data[p.id]['total_raw'] += raw

        graded_count = Submission.objects.filter(
            student=student, assessment__in=assessments,
            status__in=['graded', 'flagged'],
        ).count()

        results.append({
            'course': course, 'grade': grade,
            'avg_pct': avg_pct, 'graded_count': graded_count,
            'clo_results': clo_results,
            'plo_results': plo_results,
        })

    program_plo_attainment = sorted(
        [
            {
                'code': d['code'],
                'description': d['description'],
                'attainment': round(d['total_raw'] / d['total_max'] * 100, 1),
                'obtained': round(d['total_raw'], 1),
                'total': round(d['total_max'], 1),
            }
            for d in program_plo_data.values() if d['total_max'] > 0
        ],
        key=lambda x: x['code'],
    )

    return render(request, 'dept_head_portal/student_attainment.html', {
        'student': student,
        'results': results,
        'program_plo_attainment': program_plo_attainment,
    })


@dept_head_required
def dept_head_sections(request):
    SEMESTER_OPTIONS_DH = [
        'Fall 2023', 'Spring 2024', 'Fall 2024',
        'Spring 2025', 'Fall 2025', 'Spring 2026', 'Fall 2026',
    ]
    courses = Course.objects.all().order_by('code')
    existing_semesters = list(
        Section.objects.values_list('batch', flat=True).distinct().order_by('batch')
    )
    all_semesters = sorted(set(SEMESTER_OPTIONS_DH + existing_semesters))

    course_id  = request.GET.get('course', '').strip()
    semester_q = request.GET.get('batch', '').strip()

    sections = Section.objects.select_related('course').prefetch_related('faculty', 'students').order_by('course__code', 'batch', 'name')
    if course_id:
        sections = sections.filter(course_id=course_id)
    if semester_q:
        sections = sections.filter(batch=semester_q)

    selected_course = Course.objects.filter(id=course_id).first() if course_id else None

    return render(request, 'dept_head_portal/sections.html', {
        'sections': sections,
        'courses': courses,
        'semester_options': all_semesters,
        'selected_course': selected_course,
        'course_id': course_id,
        'batch_q': semester_q,
    })


@dept_head_required
def dept_head_create_section(request):
    if request.method == 'POST':
        course_id = request.POST.get('course_id', '').strip()
        name      = request.POST.get('name', '').strip().upper()
        batch     = request.POST.get('batch', '').strip()
        if not course_id or not name or not batch:
            messages.error(request, 'Course, section name, and batch are required.')
            return redirect('dept_head_sections')
        course = get_object_or_404(Course, id=course_id)
        if Section.objects.filter(course=course, name=name, batch=batch).exists():
            messages.error(request, f'Section {name} for batch {batch} already exists in {course.code}.')
            return redirect('dept_head_sections')
        section = Section.objects.create(course=course, name=name, batch=batch)
        messages.success(request, f'Section {name} created for {course.code} (Batch {batch}).')
        return redirect('dept_head_section_detail', section_id=section.id)
    return redirect('dept_head_sections')


@dept_head_required
def dept_head_section_detail(request, section_id):
    section = get_object_or_404(Section, id=section_id)
    faculty_list = User.objects.filter(role='faculty').order_by('full_name')
    assigned_faculty_ids = set(section.faculty.values_list('id', flat=True))
    students = section.students.all().order_by('username')
    return render(request, 'dept_head_portal/section_detail.html', {
        'section': section,
        'faculty_list': faculty_list,
        'assigned_faculty_ids': assigned_faculty_ids,
        'students': students,
    })


@dept_head_required
def dept_head_section_assign_faculty(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        faculty_ids = request.POST.getlist('faculty_ids')
        section.faculty.set(User.objects.filter(id__in=faculty_ids, role='faculty'))
        messages.success(request, f'Faculty updated for {section.course.code} — Section {section.name}.')
    return redirect('dept_head_section_detail', section_id=section_id)


@dept_head_required
def dept_head_section_assign_students(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        start_id = request.POST.get('start_reg_id', '').strip()
        end_id   = request.POST.get('end_reg_id', '').strip()
        if not start_id or not end_id:
            messages.error(request, 'Both start and end registration IDs are required.')
            return redirect('dept_head_section_detail', section_id=section_id)
        matched = User.objects.filter(role='student', username__gte=start_id, username__lte=end_id)
        if not matched.exists():
            messages.error(request, f'No students found with registration ID between "{start_id}" and "{end_id}".')
            return redirect('dept_head_section_detail', section_id=section_id)
        before = section.students.count()
        section.students.add(*matched)
        added = section.students.count() - before
        messages.success(request, f'{added} new student(s) added to Section {section.name}.')
    return redirect('dept_head_section_detail', section_id=section_id)


@dept_head_required
def dept_head_section_add_single_student(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        reg_id = request.POST.get('reg_id', '').strip()
        if not reg_id:
            messages.error(request, 'Registration ID is required.')
            return redirect('dept_head_section_detail', section_id=section_id)
        try:
            student = User.objects.get(role='student', username=reg_id)
        except User.DoesNotExist:
            messages.error(request, f'No student found with registration ID "{reg_id}".')
            return redirect('dept_head_section_detail', section_id=section_id)
        if section.students.filter(id=student.id).exists():
            messages.warning(request, f'{student.full_name or student.username} is already in this section.')
        else:
            section.students.add(student)
            messages.success(request, f'{student.full_name or student.username} added to Section {section.name}.')
    return redirect('dept_head_section_detail', section_id=section_id)


@dept_head_required
def dept_head_section_remove_student(request, section_id, student_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        student = get_object_or_404(User, id=student_id)
        section.students.remove(student)
        messages.success(request, f'{student.full_name or student.username} removed from Section {section.name}.')
    return redirect('dept_head_section_detail', section_id=section_id)


@dept_head_required
def dept_head_delete_section(request, section_id):
    if request.method == 'POST':
        section = get_object_or_404(Section, id=section_id)
        label = f"{section.course.code} — Section {section.name} (Batch {section.batch})"
        section.delete()
        messages.success(request, f'Section deleted: {label}.')
    return redirect('dept_head_sections')


@dept_head_required
def dept_head_escar(request):
    courses = Course.objects.all().order_by('code')
    selected_course = None
    course_id = request.GET.get('course')
    if course_id:
        selected_course = get_object_or_404(Course, id=course_id)
    elif courses.exists():
        selected_course = courses.first()
    rows, total_enrolled, total_participated, total_absent = _compute_escar(selected_course)
    return render(request, 'dept_head_portal/escar.html', {
        'courses': courses,
        'selected_course': selected_course,
        'rows': rows,
        'total_enrolled': total_enrolled,
        'total_participated': total_participated,
        'total_absent': total_absent,
    })


# ── Assessment Download (PDF / DOCX) ─────────────────────────────────────────

@faculty_required
def download_assessment(request, assessment_id, fmt):
    from django.http import HttpResponse
    assessment = get_object_or_404(
        Assessment, id=assessment_id, course__faculty=request.user
    )
    questions = assessment.questions.prefetch_related(
        'clos', 'plos', 'sub_questions__clos', 'sub_questions__plos'
    ).order_by('order')

    if fmt == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, HRFlowable)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        import io

        from xml.sax.saxutils import escape as _xe
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        TEAL   = colors.HexColor('#20B2AA')
        LTEAL  = colors.HexColor('#e6f7f7')
        ORANGE = colors.HexColor('#f97316')
        GRAY   = colors.HexColor('#6b7280')

        title_style = ParagraphStyle('ATitle', parent=styles['Normal'],
            fontSize=18, fontName='Helvetica-Bold', textColor=TEAL, spaceAfter=4, leading=22)
        sub_style   = ParagraphStyle('ASub', parent=styles['Normal'],
            fontSize=10, fontName='Helvetica', textColor=GRAY, spaceAfter=2)
        qnum_style  = ParagraphStyle('QNum', parent=styles['Normal'],
            fontSize=12, fontName='Helvetica-Bold', textColor=TEAL, alignment=TA_CENTER)
        qtext_style = ParagraphStyle('QText', parent=styles['Normal'],
            fontSize=11, fontName='Helvetica', leading=15)
        sq_style    = ParagraphStyle('SQText', parent=styles['Normal'],
            fontSize=10, fontName='Helvetica', leading=14,
            textColor=colors.HexColor('#374151'))

        story = []

        header_data = [[
            Paragraph(f"{_xe(assessment.course.code)}: {_xe(assessment.course.name)}", title_style),
            Paragraph(
                f"<font color='#20B2AA'><b>{_xe(assessment.get_assessment_type_display())}</b></font>",
                ParagraphStyle('th', parent=styles['Normal'], fontSize=13,
                               fontName='Helvetica-Bold', alignment=TA_CENTER)
            )
        ]]
        ht = Table(header_data, colWidths=['75%', '25%'])
        ht.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), LTEAL),
            ('BOX',        (0,0), (-1,-1), 0.5, TEAL),
            ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING',(0,0), (-1,-1), 14),
            ('RIGHTPADDING',(0,0),(-1,-1), 14),
            ('TOPPADDING', (0,0), (-1,-1), 12),
            ('BOTTOMPADDING',(0,0),(-1,-1),12),
        ]))
        story.append(ht)
        story.append(Spacer(1, 10))

        due_txt  = str(assessment.due_date) if assessment.due_date else 'No due date'
        meta_txt = (f"<b>Total Marks:</b> {assessment.total_marks} &nbsp;&nbsp; "
                    f"<b>Due Date:</b> {due_txt} &nbsp;&nbsp; "
                    f"<b>Semester:</b> {_xe(assessment.course.semester)}")
        if assessment.description:
            meta_txt += f"<br/><b>Description:</b> {_xe(assessment.description)}"
        story.append(Paragraph(meta_txt, sub_style))
        story.append(Spacer(1, 6))
        story.append(HRFlowable(width='100%', thickness=1, color=TEAL))
        story.append(Spacer(1, 12))

        for q in questions:
            clo_codes = ', '.join(c.code for c in q.clos.all()) or '—'
            plo_codes = ', '.join(p.code for p in q.plos.all()) or '—'

            q_data = [[
                Paragraph(f"Q{q.order}", qnum_style),
                Paragraph(_xe(q.text), qtext_style),
                Paragraph(
                    f"<font color='#f97316'><b>{q.max_marks} marks</b></font><br/>"
                    f"<font color='#1d4ed8' size='8'><b>CLO:</b> {_xe(clo_codes)}</font><br/>"
                    f"<font color='#9d174d' size='8'><b>PLO:</b> {_xe(plo_codes)}</font>",
                    ParagraphStyle('mk', parent=styles['Normal'],
                                   fontSize=10, fontName='Helvetica',
                                   alignment=TA_CENTER, leading=13)
                ),
            ]]
            qt = Table(q_data, colWidths=[1.0*cm, 11.8*cm, 4.2*cm])
            qt.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
                ('BOX',       (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
                ('VALIGN',    (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING',(0,0),(-1,-1), 10),
                ('RIGHTPADDING',(0,0),(-1,-1), 10),
                ('TOPPADDING',(0,0),(-1,-1), 10),
                ('BOTTOMPADDING',(0,0),(-1,-1), 10),
            ]))
            story.append(qt)

            for sq in q.sub_questions.order_by('order'):
                sq_clo = ', '.join(c.code for c in sq.clos.all()) or '—'
                sq_plo = ', '.join(p.code for p in sq.plos.all()) or '—'
                sq_data = [[
                    Paragraph(f"({sq.order})", ParagraphStyle('sqn', parent=styles['Normal'],
                        fontSize=10, fontName='Helvetica-Bold', textColor=GRAY,
                        alignment=TA_CENTER)),
                    Paragraph(_xe(sq.text), sq_style),
                    Paragraph(
                        f"<font color='#f97316'><b>{sq.max_marks} marks</b></font><br/>"
                        f"<font color='#1d4ed8' size='8'><b>CLO:</b> {_xe(sq_clo)}</font><br/>"
                        f"<font color='#9d174d' size='8'><b>PLO:</b> {_xe(sq_plo)}</font>",
                        ParagraphStyle('sqm', parent=styles['Normal'],
                                       fontSize=9, fontName='Helvetica',
                                       alignment=TA_CENTER, leading=12)
                    ),
                ]]
                stt = Table(sq_data, colWidths=[1.0*cm, 11.8*cm, 4.2*cm])
                stt.setStyle(TableStyle([
                    ('BACKGROUND',(0,0),(-1,-1), colors.HexColor('#fafbfc')),
                    ('LINEBEFORE',(0,0),(0,-1), 3, TEAL),
                    ('LINEBELOW', (0,0),(-1,-1), 0.3, colors.HexColor('#e2e8f0')),
                    ('LEFTPADDING',(0,0),(0,-1), 10),
                    ('LEFTPADDING',(1,0),(1,-1), 18),
                    ('LEFTPADDING',(2,0),(2,-1), 8),
                    ('RIGHTPADDING',(0,0),(-1,-1), 8),
                    ('TOPPADDING',(0,0),(-1,-1), 6),
                    ('BOTTOMPADDING',(0,0),(-1,-1), 6),
                    ('VALIGN',(0,0),(-1,-1), 'TOP'),
                ]))
                story.append(stt)

            story.append(Spacer(1, 10))

        doc.build(story)
        buf.seek(0)
        safe_title = assessment.title.replace(' ', '_')
        response = HttpResponse(buf, content_type='application/pdf')
        response['Content-Disposition'] = (
            f'attachment; filename="{safe_title}_{assessment.course.code}.pdf"'
        )
        return response

    elif fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io

        def hex_rgb(h):
            h = h.lstrip('#')
            return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

        def set_cell_bg(cell, hex_color):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  hex_color.lstrip('#'))
            tcPr.append(shd)

        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(2)
            sec.bottom_margin = Cm(2)
            sec.left_margin   = Cm(2.5)
            sec.right_margin  = Cm(2.5)

        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run(f"{assessment.course.code}: {assessment.course.name}")
        run.font.size = Pt(18); run.font.bold = True
        run.font.color.rgb = hex_rgb('#20B2AA')

        t2 = doc.add_paragraph()
        t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = t2.add_run(assessment.get_assessment_type_display())
        r2.font.size = Pt(13); r2.font.bold = True
        r2.font.color.rgb = hex_rgb('#20B2AA')

        due_txt = str(assessment.due_date) if assessment.due_date else 'No due date'
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = meta.add_run(
            f"Total Marks: {assessment.total_marks}   |   "
            f"Due Date: {due_txt}   |   "
            f"Semester: {assessment.course.semester}"
        )
        mr.font.size = Pt(10); mr.font.color.rgb = hex_rgb('#6b7280')

        if assessment.description:
            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dr = dp.add_run(f"Description: {assessment.description}")
            dr.font.size = Pt(10); dr.font.italic = True

        doc.add_paragraph()

        for q in questions:
            clo_codes = ', '.join(c.code for c in q.clos.all()) or '—'
            plo_codes = ', '.join(p.code for p in q.plos.all()) or '—'

            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = 'Table Grid'
            tbl.autofit = False
            tbl.allow_autofit = False
            row = tbl.rows[0]
            row.cells[0].width = Cm(1.2)
            row.cells[1].width = Cm(11.0)
            row.cells[2].width = Cm(3.8)
            set_cell_bg(row.cells[0], 'f0fafa')
            set_cell_bg(row.cells[1], 'f8fafc')
            set_cell_bg(row.cells[2], 'f8fafc')
            c0 = row.cells[0].paragraphs[0]
            r0 = c0.add_run(f"Q{q.order}")
            r0.font.bold = True; r0.font.size = Pt(12)
            r0.font.color.rgb = hex_rgb('#20B2AA')
            c1 = row.cells[1].paragraphs[0]
            r1 = c1.add_run(q.text); r1.font.size = Pt(11)
            c2 = row.cells[2].paragraphs[0]
            c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rm = c2.add_run(f"{q.max_marks} marks\n")
            rm.font.bold = True; rm.font.size = Pt(11)
            rm.font.color.rgb = hex_rgb('#f97316')
            rc = c2.add_run(f"CLO: {clo_codes}\n")
            rc.font.bold = True; rc.font.size = Pt(8)
            rc.font.color.rgb = hex_rgb('#1d4ed8')
            rp = c2.add_run(f"PLO: {plo_codes}")
            rp.font.bold = True; rp.font.size = Pt(8)
            rp.font.color.rgb = hex_rgb('#9d174d')

            for sq in q.sub_questions.order_by('order'):
                sq_clo = ', '.join(c.code for c in sq.clos.all()) or '—'
                sq_plo = ', '.join(p.code for p in sq.plos.all()) or '—'
                stbl = doc.add_table(rows=1, cols=3)
                stbl.style = 'Table Grid'
                stbl.autofit = False
                stbl.allow_autofit = False
                sr = stbl.rows[0]
                sr.cells[0].width = Cm(1.2)
                sr.cells[1].width = Cm(11.0)
                sr.cells[2].width = Cm(3.8)
                set_cell_bg(sr.cells[0], 'f0fafa')
                set_cell_bg(sr.cells[1], 'ffffff')
                set_cell_bg(sr.cells[2], 'ffffff')
                sr.cells[0].paragraphs[0].add_run(f"({sq.order})").font.size = Pt(9)
                st = sr.cells[1].paragraphs[0]
                st.add_run(sq.text).font.size = Pt(10)
                sm = sr.cells[2].paragraphs[0]
                sm.alignment = WD_ALIGN_PARAGRAPH.CENTER
                smr = sm.add_run(f"{sq.max_marks} marks\n")
                smr.font.bold = True; smr.font.size = Pt(9)
                smr.font.color.rgb = hex_rgb('#f97316')
                smrc = sm.add_run(f"CLO: {sq_clo}\n")
                smrc.font.bold = True; smrc.font.size = Pt(7)
                smrc.font.color.rgb = hex_rgb('#1d4ed8')
                smrp = sm.add_run(f"PLO: {sq_plo}")
                smrp.font.bold = True; smrp.font.size = Pt(7)
                smrp.font.color.rgb = hex_rgb('#9d174d')

            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_title = assessment.title.replace(' ', '_')
        response = HttpResponse(
            buf,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{safe_title}_{assessment.course.code}.docx"'
        )
        return response

    from django.http import Http404
    raise Http404("Invalid format")




# ── BATCH PLO ANALYTICS ────────────────────────────────────────────────────────

def _compute_batch_plo_attainment(batch_name):
    """
    Batch-level PLO attainment computation.
    Returns a dict with:
      - sections: list of section-level data
      - all_plos: sorted list of PLO codes
      - overall: overall attainment across all sections
    """
    SUB_TYPES = {'mid', 'final'}

    # All sections for this batch
    sections_qs = Section.objects.filter(batch=batch_name).select_related('course').order_by('name')

    all_plo_codes = set()
    section_results = []

    # Collect all courses in this batch (unique)
    courses_in_batch = Course.objects.filter(sections__batch=batch_name).distinct().order_by('code')

    # For each course, precompute PLO attainment logic once
    # Then we'll slice by section students

    for section in sections_qs:
        course = section.course
        section_students = list(section.students.all())
        student_ids = {s.id for s in section_students}

        if not student_ids:
            section_results.append({
                'section': section,
                'course': course,
                'plo_data': [],
                'overall_attainment': 0.0,
                'present_count': 0,
                'total_students': 0,
            })
            continue

        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related(
                'questions__clos__plos',
                'questions__plos',
                'questions__sub_questions__clos__plos',
                'questions__sub_questions__plos',
            )
            .order_by('assessment_type', 'created_at')
        )

        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({
                            'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                            'plo_ids': [p.id for p in sq.plos.all()],
                            'is_final': is_final,
                        })
                else:
                    all_columns.append({
                        'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                        'plo_ids': [p.id for p in q.plos.all()],
                        'is_final': is_final,
                    })

        plo_ids_used = set()
        for col in all_columns:
            plo_ids_used.update(col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used).order_by('code'))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}

        q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(
                question_id__in=q_ids,
                submission__assessment__course=course,
                submission__student_id__in=student_ids,
            ).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(
                sub_question_id__in=sq_ids,
                submission__assessment__course=course,
                submission__student_id__in=student_ids,
            ).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for sid in student_ids:
            sq_m = sq_grade_map.get(sid, {})
            q_m  = q_grade_map.get(sid, {})
            if has_final_cols:
                is_present = any(
                    col['entity_id'] in (sq_m if col['is_sub'] else q_m)
                    for col in all_columns if col['is_final']
                )
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(sid)
        present_count = len(present_student_ids)

        plo_achieved = {p.id: 0 for p in plos}
        for sid in student_ids:
            sq_m = sq_grade_map.get(sid, {})
            q_m  = q_grade_map.get(sid, {})
            if not sq_m and not q_m:
                continue
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(
                    (sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0))
                    for col in all_columns if p.id in col['plo_ids']
                )
                if round(raw / mx * 100, 1) >= 40:
                    plo_achieved[p.id] += 1

        plo_data = []
        for p in plos:
            achieved = plo_achieved[p.id]
            pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
            plo_data.append({
                'code': p.code,
                'description': p.description,
                'attainment': pct,
                'achieved': achieved,
                'present': present_count,
                'attained': pct >= 60,
            })
            all_plo_codes.add(p.code)

        overall_attainment = round(
            sum(d['attainment'] for d in plo_data) / len(plo_data), 1
        ) if plo_data else 0.0

        plo_map = {d['code']: d for d in plo_data}
        section_results.append({
            'section': section,
            'course': course,
            'plo_data': plo_data,
            'plo_map': plo_map,
            'overall_attainment': overall_attainment,
            'present_count': present_count,
            'total_students': len(student_ids),
        })

    # Overall across all sections
    all_attainments = [r['overall_attainment'] for r in section_results if r['present_count'] > 0]
    overall_batch_attainment = round(sum(all_attainments) / len(all_attainments), 1) if all_attainments else 0.0

    total_students_all = sum(r['total_students'] for r in section_results)
    total_present_all  = sum(r['present_count'] for r in section_results)

    return {
        'section_results': section_results,
        'all_plo_codes': sorted(all_plo_codes),
        'overall_batch_attainment': overall_batch_attainment,
        'total_students': total_students_all,
        'total_present': total_present_all,
        'batch_name': batch_name,
    }


@dao_required
def dao_batch_analytics(request):
    all_batches = list(
        Section.objects.values_list('batch', flat=True).distinct().order_by('batch')
    )
    selected_batch = request.GET.get('batch', '').strip()
    if not selected_batch and all_batches:
        selected_batch = all_batches[0]

    result = None
    if selected_batch:
        result = _compute_batch_plo_attainment(selected_batch)

    return render(request, 'dao_portal/batch_analytics.html', {
        'all_batches': all_batches,
        'selected_batch': selected_batch,
        'result': result,
    })


@dept_head_required
def dept_head_batch_analytics(request):
    all_batches = list(
        Section.objects.values_list('batch', flat=True).distinct().order_by('batch')
    )
    selected_batch = request.GET.get('batch', '').strip()
    if not selected_batch and all_batches:
        selected_batch = all_batches[0]

    result = None
    if selected_batch:
        result = _compute_batch_plo_attainment(selected_batch)

    return render(request, 'dept_head_portal/batch_analytics.html', {
        'all_batches': all_batches,
        'selected_batch': selected_batch,
        'result': result,
    })


@dept_head_required
def dept_head_plo_comparison(request):
    all_courses = Course.objects.all().order_by('code', 'semester')
    seen_codes = set()
    course_codes = []
    for c in all_courses:
        if c.code not in seen_codes:
            seen_codes.add(c.code)
            course_codes.append({'code': c.code, 'name': c.name})
    selected_code = request.GET.get('course_code', '').strip()
    selected_ids  = request.GET.getlist('sem_ids')
    courses_for_code = []
    if selected_code:
        courses_for_code = list(all_courses.filter(code=selected_code))
    elif course_codes:
        selected_code = course_codes[0]['code']
        courses_for_code = list(all_courses.filter(code=selected_code))
    if selected_ids:
        compare_set = set(selected_ids)
        courses_to_compare = [c for c in courses_for_code if str(c.id) in compare_set]
    else:
        courses_to_compare = list(courses_for_code)
    SUB_TYPES = {'mid', 'final'}
    comparison_data = []
    all_plo_meta = {}
    for course in courses_to_compare:
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related('questions__plos', 'questions__sub_questions__plos')
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'plo_ids': [p.id for p in sq.plos.all()], 'is_final': is_final})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'plo_ids': [p.id for p in q.plos.all()], 'is_final': is_final})
        if not all_columns:
            comparison_data.append({'course': course, 'semester': course.semester, 'plo_attainment': {}, 'present_count': 0})
            continue
        plo_ids_used = set(pid for col in all_columns for pid in col['plo_ids'])
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        for p in plos:
            all_plo_meta[p.code] = p.description
        q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=course).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=course).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained
        students = list(User.objects.filter(enrollments__course=course).distinct())
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(col['entity_id'] in (sq_m if col['is_sub'] else q_m) for col in all_columns if col['is_final'])
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)
        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue
            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if present_count > 0 and (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1
        plo_attainment = {p.code: (round(plo_achieved[p.id] / present_count * 100, 1) if present_count > 0 else 0.0) for p in plos}
        comparison_data.append({'course': course, 'semester': course.semester, 'plo_attainment': plo_attainment, 'present_count': present_count})
    def _sem_key(entry):
        parts = entry['semester'].split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))
    comparison_data.sort(key=_sem_key)
    plo_codes = sorted(all_plo_meta.keys())
    table_rows = []
    for entry in comparison_data:
        cells = [{'plo_code': code, 'attainment': entry['plo_attainment'].get(code),
                  'has_data': entry['plo_attainment'].get(code) is not None} for code in plo_codes]
        table_rows.append({'semester': entry['semester'], 'course_id': entry['course'].id,
                           'is_archived': entry['course'].is_archived, 'present_count': entry['present_count'], 'cells': cells})
    chart_json = json.dumps({
        'plo_codes': plo_codes, 'plo_descriptions': all_plo_meta,
        'semesters': [e['semester'] for e in comparison_data],
        'datasets': [{'semester': e['semester'], 'values': [e['plo_attainment'].get(code) for code in plo_codes]} for e in comparison_data],
    })
    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]
    return render(request, 'dept_head_portal/plo_comparison.html', {
        'course_codes': course_codes, 'selected_code': selected_code,
        'courses_for_code': courses_for_code, 'selected_ids': [int(i) for i in selected_ids if i.isdigit()],
        'plo_codes': plo_codes, 'plo_list': plo_list, 'table_rows': table_rows,
        'chart_json': chart_json, 'has_data': bool(plo_codes and table_rows),
    })


@dept_head_required
def dept_head_semester_plo_comparison(request):
    SUB_TYPES = {'mid', 'final'}
    all_courses = list(Course.objects.all().order_by('semester', 'code'))
    def _sem_key(s):
        parts = s.split()
        year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
        order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
        return (year, order.get(parts[0].lower() if parts else '', 5))
    all_semesters = sorted(set(c.semester for c in all_courses), key=_sem_key)
    selected_sems = request.GET.getlist('sems')
    if not selected_sems:
        selected_sems = list(all_semesters)
    selected_sems_set = set(selected_sems)
    semester_plo_data = {}
    all_plo_meta = {}
    semester_present_students = {}
    for course in all_courses:
        if course.semester not in selected_sems_set:
            continue
        assessments = list(
            Assessment.objects.filter(course=course, status='published')
            .prefetch_related('questions__plos', 'questions__sub_questions__plos')
            .order_by('assessment_type', 'created_at')
        )
        all_columns = []
        for a in assessments:
            use_subs = a.assessment_type in SUB_TYPES
            is_final = a.assessment_type == 'final'
            for q in a.questions.all().order_by('order'):
                subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                if use_subs and subs:
                    for sq in subs:
                        all_columns.append({'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                            'plo_ids': [p.id for p in sq.plos.all()], 'is_final': is_final})
                else:
                    all_columns.append({'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                                        'plo_ids': [p.id for p in q.plos.all()], 'is_final': is_final})
        if not all_columns:
            continue
        plo_ids_used = set(pid for col in all_columns for pid in col['plo_ids'])
        if not plo_ids_used:
            continue
        plos = list(PLO.objects.filter(id__in=plo_ids_used))
        plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}
        for p in plos:
            all_plo_meta[p.code] = p.description
        q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
        sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
        q_grade_map, sq_grade_map = {}, {}
        if q_ids:
            for g in QuestionGrade.objects.filter(question_id__in=q_ids, submission__assessment__course=course).select_related('submission__student'):
                q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
        if sq_ids:
            for g in SubQuestionGrade.objects.filter(sub_question_id__in=sq_ids, submission__assessment__course=course).select_related('submission__student'):
                sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained
        students = list(User.objects.filter(enrollments__course=course).distinct())
        has_final_cols = any(col['is_final'] for col in all_columns)
        present_student_ids = set()
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if has_final_cols:
                is_present = any(col['entity_id'] in (sq_m if col['is_sub'] else q_m) for col in all_columns if col['is_final'])
            else:
                is_present = bool(sq_m or q_m)
            if is_present:
                present_student_ids.add(student.id)
        present_count = len(present_student_ids)
        if present_count == 0:
            continue
        plo_achieved = {p.id: 0 for p in plos}
        for student in students:
            sq_m = sq_grade_map.get(student.id, {})
            q_m  = q_grade_map.get(student.id, {})
            if not sq_m and not q_m:
                continue
            def _mark(col, _sq=sq_m, _q=q_m):
                return _sq.get(col['entity_id'], 0) if col['is_sub'] else _q.get(col['entity_id'], 0)
            for p in plos:
                mx = plo_max[p.id]
                if mx == 0:
                    continue
                raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
                if (raw / mx * 100) >= 40:
                    plo_achieved[p.id] += 1
        semester = course.semester
        if semester not in semester_plo_data:
            semester_plo_data[semester] = {}
        semester_present_students.setdefault(semester, set()).update(present_student_ids)
        for p in plos:
            if plo_max[p.id] == 0:
                continue
            code = p.code
            if code not in semester_plo_data[semester]:
                semester_plo_data[semester][code] = {'description': p.description, 'total_achieved': 0, 'total_present': 0}
            semester_plo_data[semester][code]['total_achieved'] += plo_achieved[p.id]
            semester_plo_data[semester][code]['total_present'] += present_count
    plo_codes = sorted(all_plo_meta.keys())
    semesters_sorted = sorted(semester_plo_data.keys(), key=_sem_key)
    plo_list = [{'code': c, 'description': all_plo_meta[c]} for c in plo_codes]
    table_rows = []
    for sem in semesters_sorted:
        cells = []
        for code in plo_codes:
            pe = semester_plo_data.get(sem, {}).get(code)
            if pe and pe['total_present'] > 0:
                att = round(pe['total_achieved'] / pe['total_present'] * 100, 1)
                cells.append({'plo_code': code, 'attainment': att, 'has_data': True})
            else:
                cells.append({'plo_code': code, 'attainment': 0, 'has_data': False})
        table_rows.append({'semester': sem, 'present_count': len(semester_present_students.get(sem, set())), 'cells': cells})
    chart_json = json.dumps({
        'plo_codes': plo_codes, 'semesters': semesters_sorted,
        'datasets': [
            {'plo': code, 'values': [
                (round(semester_plo_data[sem][code]['total_achieved'] / semester_plo_data[sem][code]['total_present'] * 100, 1)
                 if (sem in semester_plo_data and code in semester_plo_data[sem] and semester_plo_data[sem][code]['total_present'] > 0)
                 else None)
                for sem in semesters_sorted
            ]}
            for code in plo_codes
        ],
    })
    return render(request, 'dept_head_portal/semester_plo_comparison.html', {
        'all_semesters': all_semesters, 'selected_sems': selected_sems,
        'plo_codes': plo_codes, 'plo_list': plo_list, 'table_rows': table_rows,
        'semesters': semesters_sorted, 'chart_json': chart_json,
        'has_data': bool(plo_codes and semesters_sorted),
    })


# ── FACULTY BATCH PLO ANALYTICS ───────────────────────────────────────────────


def faculty_batch_analytics(request):
    """Faculty: batch-wise PLO attainment for their own assigned sections."""
    if not request.user.is_authenticated or request.user.role != 'faculty':
        return redirect('sign_in_html')

    SUB_TYPES = {'mid', 'final'}

    # Only sections assigned to this faculty
    my_sections = Section.objects.filter(faculty=request.user).select_related('course').order_by('batch', 'name')

    # All batches this faculty has
    all_batches = list(my_sections.values_list('batch', flat=True).distinct().order_by('batch'))

    selected_batch = request.GET.get('batch', '').strip()
    if not selected_batch and all_batches:
        selected_batch = all_batches[0]

    result = None
    if selected_batch:
        sections_qs = my_sections.filter(batch=selected_batch)

        all_plo_codes = set()
        section_results = []

        for section in sections_qs:
            course = section.course
            # Use Enrollment (same as faculty_analytics) — section.students ManyToMany may be empty
            section_student_ids_mm = set(section.students.values_list('id', flat=True))
            enrolled_ids = set(
                User.objects.filter(enrollments__course=course).values_list('id', flat=True)
            )
            # If section has students assigned via ManyToMany use that, else fall back to all enrolled
            student_ids = section_student_ids_mm if section_student_ids_mm else enrolled_ids

            if not student_ids:
                section_results.append({
                    'section': section,
                    'course': course,
                    'plo_data': [],
                    'plo_map': {},
                    'overall_attainment': 0.0,
                    'present_count': 0,
                    'total_students': 0,
                })
                continue

            assessments = list(
                Assessment.objects.filter(course=course, status='published')
                .prefetch_related(
                    'questions__plos',
                    'questions__sub_questions__plos',
                )
                .order_by('assessment_type', 'created_at')
            )

            all_columns = []
            for a in assessments:
                use_subs = a.assessment_type in SUB_TYPES
                is_final = a.assessment_type == 'final'
                for q in a.questions.all().order_by('order'):
                    subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
                    if use_subs and subs:
                        for sq in subs:
                            all_columns.append({
                                'is_sub': True, 'entity_id': sq.id, 'max_marks': sq.max_marks,
                                'plo_ids': [p.id for p in sq.plos.all()],
                                'is_final': is_final,
                            })
                    else:
                        all_columns.append({
                            'is_sub': False, 'entity_id': q.id, 'max_marks': q.max_marks,
                            'plo_ids': [p.id for p in q.plos.all()],
                            'is_final': is_final,
                        })

            plo_ids_used = set()
            for col in all_columns:
                plo_ids_used.update(col['plo_ids'])
            plos = list(PLO.objects.filter(id__in=plo_ids_used).order_by('code'))
            plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}

            q_ids  = [col['entity_id'] for col in all_columns if not col['is_sub']]
            sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
            q_grade_map, sq_grade_map = {}, {}
            if q_ids:
                for g in QuestionGrade.objects.filter(
                    question_id__in=q_ids,
                    submission__assessment__course=course,
                    submission__student_id__in=student_ids,
                ).select_related('submission__student'):
                    q_grade_map.setdefault(g.submission.student_id, {})[g.question_id] = g.marks_obtained
            if sq_ids:
                for g in SubQuestionGrade.objects.filter(
                    sub_question_id__in=sq_ids,
                    submission__assessment__course=course,
                    submission__student_id__in=student_ids,
                ).select_related('submission__student'):
                    sq_grade_map.setdefault(g.submission.student_id, {})[g.sub_question_id] = g.marks_obtained

            has_final_cols = any(col['is_final'] for col in all_columns)
            present_student_ids = set()
            for sid in student_ids:
                sq_m = sq_grade_map.get(sid, {})
                q_m  = q_grade_map.get(sid, {})
                if has_final_cols:
                    is_present = any(
                        col['entity_id'] in (sq_m if col['is_sub'] else q_m)
                        for col in all_columns if col['is_final']
                    )
                else:
                    is_present = bool(sq_m or q_m)
                if is_present:
                    present_student_ids.add(sid)
            present_count = len(present_student_ids)

            plo_achieved = {p.id: 0 for p in plos}
            for sid in student_ids:
                sq_m = sq_grade_map.get(sid, {})
                q_m  = q_grade_map.get(sid, {})
                if not sq_m and not q_m:
                    continue
                for p in plos:
                    mx = plo_max[p.id]
                    if mx == 0:
                        continue
                    raw = sum(
                        (sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0))
                        for col in all_columns if p.id in col['plo_ids']
                    )
                    if round(raw / mx * 100, 1) >= 40:
                        plo_achieved[p.id] += 1

            plo_data = []
            for p in plos:
                achieved = plo_achieved[p.id]
                pct = round(achieved / present_count * 100, 1) if present_count > 0 else 0.0
                plo_data.append({
                    'code': p.code,
                    'description': p.description,
                    'attainment': pct,
                    'achieved': achieved,
                    'present': present_count,
                    'attained': pct >= 60,
                })
                all_plo_codes.add(p.code)

            plo_map = {d['code']: d for d in plo_data}
            overall_attainment = round(
                sum(d['attainment'] for d in plo_data) / len(plo_data), 1
            ) if plo_data else 0.0

            section_results.append({
                'section': section,
                'course': course,
                'plo_data': plo_data,
                'plo_map': plo_map,
                'overall_attainment': overall_attainment,
                'present_count': present_count,
                'total_students': len(student_ids),
            })

        all_attainments = [r['overall_attainment'] for r in section_results if r['present_count'] > 0]
        overall_batch_attainment = round(sum(all_attainments) / len(all_attainments), 1) if all_attainments else 0.0

        result = {
            'section_results': section_results,
            'all_plo_codes': sorted(all_plo_codes),
            'overall_batch_attainment': overall_batch_attainment,
            'total_students': sum(r['total_students'] for r in section_results),
            'total_present': sum(r['present_count'] for r in section_results),
            'batch_name': selected_batch,
        }

    return render(request, 'faculty/batch_analytics.html', {
        'all_batches': all_batches,
        'selected_batch': selected_batch,
        'result': result,
    })


# ── PLO Track ──────────────────────────────────────────────────────────────

def _sem_sort_key(s):
    parts = s.split()
    year = int(parts[-1]) if len(parts) > 1 and parts[-1].isdigit() else 0
    term_order = {'spring': 1, 'summer': 2, 'fall': 3, 'winter': 0}
    return (year, term_order.get(parts[0].lower() if parts else '', 5))


def _compute_plo_attainments_for_student(student, course):
    SUB_TYPES = {'mid', 'final'}
    assessments = list(
        Assessment.objects.filter(course=course, status='published')
        .prefetch_related(
            'questions__clos', 'questions__plos',
            'questions__sub_questions__clos', 'questions__sub_questions__plos',
        )
        .order_by('assessment_type', 'created_at')
    )
    if not assessments:
        return [], False

    all_columns = []
    for a in assessments:
        use_subs = a.assessment_type in SUB_TYPES
        for q in a.questions.all().order_by('order'):
            subs = list(q.sub_questions.all().order_by('order')) if use_subs else []
            if use_subs and subs:
                for sq in subs:
                    all_columns.append({
                        'is_sub': True, 'entity_id': sq.id,
                        'max_marks': sq.max_marks,
                        'plo_ids': [p.id for p in sq.plos.all()],
                    })
            else:
                all_columns.append({
                    'is_sub': False, 'entity_id': q.id,
                    'max_marks': q.max_marks,
                    'plo_ids': [p.id for p in q.plos.all()],
                })

    plo_ids_used = set()
    for col in all_columns:
        plo_ids_used.update(col['plo_ids'])
    if not plo_ids_used:
        return [], False

    plos = list(PLO.objects.filter(id__in=plo_ids_used))
    plo_max = {p.id: sum(col['max_marks'] for col in all_columns if p.id in col['plo_ids']) for p in plos}

    q_ids = [col['entity_id'] for col in all_columns if not col['is_sub']]
    sq_ids = [col['entity_id'] for col in all_columns if col['is_sub']]
    q_m, sq_m = {}, {}
    if q_ids:
        for g in QuestionGrade.objects.filter(
            question_id__in=q_ids,
            submission__assessment__course=course,
            submission__student=student,
        ):
            q_m[g.question_id] = g.marks_obtained
    if sq_ids:
        for g in SubQuestionGrade.objects.filter(
            sub_question_id__in=sq_ids,
            submission__assessment__course=course,
            submission__student=student,
        ):
            sq_m[g.sub_question_id] = g.marks_obtained

    has_grades = bool(q_m or sq_m)

    def _mark(col):
        return sq_m.get(col['entity_id'], 0) if col['is_sub'] else q_m.get(col['entity_id'], 0)

    results = []
    for p in plos:
        mx = plo_max[p.id]
        if mx == 0:
            continue
        raw = sum(_mark(col) for col in all_columns if p.id in col['plo_ids'])
        att = round(raw / mx * 100, 1)
        results.append({'plo_id': p.id, 'code': p.code, 'description': p.description, 'attainment': att})
    return results, has_grades


@student_required
def student_plo_track(request):
    student = request.user
    enrollments = Enrollment.objects.filter(student=student).select_related('course')

    raw_records = []
    all_courses = []

    for e in enrollments:
        course = e.course
        all_courses.append(course)
        plo_results, has_grades = _compute_plo_attainments_for_student(student, course)
        if not has_grades:
            continue
        for p in plo_results:
            raw_records.append({
                'course': course,
                'course_id': course.id,
                'course_code': course.code,
                'course_name': course.name,
                'semester': course.semester,
                'plo_id': p['plo_id'],
                'plo_code': p['code'],
                'plo_desc': p['description'],
                'attainment': p['attainment'],
            })

    sorted_records = sorted(raw_records, key=lambda r: (_sem_sort_key(r['semester']), r['course_code']))

    # Build per-PLO trend data
    plo_info = {}
    for r in sorted_records:
        plo_info[r['plo_id']] = {'code': r['plo_code'], 'description': r['plo_desc']}

    plo_trends = []
    for plo_id, info in sorted(plo_info.items(), key=lambda x: x[1]['code']):
        records = [r for r in sorted_records if r['plo_id'] == plo_id]
        prev_att = None
        records_with_trend = []
        for r in records:
            if prev_att is None:
                trend = 'first'
            elif r['attainment'] > prev_att + 2:
                trend = 'up'
            elif r['attainment'] < prev_att - 2:
                trend = 'down'
            else:
                trend = 'stable'
            records_with_trend.append({**r, 'trend': trend})
            prev_att = r['attainment']

        if len(records) >= 2:
            if records[-1]['attainment'] > records[0]['attainment'] + 2:
                overall = 'up'
            elif records[-1]['attainment'] < records[0]['attainment'] - 2:
                overall = 'down'
            else:
                overall = 'stable'
        else:
            overall = 'first'

        plo_trends.append({
            'plo_id': plo_id,
            'code': info['code'],
            'description': info['description'],
            'records': records_with_trend,
            'overall_trend': overall,
            'min_att': min(r['attainment'] for r in records) if records else 0,
            'max_att': max(r['attainment'] for r in records) if records else 0,
            'latest_att': records[-1]['attainment'] if records else 0,
            'is_weak': (records[-1]['attainment'] if records else 0) < 40,
        })

    all_semesters = sorted(set(r['semester'] for r in raw_records), key=_sem_sort_key)
    unique_courses = list({c.id: c for c in all_courses}.values())
    all_plo_opts = sorted(
        [{'id': pid, 'code': info['code'], 'description': info['description']}
         for pid, info in plo_info.items()],
        key=lambda x: x['code'],
    )

    # Per-course chart data
    course_chart_data = []
    for course in unique_courses:
        recs = sorted([r for r in sorted_records if r['course_id'] == course.id], key=lambda r: r['plo_code'])
        if not recs:
            continue
        course_chart_data.append({
            'course': course,
            'plo_codes': [r['plo_code'] for r in recs],
            'attainments': [r['attainment'] for r in recs],
            'colors': [
                '#10b981' if r['attainment'] >= 80
                else '#20b2aa' if r['attainment'] >= 60
                else '#f59e0b' if r['attainment'] >= 40
                else '#ef4444'
                for r in recs
            ],
        })

    weak_plos = [p for p in plo_trends if p['is_weak']]

    return render(request, 'student/plo_track.html', {
        'plo_trends': plo_trends,
        'sorted_records': sorted_records,
        'all_semesters': all_semesters,
        'all_courses': unique_courses,
        'all_plo_opts': all_plo_opts,
        'course_chart_data': course_chart_data,
        'weak_plos': weak_plos,
    })


@faculty_required
def faculty_plo_track(request):
    faculty = request.user
    courses = list(Course.objects.filter(faculty=faculty, is_archived=False).order_by('-created_at'))

    selected_course_id = request.GET.get('course_id')
    selected_course = None
    if selected_course_id:
        try:
            selected_course = Course.objects.get(id=int(selected_course_id), faculty=faculty)
        except (Course.DoesNotExist, ValueError):
            pass
    if selected_course is None and courses:
        selected_course = courses[0]

    student_plo_data = []
    plos_for_template = []

    if selected_course:
        enrollments = Enrollment.objects.filter(course=selected_course).select_related('student')
        students = [e.student for e in enrollments if e.student.role == 'student']

        plo_id_set = set()
        student_results = []
        for student in students:
            plo_results, has_grades = _compute_plo_attainments_for_student(student, selected_course)
            for p in plo_results:
                plo_id_set.add(p['plo_id'])
            student_results.append((student, plo_results, has_grades))

        all_plos = list(PLO.objects.filter(id__in=plo_id_set).order_by('code'))
        plos_for_template = [{'id': p.id, 'code': p.code, 'description': p.description} for p in all_plos]

        for student, plo_results, has_grades in student_results:
            att_map = {p['plo_id']: p for p in plo_results}
            ordered = []
            for plo in all_plos:
                entry = att_map.get(plo.id, {'plo_id': plo.id, 'code': plo.code, 'description': plo.description, 'attainment': 0.0})
                entry['is_weak'] = entry['attainment'] < 40
                ordered.append(entry)
            weak_plos = [p for p in ordered if p['is_weak'] and has_grades]
            student_plo_data.append({
                'student': student,
                'plo_attainments': ordered,
                'weak_plos': weak_plos,
                'has_grades': has_grades,
                'avg_attainment': round(sum(p['attainment'] for p in ordered) / len(ordered), 1) if ordered else 0.0,
            })

    return render(request, 'faculty/plo_track.html', {
        'courses': courses,
        'selected_course': selected_course,
        'student_plo_data': student_plo_data,
        'plos': plos_for_template,
    })


@faculty_required
def faculty_send_plo_notification(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, Exception):
        return JsonResponse({'error': 'Invalid request body'}, status=400)

    student_id = data.get('student_id')
    course_id  = data.get('course_id')
    plo_code   = data.get('plo_code', '')
    message    = data.get('message', '').strip()
    notif_type = data.get('notif_type', 'plo_feedback')

    if notif_type not in ('plo_feedback', 'plo_alert'):
        notif_type = 'plo_feedback'

    if not all([student_id, course_id, plo_code, message]):
        return JsonResponse({'error': 'Please fill in all fields (PLO and message are required).'}, status=400)

    try:
        student = User.objects.get(id=int(student_id), role='student')
    except (User.DoesNotExist, ValueError, TypeError):
        return JsonResponse({'error': 'Student not found.'}, status=404)

    try:
        course = Course.objects.filter(id=int(course_id), faculty=request.user).first()
        if course is None:
            return JsonResponse({'error': 'Course not found or access denied.'}, status=404)
    except (ValueError, TypeError):
        return JsonResponse({'error': 'Invalid course.'}, status=400)

    # Verify student is enrolled in this course
    if not Enrollment.objects.filter(student=student, course=course).exists():
        return JsonResponse({'error': 'This student is not enrolled in the selected course.'}, status=403)

    title = f"{'PLO Alert' if notif_type == 'plo_alert' else 'PLO Feedback'}: {plo_code} — {course.code}"
    full_message = (
        f"Course: {course.code} - {course.name}\n"
        f"PLO: {plo_code}\n\n"
        f"From {request.user.full_name or request.user.username}:\n{message}"
    )

    try:
        Notification.objects.create(
            recipient=student,
            notif_type=notif_type,
            title=title,
            message=full_message,
            course=course,
        )
    except Exception as e:
        return JsonResponse({'error': f'Failed to create notification: {str(e)}'}, status=500)

    return JsonResponse({'success': True})